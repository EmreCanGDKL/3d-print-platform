# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_thesis as bt  # noqa: E402


WORKSPACE = Path(__file__).resolve().parents[1]
OUT_PATH = WORKSPACE / "PrintForge_Bitirme_Tezi_40_50_Sayfa.docx"
SCREENSHOT_DIR = WORKSPACE / "thesis_work" / "site_screenshots"


PAGE_MAP = {
    "ozet": "v",
    "abstract": "vi",
    "onsoz": "vii",
    "kisaltmalar": "ix",
    "tablolar": "x",
    "sekiller": "xi",
    "ekler_listesi": "xii",
    "giris": "1",
    "ch1": "3",
    "s11": "3",
    "s12": "5",
    "s13": "7",
    "s14": "9",
    "s15": "11",
    "ch2": "13",
    "s21": "13",
    "s22": "15",
    "s23": "17",
    "s24": "19",
    "s25": "21",
    "s26": "23",
    "s27": "25",
    "s28": "32",
    "s29": "34",
    "s210": "39",
    "ch3": "36",
    "s31": "36",
    "s32": "38",
    "s33": "43",
    "s34": "44",
    "s35": "45",
    "s36": "46",
    "s37": "47",
    "sonuc": "48",
    "kaynakca": "49",
    "ekler": "50",
    "ozgecmis": "50",
    "tab_1_1": "11",
    "tab_2_1": "14",
    "tab_2_2": "16",
    "tab_2_3": "18",
    "tab_2_4": "22",
    "tab_2_5": "31",
    "tab_2_6": "33",
    "tab_2_7": "38",
    "tab_2_8": "42",
    "tab_3_1": "37",
    "tab_3_2": "44",
    "tab_3_3": "45",
    "tab_3_4": "47",
    "fig_1_1": "6",
    "fig_2_1": "15",
    "fig_2_2": "16",
    "fig_2_3": "18",
    "fig_2_4": "23",
    "fig_3_1": "37",
    "fig_3_2": "38",
    "fig_3_3": "39",
    "fig_3_4": "40",
    "fig_3_5": "41",
    "fig_3_6": "42",
    "fig_3_7": "43",
    "fig_3_8": "43",
    "fig_3_9": "44",
    "fig_3_10": "44",
    "fig_3_11": "45",
    "fig_3_12": "46",
    "app_a": "50",
    "app_b": "50",
    "app_c": "50",
    "app_d": "50",
}


def configure_lists() -> None:
    bt.DOCX_PATH = OUT_PATH
    bt.TOC_ENTRIES = [
        ("ÖZET", "ozet", 0),
        ("ABSTRACT", "abstract", 0),
        ("ÖNSÖZ", "onsoz", 0),
        ("KISALTMALAR", "kisaltmalar", 0),
        ("TABLOLAR LİSTESİ", "tablolar", 0),
        ("ŞEKİLLER LİSTESİ", "sekiller", 0),
        ("EKLER LİSTESİ", "ekler_listesi", 0),
        ("GİRİŞ", "giris", 0),
        ("BİRİNCİ BÖLÜM TEORİK ALT YAPI", "ch1", 0),
        ("1.1. Eklemeli İmalat ve 3D Baskı", "s11", 1),
        ("1.2. AI Destekli 3D Model Üretimi", "s12", 1),
        ("1.3. Web Tabanlı Pazaryeri Yaklaşımı", "s13", 1),
        ("1.4. Yazılım Mimarisi, Güvenlik ve Veri Yönetimi", "s14", 1),
        ("1.5. Literatür Değerlendirmesi ve Özgün Değer", "s15", 1),
        ("İKİNCİ BÖLÜM MATERYAL VE METOT", "ch2", 0),
        ("2.1. Kullanılan Materyaller ve Yazılım Araçları", "s21", 1),
        ("2.2. Sistem Mimarisi", "s22", 1),
        ("2.3. Veri Modeli ve Veritabanı Tasarımı", "s23", 1),
        ("2.4. Uygulama Geliştirme Süreci", "s24", 1),
        ("2.5. Yapılan Testler", "s25", 1),
        ("2.6. Güvenlik, Etik Standartlar ve Varsayımlar", "s26", 1),
        ("2.7. Ayrıntılı Modül Tasarımı", "s27", 1),
        ("2.8. İş Paketleri ve Zaman Yönetimi", "s28", 1),
        ("2.9. Senaryo Tabanlı Sistem Tasarımı", "s29", 1),
        ("2.10. Kalite, Güvenlik ve Sürdürülebilirlik Değerlendirmesi", "s210", 1),
        ("ÜÇÜNCÜ BÖLÜM BULGULAR VE TARTIŞMA", "ch3", 0),
        ("3.1. Gerçekleştirilen Modüller", "s31", 1),
        ("3.2. Arayüz ve Kullanıcı Akışı Bulguları", "s32", 1),
        ("3.3. Derleme ve Teknik Doğrulama Bulguları", "s33", 1),
        ("3.4. Kullanıcı Akışlarına İlişkin Bulgular", "s34", 1),
        ("3.5. Güvenlik, Sınırlılıklar ve Geliştirme Önerileri", "s35", 1),
        ("3.6. Tartışma", "s36", 1),
        ("3.7. Demo ve Kabul Değerlendirmesi", "s37", 1),
        ("SONUÇ VE ÖNERİLER", "sonuc", 0),
        ("KAYNAKÇA", "kaynakca", 0),
        ("EKLER", "ekler", 0),
        ("ÖZGEÇMİŞ", "ozgecmis", 0),
    ]
    bt.TABLES = [
        ("Tablo 1.1. Literatür çalışmaları ve PrintForge ile ilişkisi", "tab_1_1"),
        ("Tablo 2.1. Projede kullanılan yazılım bileşenleri", "tab_2_1"),
        ("Tablo 2.2. Backend API modülleri ve görevleri", "tab_2_2"),
        ("Tablo 2.3. Temel veritabanı varlıkları", "tab_2_3"),
        ("Tablo 2.4. Test senaryoları ve kabul ölçütleri", "tab_2_4"),
        ("Tablo 2.5. Ayrıntılı mimari katmanlar ve kontrol noktaları", "tab_2_5"),
        ("Tablo 2.6. İş paketleri ve çıktıları", "tab_2_6"),
        ("Tablo 2.7. Senaryo bazlı kabul matrisi", "tab_2_7"),
        ("Tablo 2.8. Teknik kalite ve risk kontrol matrisi", "tab_2_8"),
        ("Tablo 3.1. Derleme doğrulama sonuçları", "tab_3_1"),
        ("Tablo 3.2. Gerçekleştirilen kullanıcı akışları", "tab_3_2"),
        ("Tablo 3.3. Sınırlılıklar ve geliştirme önerileri", "tab_3_3"),
        ("Tablo 3.4. Savunma demosu ve kanıt matrisi", "tab_3_4"),
    ]
    bt.FIGURES = [
        ("Şekil 1.1. Literatürden projeye uzanan kavramsal konumlandırma", "fig_1_1"),
        ("Şekil 2.1. PrintForge sistem mimarisi", "fig_2_1"),
        ("Şekil 2.2. Referans görselden AI üretime veri akışı", "fig_2_2"),
        ("Şekil 2.3. Veri modeli ilişkileri", "fig_2_3"),
        ("Şekil 2.4. Kimlik doğrulama ve güvenli erişim akışı", "fig_2_4"),
        ("Şekil 3.1. Modül tamamlama durumu", "fig_3_1"),
        ("Şekil 3.2. Ana sayfa ve değer önerisi ekranı", "fig_3_2"),
        ("Şekil 3.3. Katalog ve filtreleme ekranı", "fig_3_3"),
        ("Şekil 3.4. Örnek görsel arama ekranı", "fig_3_4"),
        ("Şekil 3.5. AI model üretim ekranı", "fig_3_5"),
        ("Şekil 3.6. Giriş ekranı", "fig_3_6"),
        ("Şekil 3.7. Satıcı ürün yönetimi ekranı", "fig_3_7"),
        ("Şekil 3.8. Satıcı ürün ekleme ekranı", "fig_3_8"),
        ("Şekil 3.9. Favoriler ekranı", "fig_3_9"),
        ("Şekil 3.10. Mesajlar ekranı", "fig_3_10"),
        ("Şekil 3.11. Sipariş takip ekranı", "fig_3_11"),
        ("Şekil 3.12. Derleme ve doğrulama özeti", "fig_3_12"),
    ]
    bt.APPENDICES = [
        ("Ek A. Kurulum ve çalıştırma adımları", "app_a"),
        ("Ek B. Backend ortam değişkenleri", "app_b"),
        ("Ek C. API uç noktaları özeti", "app_c"),
        ("Ek D. Manuel test kontrol listesi", "app_d"),
    ]


def add_screen_figure(doc: Document, filename: str, caption: str, note: str, finding: str) -> None:
    path = SCREENSHOT_DIR / filename
    bt.add_para(doc, note)
    bt.add_figure(doc, path, caption, 12.8)
    bt.add_para(
        doc,
        "Kaynak: PrintForge uygulamasından alınan ekran görüntüsü, 06.06.2026.",
        "TezSource",
    )
    bt.add_para(doc, finding)


def add_extended_technical_design_compact(doc: Document) -> None:
    bt.add_section_heading(doc, "2.7. Ayrıntılı Modül Tasarımı")
    bt.add_para(
        doc,
        "Bu alt bölüm, PrintForge uygulamasının yalnızca hangi teknolojilerle geliştirildiğini değil, modüllerin hangi sorumluluklarla ayrıldığını ve bu ayrımın neden tercih edildiğini açıklamak amacıyla hazırlanmıştır. Projede her ekran bağımsız bir sayfa olarak değil, fikir seçimi, AI üretimi, katalog inceleme, satıcıyla iletişim ve sipariş takibi zincirinin bir parçası olarak ele alınmıştır. Bu nedenle modül tasarımında kullanıcı deneyimi, veri modeli, güvenlik ve sürdürülebilirlik kararları birlikte değerlendirilmiştir.",
    )
    bt.add_para(
        doc,
        "Modül ayrımı yapılırken iki temel ölçüt kullanılmıştır. İlk ölçüt, kullanıcının ekranda tamamladığı görevin sınırıdır; örneğin örnek görsel arama, AI model üretme, katalog ürünü inceleme veya satıcıya mesaj gönderme ayrı kullanıcı niyetleridir. İkinci ölçüt ise backend tarafında yürütülen iş kuralıdır; kimlik doğrulama, görsel arama, model dosyası erişimi ve mesajlaşma gibi işlemler ayrı rota dosyaları ve doğrulama adımlarıyla yönetilmiştir. Bu yaklaşım kodun okunabilirliğini artırmış, hata ayıklama ve geliştirme sürecini daha kontrollü hale getirmiştir.",
    )
    sections = [
        (
            "2.7.1. Frontend Sayfa Yapısı ve Kullanıcı Deneyimi",
            [
                "Frontend katmanında ana sayfa, örnekler, AI üretim stüdyosu, pazaryeri, sepet, favoriler, mesajlar, siparişler, satıcı ürün paneli ve admin örnek yönetimi gibi sayfalar bulunmaktadır. Bu sayfaların her biri farklı kullanıcı rolünün bir görevi tamamlamasına hizmet eder. Kullanıcı için en kritik akış örnek görselden AI üretim sayfasına geçiştir; satıcı için ürün ekleme ve gelen mesaj/sipariş takibidir; admin için ise platformdaki referans içeriklerinin düzenli tutulmasıdır.",
                "Arayüz tasarımında rol tabanlı görünürlük tercih edilmiştir. Satıcı hesabı ürün ekleme ve ürün yönetme bağlantılarını görürken, normal kullanıcı katalog inceleme, AI üretim ve satıcıya ulaşma akışlarına yönlendirilir. Bu yaklaşım kullanıcıyı ilgisiz kontrollerle karşılaştırmadığı için hata olasılığını azaltır ve arayüzü daha anlaşılır hale getirir.",
            ],
        ),
        (
            "2.7.2. Örnek Görsel ve Prompt Aktarım Mekanizması",
            [
                "Örnekler sayfasında iki farklı referans kaynağı vardır: admin tarafından yönetilen kontrollü örnekler ve internet görsel arama sonuçları. Kontrollü örnekler, dış servis kotası veya arama kalitesi sorunlarında sistemin temel akışını sürdürebilmesini sağlar. İnternet araması ise kullanıcının daha geniş fikir havuzuna ulaşmasına yardımcı olur.",
                "Kullanıcı bir görsel kartını seçtiğinde yalnızca görsel URL'si aktarılmaz; başlık, kaynak, arama kelimesi ve oluşturulan prompt da query parametreleriyle AI üretim sayfasına taşınır. Bu karar, kullanıcının boş bir prompt alanıyla karşılaşmasını önler ve referans görselin üretim bağlamını korur.",
            ],
        ),
        (
            "2.7.3. AI Üretim Görevi ve Durum İzleme",
            [
                "AI üretim servisleri genellikle anlık sonuç döndürmek yerine görev tabanlı çalışır. Bu nedenle PrintForge, model üretimini tek istek-tek cevap şeklinde değil, görev başlatma ve görev durumunu izleme şeklinde tasarlamıştır. Backend üretim isteğini aldıktan sonra harici servisten taskId alır ve bunu Model kaydıyla ilişkilendirir.",
                "Frontend tarafındaki üretim durumu yönetimi, aktif üretim görevini kullanıcının oturum deneyimi içinde görünür tutar. Bu sayede kullanıcı sayfadan ayrılsa veya katalogda gezinmeye devam etse bile üretim görevi tamamen unutulmaz. Bu yaklaşım özellikle 3D model üretiminin saniyeler değil dakikalar sürebileceği durumlarda önemlidir.",
            ],
        ),
        (
            "2.7.4. Backend Doğrulama ve Hata Yönetimi",
            [
                "Backend katmanında gelen isteklerin biçimsel doğruluğu Zod şemalarıyla denetlenmiştir. Kayıt, giriş, profil güncelleme, ürün ekleme, yorum, soru, cevap, mesaj, sipariş ve görsel arama isteklerinin her biri beklenen tip, uzunluk ve zorunlu alanlara göre kontrol edilir. Bu yaklaşım hatalı verinin veritabanına ulaşmasını engeller.",
                "Hata yönetiminde kullanıcıya teknik ayrıntı yerine anlaşılır mesaj döndürme yaklaşımı kullanılmıştır. Görsel arama kotası, geçersiz API anahtarı, üretim servisinin yanıt vermemesi veya dosya formatı uyuşmazlığı gibi durumlar ayrı ayrı ele alınmıştır. Bu tasarım sistemin hata anında sessiz kalmasını önler.",
            ],
        ),
        (
            "2.7.5. Model Dosyası ve 3D Önizleme Tasarımı",
            [
                "AI üretim görevi başarıyla tamamlandığında backend model dosyasını indirir ve dosya türünü kontrol eder. GLB formatındaki modeller doğrudan önizlenebilir dosya olarak saklanırken, STL formatındaki modeller için güvenli önizleme verisi üretilebilir. Bu ayrım, farklı servislerden gelen çıktı formatlarının aynı uygulamada yönetilebilmesi için gereklidir.",
                "3D önizleme bileşeni Three.js ekosistemi üzerine kurulmuştur. ModelViewer bileşeni modeli yükler, merkezler, kamera sınırlarını ayarlar ve kullanıcıya döndürme/inceleme olanağı verir. Model dosyalarına erişim ise güvenli uç nokta üzerinden yapılır; kullanıcı yalnızca kendi modeline veya tarafı olduğu konuşmaya bağlı modele erişebilir.",
            ],
        ),
        (
            "2.7.6. Pazaryeri ve Satıcı Ürün Yönetimi",
            [
                "Pazaryeri modülü, AI üretimden bağımsız olarak çalışan hazır ürün katalogunu temsil eder. Satıcı; ürün adı, açıklama, kategori, fiyat ve görseller ile katalog ürünü oluşturur. Kullanıcılar bu ürünleri arayabilir, kategori ve maksimum fiyat filtreleriyle daraltabilir, ürün görsellerini inceleyebilir ve satıcıya mesaj gönderebilir.",
                "Ürün yönetiminde pasif hale getirme yaklaşımı tercih edilmiştir. Satıcı bir ürünü kaldırdığında kayıt tamamen yok edilmek yerine görünürlükten çıkarılabilir. Bu yaklaşım, ileride sipariş geçmişi, konuşma kayıtları ve raporlama ihtiyaçları için veri bütünlüğünün korunmasına yardımcı olur.",
            ],
        ),
        (
            "2.7.7. Mesajlaşma ve Sipariş Durum Yaşam Döngüsü",
            [
                "Mesajlaşma modülü kullanıcı ile satıcı arasında ürün veya AI modeli bağlamında konuşma oluşturur. Katalog ürünleri için konuşma satıcının ürün sahibi olması üzerinden açılırken, AI modeli için kullanıcı mesaj göndermek istediği satıcıyı seçebilir. Bu fark, hazır ürün alımı ile özel üretim talebi arasındaki iş akışı farkını yansıtır.",
                "Konuşma kayıtlarında buyerId, sellerId, modelId, modelType ve status alanları bulunur. Status alanı yalnızca mesajlaşma durumunu değil, sipariş yaşam döngüsünü de temsil eder. ORDERED, PREPARING, SHIPPED, COMPLETED ve CANCELLED değerleri alıcı ve satıcı arasında sipariş takibinin izlenebilir bir süreç olarak yürütülmesini sağlar.",
            ],
        ),
        (
            "2.7.8. Dağıtım, Ortam Değişkenleri ve Sürdürme Yaklaşımı",
            [
                "Frontend dağıtımı için Netlify yapılandırması hazırlanmıştır. Netlify yalnızca frontend uygulamasını yayınladığı için backend'in ayrı bir web servisi olarak çalışması gerekir. Bu ayrım README içinde açık biçimde belirtilmiş ve canlı ortamda localhost adreslerinin kullanılmaması gerektiği vurgulanmıştır.",
                "Backend tarafında DATABASE_URL, JWT_SECRET, FRONTEND_URLS, SERPAPI_API_KEY, TRIPO_API_KEY ve HITEM3D anahtarları gibi değerler ortam değişkenleriyle yönetilir. Böylece gizli bilgiler kaynak koduna yazılmaz. Modüler rota yapısı ise Auth, AI, models, chat, examples ve images alanlarının ayrı dosyalarda sürdürülebilmesini sağlar.",
            ],
        ),
    ]
    for heading, paragraphs in sections:
        bt.add_section_heading(doc, heading)
        for paragraph in paragraphs:
            bt.add_para(doc, paragraph)

    bt.add_para(
        doc,
        "Ayrıntılı modül tasarımında ortaya çıkan en önemli sonuç, PrintForge'un tek bir servis entegrasyonundan ibaret olmadığıdır. Sistem, harici AI üretimini kullanıcı deneyimi, güvenlik, veri modeli, pazaryeri ve mesajlaşma katmanlarıyla birlikte ele almaktadır. Bu nedenle proje kapsamı, klasik bir katalog uygulaması veya yalnızca görselden model üreten bir arayüzden daha geniştir.",
    )
    rows = [
        ["Katman", "Sorumluluk", "Kontrol noktası"],
        ["Frontend", "Kullanıcı etkileşimi, yönlendirme, 3D önizleme ve üretim durumu", "Rol tabanlı görünürlük ve anlaşılır durum mesajı"],
        ["Backend", "Doğrulama, kimlik, iş kuralları ve harici servis çağrıları", "Zod şemaları, JWT doğrulaması ve hata kodları"],
        ["Veri modeli", "Kullanıcı, model, konuşma, mesaj ve katalog ilişkileri", "Prisma ilişkileri ve indeksler"],
        ["Dosya yönetimi", "Model dosyaları ve görsel bağlantıları", "Uploads dizini sınırı ve yetki kontrolü"],
        ["Harici servisler", "Görsel arama, görsel yükleme ve AI üretim", "Timeout, kota ve anahtar hatası yönetimi"],
    ]
    bt.add_table(doc, "Tablo 2.5. Ayrıntılı mimari katmanlar ve kontrol noktaları", rows, [3.2, 6.2, 6.1], "Proje kaynak kodu ve modül sorumlulukları esas alınarak hazırlanmıştır.", 9)

    bt.add_section_heading(doc, "2.8. İş Paketleri ve Zaman Yönetimi")
    bt.add_para(
        doc,
        "Bitirme projesi süreci, yalnızca kod yazımı olarak değil; ihtiyaç analizi, mimari kararlar, uygulama geliştirme, entegrasyon, test ve raporlama adımlarını kapsayan bir mühendislik süreci olarak planlanmıştır. Bu nedenle iş paketleri çıktıya göre ayrılmıştır. Her iş paketinin sonunda kontrol edilebilir bir çıktı üretilmesi hedeflenmiştir.",
    )
    bt.add_para(
        doc,
        "İş paketlerinin sıralı tasarlanması, projenin son aşamada bütünleşik görünmesini sağlamıştır. Örneğin AI üretim sayfası katalog ve mesajlaşma modülleri olmadan yalnızca model oluşturan bir araç olarak kalacaktı. Mesajlaşma ve sipariş modüllerinin eklenmesiyle AI çıktısının gerçek üretim talebine dönüşmesi mümkün hale gelmiştir.",
    )
    rows = [
        ["İş paketi", "Amaç", "Çıktı"],
        ["İhtiyaç analizi", "3D baskı kullanıcılarının model ve satıcı bulma problemini tanımlamak", "Problem tanımı, araştırma sorusu ve kapsam"],
        ["Mimari tasarım", "Frontend, backend, veri modeli ve servis ayrımını planlamak", "Sistem mimarisi ve Prisma şeması"],
        ["Katalog geliştirme", "Satıcı ürünlerini ve kullanıcı katalog deneyimini oluşturmak", "Pazaryeri, ürün kartları, filtreler, detaylar"],
        ["AI akışı", "Referans görselden üretim görevine geçişi sağlamak", "Örnekler, prompt aktarımı, görev izleme"],
        ["Mesaj/sipariş", "Satıcı-kullanıcı iletişimini ve sipariş durumlarını izlemek", "Konuşma, mesaj, sipariş ekranları"],
        ["Test ve raporlama", "Derleme, senaryo ve hata durumlarını doğrulamak", "Derleme çıktıları, test tabloları, tez metni"],
    ]
    bt.add_table(doc, "Tablo 2.6. İş paketleri ve çıktıları", rows, [3.2, 6.1, 6.2], "Proje geliştirme süreci esas alınarak hazırlanmıştır.", 9)


def add_scenario_and_quality_sections(doc: Document) -> None:
    bt.add_section_heading(doc, "2.9. Senaryo Tabanlı Sistem Tasarımı")
    bt.add_para(
        doc,
        "Senaryo tabanlı tasarım değerlendirmesi, geliştirilen web uygulamasının yalnızca teknik bileşenlerden oluşmadığını, gerçek kullanıcı görevlerini tamamlayacak biçimde kurgulandığını göstermek için yapılmıştır. Bu yaklaşımda her senaryo bir kullanıcının sistemde tamamlamak istediği iş adımına karşılık gelir. Böylece arayüz, backend rotaları, veritabanı kayıtları ve hata durumları aynı bütün içinde incelenebilir.",
    )
    bt.add_para(
        doc,
        "PrintForge için belirlenen senaryolar üç kullanıcı rolüne göre gruplandırılmıştır: normal kullanıcı, satıcı ve admin. Normal kullanıcı; fikir bulma, AI model üretme, katalog inceleme, favori/sepet kullanma ve satıcıyla iletişime geçme görevlerini yürütür. Satıcı; ürün ekleme, ürünleri yönetme, mesajlara dönme ve sipariş durumunu takip etme görevlerini yürütür. Admin ise örnek içeriklerin sürdürülebilir biçimde yönetilmesini sağlar.",
    )

    scenarios = [
        (
            "2.9.1. Hesap Oluşturma ve Giriş Senaryosu",
            "kullanıcının sisteme güvenli biçimde dahil olması",
            "registerSchema ve loginSchema ile e-posta, parola, ad, rol ve satıcı bilgileri doğrulanır; parola bcrypt ile özetlenir ve başarılı girişte JWT üretilir",
            "kullanıcının favori, sepet, mesaj, AI geçmişi ve sipariş gibi kişisel verilerinin tek hesaba bağlanabilmesi",
            "eksik e-posta, zayıf parola veya hatalı kullanıcı bilgisi veritabanına yazılmadan reddedilir",
        ),
        (
            "2.9.2. Rol Tabanlı Menü ve Yetki Senaryosu",
            "kullanıcının kendi rolüne uygun işlemleri görmesi",
            "Navbar ve korunan sayfalar USER, SELLER ve ADMIN rollerine göre farklı bağlantılar ve eylemler gösterir",
            "satıcı paneli, admin örnek yönetimi ve kullanıcı katalog akışlarının birbirine karışmaması",
            "yetkisiz kullanıcı ürün ekleme veya admin örneği düzenleme sayfalarına yönlendirildiğinde erişim sınırı uygulanır",
        ),
        (
            "2.9.3. Ana Sayfadan Üretim veya Katalog Akışına Geçiş",
            "ziyaretçinin platformun temel değerini ilk ekranda anlaması",
            "ana sayfada katalog inceleme ve AI ile model oluşturma çağrıları birlikte sunulur; ürün önizleme alanı katalog mantığını görsel olarak açıklar",
            "kullanıcının tek sistem içinde hem hazır ürün hem de özel üretim fikrine yönelebilmesi",
            "ana sayfa, sistemin yalnızca tanıtım sayfası değil, iş akışına başlangıç noktası olduğunu doğrular",
        ),
        (
            "2.9.4. Katalog Filtreleme ve Boş Durum Senaryosu",
            "katalog ürünlerinin aranabilir ve anlaşılır biçimde sunulması",
            "ürün adı, açıklama, satıcı adı, kategori ve maksimum fiyat filtresi aynı katalog ekranında işlenir",
            "çok sayıda ürün olduğunda kullanıcının fiyat ve kategoriye göre daha hızlı karar verebilmesi",
            "katalog boş olduğunda kullanıcı pasif bırakılmaz; satıcı katılımı veya AI üretim yönlendirmesi gösterilir",
        ),
        (
            "2.9.5. Örnek Görsel Arama ve AI Sayfasına Aktarım",
            "modelleme bilgisi olmayan kullanıcının üretim fikri bulması",
            "admin örnekleri ve internet görsel arama sonuçları aynı sayfada gösterilir; seçilen görselin başlığı, kaynağı ve prompt bağlamı AI üretim ekranına taşınır",
            "kullanıcının boş bir prompt alanıyla başlamaması ve referans görsel üzerinden üretim niyetini somutlaştırması",
            "görsel arama servisi çalışmasa bile admin örnekleriyle demo ve temel kullanım akışı sürdürülebilir",
        ),
        (
            "2.9.6. AI Model Üretim Görevi Senaryosu",
            "referans görselden 3D model üretim sürecinin başlatılması",
            "görsel dosyası, model açıklaması ve üretim türü backend tarafından doğrulanır; harici AI servisine görev tabanlı istek gönderilir",
            "uzun süren üretimlerin tek ekranda donmuş gibi görünmemesi ve görev durumunun izlenebilir olması",
            "dosya tipi, eksik prompt veya servis anahtarı hatası kullanıcıya anlaşılır biçimde bildirilir",
        ),
        (
            "2.9.7. Model Geçmişi ve 3D Önizleme Senaryosu",
            "tamamlanan AI üretimlerinin tekrar incelenebilmesi",
            "Model kayıtları kullanıcı kimliği ve model türüyle ilişkilendirilir; GLB çıktıları Three.js tabanlı önizleme bileşeniyle görüntülenir",
            "kullanıcının model dosyasını indirmeden önce görsel kontrol yapabilmesi",
            "yetkisiz dosya erişiminde model sahibinin veya ilgili konuşma tarafının doğrulanması gerekir",
        ),
        (
            "2.9.8. Satıcı Ürün Ekleme Senaryosu",
            "satıcının ürününü teknik destek almadan katalogda yayınlaması",
            "ürün adı, açıklama, kategori, fiyat ve en fazla beş görsel alanı doğrulanır; ürün CATALOG türünde Model kaydı olarak saklanır",
            "satıcının ürününü tek fiyat modeliyle yayınlayarak kullanıcıya net başlangıç bilgisi vermesi",
            "satıcı olmayan kullanıcıların ürün yayınlama akışına erişmesi engellenir",
        ),
        (
            "2.9.9. Ürün Detayı, Yorum ve Soru-Cevap Senaryosu",
            "kullanıcının katalog ürününe ilişkin güven oluşturması",
            "ProductReview ve ProductQuestion varlıkları ürünle ilişkilendirilir; cevap yazma yetkisi ürün satıcısıyla sınırlandırılır",
            "ürün bilgisinin yalnızca satıcının tek taraflı açıklamasından ibaret kalmaması",
            "aynı kullanıcının aynı ürüne tekrarlı bağımsız puan vermesi sınırlandırılarak yorum güvenilirliği artırılır",
        ),
        (
            "2.9.10. Favori ve Sepet Senaryosu",
            "kullanıcının karar verme sürecini zamana yayabilmesi",
            "favori ve sepet yardımcı fonksiyonları istemci tarafında kalıcı kullanıcı davranışı oluşturur; katalog kartları bu durumlara göre tepki verir",
            "kullanıcının ilgilendiği ürünleri kaybetmeden daha sonra tekrar inceleyebilmesi",
            "boş favori veya boş sepet ekranları kullanıcıyı katalog akışına geri yönlendirir",
        ),
        (
            "2.9.11. Katalog Ürününden Mesaj Başlatma Senaryosu",
            "alıcı ile satıcının ürün bağlamında iletişime geçmesi",
            "ürün kimliği üzerinden satıcı belirlenir ve Conversation kaydı buyerId, sellerId, modelId ve modelType alanlarıyla oluşturulur",
            "mesajların genel sohbetten farklı olarak belirli ürün veya model bağlamında izlenebilmesi",
            "kullanıcının satıcıya boş ve bağlamsız mesaj göndermesi yerine ürün üzerinden konuşma açması sağlanır",
        ),
        (
            "2.9.12. AI Modeli İçin Satıcı Seçimi Senaryosu",
            "özel üretim modelinin pazaryeri sürecine bağlanması",
            "AI modeli tamamlandığında kullanıcı uygun satıcıya mesaj gönderebilir; konuşma hazır katalog ürününden değil özel modelden başlatılır",
            "AI çıktısının sadece dosya olarak kalmaması, gerçek üretim talebine dönüşebilmesi",
            "model dosyasına erişim konuşmanın taraflarıyla sınırlandırılarak özel üretim çıktısının korunması sağlanır",
        ),
        (
            "2.9.13. Sipariş Durumu Takip Senaryosu",
            "iletişimden üretim sürecine geçen işlerin izlenmesi",
            "Conversation status alanı ORDERED, PREPARING, SHIPPED, COMPLETED ve CANCELLED değerleriyle sipariş yaşam döngüsünü temsil eder",
            "alıcı ve satıcının mesaj metinleri arasında kaybolmadan siparişin güncel aşamasını takip etmesi",
            "mesajlaşma ekranı ile sipariş ekranının ayrılması iletişim ve üretim takibinin karışmasını önler",
        ),
        (
            "2.9.14. Admin Örnek İçerik Yönetimi Senaryosu",
            "platformun kontrollü başlangıç içerikleriyle sürdürülebilmesi",
            "ExampleItem kayıtları title, category, imageUrl, prompt ve tags alanlarıyla doğrulanır; admin yetkisi olmayan kullanıcıların bu alanı düzenlemesi engellenir",
            "dış görsel arama servisi çalışmadığında bile kullanıcının AI üretime başlayabilmesi",
            "demo ve jüri sunumu sırasında dış servis kaynaklı risklerin azaltılması",
        ),
        (
            "2.9.15. Görsel Arama, Önbellek ve Proxy Senaryosu",
            "harici görsellerin güvenli ve kontrollü biçimde kullanılabilmesi",
            "arama sorgusu temizlenir, sağlayıcı seçimi ortam değişkenlerine göre yapılır, sonuçlar kısa süreli cache ile tutulur ve dış görseller proxy üzerinden alınır",
            "API anahtarlarının istemci tarafına taşınmaması ve kota kullanımının azaltılması",
            "image/* olmayan içerikler üretim akışına girmeden reddedilir",
        ),
        (
            "2.9.16. Dağıtım ve Ortam Yapılandırma Senaryosu",
            "uygulamanın geliştirici bilgisayarı dışında da çalıştırılabilmesi",
            "frontend Netlify yapılandırmasıyla hazırlanır; backend ayrı web servisi olarak çalışır ve FRONTEND_URLS, DATABASE_URL, JWT_SECRET gibi değişkenlerle yönetilir",
            "yerel geliştirme ile canlı ortam arasındaki adres, CORS ve servis anahtarı farklarının açıkça ayrılması",
            "localhost değerlerinin canlı ortamda unutulması durumunda bağlantı hatası oluşacağı için README ve ortam dosyaları kritik kabul edilir",
        ),
        (
            "2.9.17. Hata Mesajı ve Kullanıcı Geri Bildirimi Senaryosu",
            "kullanıcının başarısız işlemde ne olduğunu anlayabilmesi",
            "backend hata kodları ve frontend durum mesajları eksik alan, yetkisiz erişim, servis hatası ve dosya formatı sorunlarını ayrıştırır",
            "kullanıcının belirsiz bekleme ekranında kalmaması ve doğru bir sonraki adıma yönelmesi",
            "teknik log ayrıntıları kullanıcıya gösterilmez; kullanıcıya anlaşılır özet bilgi sunulur",
        ),
        (
            "2.9.18. Demo Sürekliliği ve Yedek Akış Senaryosu",
            "savunma ve canlı gösterim sırasında uygulamanın temel akışının korunması",
            "admin örnekleri, boş durum ekranları, manuel katalog ürünü ekleme ve ekran görüntüsüyle desteklenen bulgular birlikte kullanılır",
            "harici servis gecikmesi yaşansa bile proje mimarisi ve kullanıcı akışlarının jüriye gösterilebilmesi",
            "dış servis bağımlılığının bitirme projesinin tamamını başarısız göstermemesi için yedek içerik akışı önemlidir",
        ),
    ]
    for heading, focus, implementation, contribution, acceptance in scenarios:
        bt.add_section_heading(doc, heading)
        bt.add_para(
            doc,
            f"Bu senaryo, {focus} amacıyla değerlendirilmiştir. PrintForge'un hedef kitlesi teknik modelleme bilgisi sınırlı kullanıcılar ve ürün yayınlamak isteyen satıcılar olduğu için her akışın açık, izlenebilir ve hata durumunda anlaşılır olması gerekir. Senaryonun tasarım değeri, kullanıcının tek bir ekranda ne yapacağını anlaması ve işlemin sonucunu sistem içinde takip edebilmesidir.",
        )
        bt.add_para(
            doc,
            f"Uygulama tarafında bu akış {implementation}. Bu karar, frontend ekranı ile backend iş kuralı arasında doğrudan bağ kurar. Böylece arayüzde görülen eylem yalnızca görsel bir düğme olmaktan çıkar; veritabanı kaydı, doğrulama, yetki kontrolü ve kullanıcıya dönen sonuçla birlikte tamamlanan bir işlem haline gelir.",
        )
        bt.add_para(
            doc,
            f"Senaryonun kullanıcıya katkısı {contribution} olarak özetlenebilir. Kabul ölçütü ise {acceptance}. Bu ölçüt sağlanmadığında senaryo tamamlanmış kabul edilmez; çünkü bitirme projesinde yalnızca mutlu yolun çalışması değil, hatalı veya eksik girişlerin de kontrollü yönetilmesi beklenir.",
        )
        bt.add_para(
            doc,
            "Bu değerlendirme, proje kapsamının yalnızca ekran tasarımından ibaret olmadığını göstermektedir. Her senaryo, kullanıcı arayüzü, API katmanı, veri modeli ve hata yönetimi arasında tutarlı bir ilişki kurulup kurulmadığını sınar. Bu nedenle senaryo tabanlı yaklaşım, hem geliştirme sürecinde kontrol listesi hem de tez kapsamında yöntemsel açıklama aracı olarak kullanılmıştır.",
        )

    rows = [
        ["Senaryo grubu", "Ana modül", "Kabul göstergesi"],
        ["Hesap ve rol", "Auth, Navbar, protected routes", "Kullanıcı yalnızca yetkili olduğu ekrana erişir."],
        ["Fikirden AI üretime", "Examples, image search, AI generation", "Referans görsel ve prompt üretim ekranına taşınır."],
        ["Pazaryeri", "Marketplace, product detail, favorites, cart", "Ürünler aranır, incelenir ve kullanıcı listelerine eklenir."],
        ["Satıcı operasyonu", "Seller products, add product, uploads", "Satıcı ürününü görsellerle katalogda yayınlayabilir."],
        ["İletişim ve sipariş", "Chat, orders, conversation status", "Ürün veya AI modeli bağlamında takip edilebilir konuşma oluşur."],
        ["Dayanıklılık", "Admin examples, cache, proxy, error handling", "Dış servis sorunu temel demo akışını tamamen kesmez."],
    ]
    bt.add_table(doc, "Tablo 2.7. Senaryo bazlı kabul matrisi", rows, [4.0, 5.8, 6.7], "Senaryo tabanlı tasarım değerlendirmesi esas alınarak hazırlanmıştır.", 9)

    bt.add_section_heading(doc, "2.10. Kalite, Güvenlik ve Sürdürülebilirlik Değerlendirmesi")
    bt.add_para(
        doc,
        "Kalite değerlendirmesi, sistemin yalnızca çalışıp çalışmadığını değil, sürdürülebilir biçimde geliştirilebilir olup olmadığını da ele alır. PrintForge'da kalite; kodun modülerliği, veri modelinin genişletilebilirliği, güvenlik kontrollerinin açık olması, kullanıcı akışlarının anlaşılır kalması ve dış servis bağımlılıklarının yönetilebilir olması üzerinden incelenmiştir.",
    )
    quality_topics = [
        (
            "2.10.1. Veri Güvenliği ve Kişisel Veri Yönetimi",
            "Kullanıcı hesapları, mesajlar, favoriler ve sipariş durumları kişisel veri niteliği taşıyan kayıtlar oluşturur. Bu nedenle uygulamada parolalar düz metin olarak saklanmaz; bcrypt ile özetlenir. Oturum doğrulaması JWT üzerinden yürütülür ve korunan rotalarda kullanıcının kimliği doğrulanmadan işlem yapılmaz.",
            "Mesajlaşma ve model dosyası erişiminde de yetki kontrolü önemlidir. Kullanıcı yalnızca kendi oluşturduğu modele veya tarafı olduğu konuşmaya bağlı modele erişebilmelidir. Bu kontrol yapılmadığında özel üretim çıktılarının başka kullanıcılar tarafından görüntülenmesi riski doğar.",
            "Bu güvenlik yaklaşımı, bitirme projesi ölçeği için temel ama kritik bir gerekliliktir. Proje gerçek ödeme veya hassas üretim sırrı entegrasyonu içermese de kullanıcı hesabı ve model dosyası gizliliğini korumak akademik değerlendirme açısından önemlidir.",
        ),
        (
            "2.10.2. Telif, Etik ve Harici Görsel Kullanımı",
            "Görsel arama sonuçları kullanıcıya üretim fikri vermek için kullanılır; ancak internette bulunan her görselin üretim veya ticari kullanım hakkı bulunmayabilir. Bu nedenle tez kapsamında harici görsellerin kullanımında lisans ve telif sorumluluğunun dikkate alınması gerektiği açıkça belirtilmiştir.",
            "Admin tarafından yönetilen örnek içeriklerin bulunması bu riski azaltan bir yaklaşımdır. Kontrollü örnekler, hem demo sürecinde hem de kullanıcıya güvenli başlangıç içeriği sunmada yararlıdır. Bu yapı ileride lisans bilgisi, kaynak etiketi ve kullanım izni gibi alanlarla genişletilebilir.",
            "Etik değerlendirme yalnızca görsel telifiyle sınırlı değildir. AI üretim servisleriyle çalışırken kullanıcının yüklediği görselin hangi amaçla kullanılacağı, üretim sonucunun baskıya uygun olup olmadığı ve çıktı kalitesinin garanti edilmediği de kullanıcıya açık biçimde anlatılmalıdır.",
        ),
        (
            "2.10.3. Performans ve Kullanıcı Bekleme Süresi",
            "AI destekli 3D model üretimi çoğu zaman anlık sonuç üretmeyen, servis yoğunluğuna ve model karmaşıklığına bağlı olarak uzayabilen bir süreçtir. Bu nedenle PrintForge'da üretim akışının görev tabanlı izlenmesi tercih edilmiştir. Kullanıcıya görevin başlatıldığı, sürdüğü veya tamamlandığı bilgisinin verilmesi bekleme deneyimini daha anlaşılır hale getirir.",
            "Görsel arama tarafında cache kullanılması performans açısından önemli bir karardır. Aynı sorgunun kısa sürede tekrar tekrar dış servise gönderilmesi hem gecikmeye hem de kota tüketimine neden olabilir. Sorguya ve sağlayıcıya göre oluşturulan önbellek anahtarı, kullanıcıya daha hızlı cevap verilmesini sağlar.",
            "Performans değerlendirmesi yalnızca hızla sınırlı değildir. Uygulamanın boş durum ekranları, hata mesajları ve görev durumu bildirimleri de algılanan performansı etkiler. Kullanıcı sistemin ne yaptığını gördüğünde işlem süresi uzasa bile deneyim daha güvenilir görünür.",
        ),
        (
            "2.10.4. Kodun Sürdürülebilirliği",
            "Kod tabanının frontend ve backend olarak ayrılması, değişikliklerin etkisini sınırlandırmaktadır. Frontend tarafında sayfalar ve bileşenler kullanıcı akışlarına göre düzenlenirken, backend tarafında auth, ai, models, chat, examples ve images gibi rota dosyaları iş alanlarına göre ayrılmıştır.",
            "Bu modüler yapı yeni özellik eklemeyi kolaylaştırır. Örneğin ödeme entegrasyonu eklenecekse chat veya AI üretim rotalarının tamamını değiştirmek gerekmez; sipariş ve ödeme alanında yeni bir servis katmanı oluşturulabilir. Benzer şekilde maliyet tahmini modülü, model dosyası işleme tarafına eklenebilir.",
            "Sürdürülebilirlik açısından ortam değişkenlerinin belgelenmesi de önemlidir. DATABASE_URL, JWT_SECRET, FRONTEND_URLS ve harici servis anahtarları kaynak koddan ayrıldığı için proje farklı geliştirme ve canlı ortamlarında daha güvenli biçimde çalıştırılabilir.",
        ),
        (
            "2.10.5. Kullanılabilirlik ve Arayüz Tutarlılığı",
            "Arayüzde kullanılan menü yapısı, butonlar, boş durum ekranları ve yönlendirmeler kullanıcıya aynı görsel dil içinde sunulmaktadır. Katalog, örnekler ve AI üretim sayfası aynı üst gezinme yapısını paylaşır; bu da kullanıcının nerede olduğunu anlamasını kolaylaştırır.",
            "Boş durum ekranlarının tasarımı özellikle önemlidir. Katalog, favoriler, sepet, mesajlar ve siparişler ekranlarında veri yokken yalnızca boş beyaz alan bırakılmamış; kullanıcıyı ilgili bir sonraki adıma yönlendiren açıklama ve butonlar eklenmiştir. Bu yaklaşım, uygulamanın ilk kullanımda da anlamlı görünmesini sağlar.",
            "Kullanılabilirlik açısından sade form alanları tercih edilmiştir. Satıcı ürün ekleme ekranında ürün adı, açıklama, kategori, fiyat ve görsel yükleme alanları aynı akışta sunulur. Bu yapı, satıcının karmaşık stok yönetimiyle uğraşmadan ürün yayınlamasını sağlar.",
        ),
        (
            "2.10.6. Test Edilebilirlik",
            "Proje test yaklaşımı derleme doğrulaması, manuel senaryo kontrolü ve hata durumlarının gözlenmesi şeklinde ele alınmıştır. Backend TypeScript derlemesinin başarılı olması API katmanındaki tip hatalarının azaltıldığını; frontend üretim derlemesinin başarılı olması ise sayfa yapısının üretim modunda oluşturulabildiğini gösterir.",
            "Manuel senaryo testleri, bitirme projesi ölçeğinde önemlidir çünkü sistem çok sayıda kullanıcı akışını içerir. Kayıt, giriş, ürün ekleme, görsel arama, AI üretim sayfasına geçiş, mesajlaşma ve sipariş takibi tek tek kontrol edilmelidir. Bu testler, yalnızca kodun derlenmesini değil, kullanıcı iş akışının tamamlanmasını doğrular.",
            "İleride otomatik test kapsamı genişletilebilir. Backend için rota testleri, frontend için bileşen ve kullanıcı akışı testleri, model dosyası işleme için örnek dosya tabanlı testler eklenebilir. Bu geliştirmeler projenin akademik prototipten daha olgun bir yazılım ürününe dönüşmesini sağlar.",
        ),
        (
            "2.10.7. Genişletilebilirlik ve Gelecek Çalışmalar",
            "PrintForge'un veri modeli ve modüler mimarisi yeni özelliklerin eklenebilmesine uygundur. Mevcut sistemde Model varlığı hem AI üretimi hem de katalog ürünü için kullanılmaktadır. Bu karar ortak yorum, soru, mesaj ve sipariş ilişkilerinin tek kavram üzerinden yönetilmesini sağlar.",
            "Gelecekte baskıya uygunluk analizi eklenirse model dosyasından hacim, yüzey alanı, sınır kutusu ve tahmini malzeme kullanımı hesaplanabilir. Bu bilgiler satıcı teklif sürecini destekler ve kullanıcıya üretim öncesi daha gerçekçi fiyat beklentisi verir.",
            "Ödeme, kargo, satıcı puanı, lisans filtresi, teklif karşılaştırma ve çoklu AI sağlayıcı desteği de sistemin doğal genişleme alanlarıdır. Bu özellikler eklendiğinde PrintForge yalnızca model üretim ve katalog uygulaması olmaktan çıkarak daha kapsamlı bir 3D baskı hizmet platformuna dönüşebilir.",
        ),
        (
            "2.10.8. Kabul Değerlendirmesine Uygunluk",
            "Bir bitirme tezinde kabul edilebilirlik yalnızca sayfa sayısına bağlı değildir. Projenin çalışır olması, tez metninin projeyle tutarlı olması, yöntem ve bulguların açık yazılması, kaynakçanın akademik düzende verilmesi ve jüriye gösterilebilir ekranların bulunması önemlidir.",
            "PrintForge tezi bu ölçütler açısından değerlendirildiğinde, çalışan site ekran görüntüleri, sistem mimarisi, veri modeli, modül açıklamaları, test senaryoları ve kaynakça ile desteklenen bütünlüklü bir yapı sunmaktadır. Tezin güçlü tarafı, AI üretim fikrini tek başına bırakmayıp pazaryeri, mesajlaşma ve satıcı yönetimiyle tamamlamasıdır.",
            "Kabul riskini azaltmak için savunmada özellikle ana sayfa, örnek görsel arama, AI üretim ekranı, satıcı ürün ekleme, katalog ve mesaj/sipariş akışları sırayla gösterilmelidir. Böyle bir demo akışı, tezde anlatılan mimarinin gerçek uygulama üzerinde karşılığı olduğunu açık biçimde kanıtlar.",
        ),
    ]
    for heading, p1, p2, p3 in quality_topics:
        bt.add_section_heading(doc, heading)
        bt.add_para(doc, p1)
        bt.add_para(doc, p2)
        bt.add_para(doc, p3)

    rows = [
        ["Kalite alanı", "Projede uygulanan yaklaşım", "Geliştirme fırsatı"],
        ["Güvenlik", "JWT, bcrypt, Zod doğrulama ve dosya erişim kontrolü", "Rol bazlı daha ayrıntılı yetki politikaları"],
        ["Performans", "Görsel arama önbelleği ve görev tabanlı AI izleme", "Kuyruk sistemi ve arka plan işleyici"],
        ["Kullanılabilirlik", "Boş durum ekranları, net menü yapısı ve yönlendirici butonlar", "Kullanıcı testi ve erişilebilirlik iyileştirmeleri"],
        ["Sürdürülebilirlik", "Modüler route yapısı ve ortam değişkenleri", "Servis katmanı ve otomatik test kapsamı"],
        ["Genişletilebilirlik", "Ortak Model varlığı, chat ve sipariş ilişkileri", "Maliyet tahmini, ödeme ve kargo entegrasyonu"],
    ]
    bt.add_table(doc, "Tablo 2.8. Teknik kalite ve risk kontrol matrisi", rows, [3.5, 6.0, 6.2], "Kalite değerlendirmesi ve proje geliştirme önerileri esas alınarak hazırlanmıştır.", 9)


def add_chapter_three(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_page_break()
    bt.add_main_heading(doc, "ÜÇÜNCÜ BÖLÜM")
    bt.add_main_heading(doc, "BULGULAR VE TARTIŞMA")
    bt.add_section_heading(doc, "3.1. Gerçekleştirilen Modüller")
    bt.add_para(
        doc,
        "PrintForge projesi kapsamında kullanıcı, satıcı ve admin rollerini destekleyen bütünleşik bir web uygulaması geliştirilmiştir. Sistem; ürün katalogu, örnek görsel arama, AI destekli model üretimi, satıcı ürün yönetimi, favoriler, sepet, mesajlaşma ve sipariş takibi modüllerinden oluşmaktadır. Bulgular, uygulamanın yalnızca tek bir ekran prototipi olmadığını, farklı kullanıcı rollerinin gerçek iş akışlarını destekleyen uçtan uca bir sistem yapısına ulaştığını göstermektedir.",
    )
    bt.add_para(
        doc,
        "Geliştirilen modüller kullanıcı tarafında fikir seçme, model oluşturma ve satıcıyla iletişime geçme ihtiyacını; satıcı tarafında ise ürün ekleme, katalogda yayınlama ve gelen talepleri takip etme ihtiyacını karşılamaktadır. Admin örnekleri ve internet görsel araması, AI üretim akışının başlangıç noktasını güçlendirmiştir. Bu yapı, 3D baskı hizmetini yalnızca teknik dosya üretimi olarak değil, pazaryeri ve iletişim akışıyla birlikte ele alan bir uygulama yaklaşımı sunmaktadır.",
    )
    bt.add_figure(doc, diagrams["fig_3_1"], "Şekil 3.1. Modül tamamlama durumu", 14.0)
    rows = [
        ["Bileşen", "Çalıştırılan doğrulama", "Sonuç"],
        ["Backend", "npm run build / tsc", "Başarılı. TypeScript derlemesi tamamlandı."],
        ["Frontend", "npm run build / next build", "Başarılı. Next.js üretim derlemesi tamamlandı."],
        ["Arayüz ekranları", "Manuel gezinme ve ekran görüntüsü kontrolü", "Ana akış ekranları başarıyla görüntülendi."],
        ["Kullanıcı rolleri", "USER, SELLER ve ADMIN ayrımı", "Rol tabanlı menü ve sayfa akışları doğrulandı."],
    ]
    bt.add_table(
        doc,
        "Tablo 3.1. Derleme doğrulama sonuçları",
        rows,
        [3.2, 5.3, 7.0],
        "Proje üzerinde çalıştırılan derleme komutları ve manuel arayüz kontrolleri esas alınmıştır.",
        9,
    )

    doc.add_page_break()
    bt.add_section_heading(doc, "3.2. Arayüz ve Kullanıcı Akışı Bulguları")
    bt.add_para(
        doc,
        "Bu alt bölümde uygulamanın çalışan web arayüzünden alınan ekran görüntüleri sunulmaktadır. Ekran görüntülerinin amacı, geliştirilen modüllerin yalnızca metin düzeyinde açıklanmadığını, gerçek uygulama üzerinde çalışır biçimde gözlemlendiğini göstermektir. Görseller, tezdeki mimari ve yöntem açıklamalarını somutlaştırmak için kullanılmıştır.",
    )
    add_screen_figure(
        doc,
        "site_home.png",
        "Şekil 3.2. Ana sayfa ve değer önerisi ekranı",
        "Ana sayfa, platformun temel değer önerisini kullanıcıya ilk ekranda sunmaktadır. Sabit fiyatlı katalog ve AI ile model oluşturma seçeneklerinin aynı yüzeyde verilmesi, kullanıcının platformu pazaryeri ve üretim aracı olarak birlikte algılamasını sağlamaktadır.",
        "Bu ekran, çalışmanın problem tanımında belirtilen dağınık hizmet akışını tek giriş noktasında toplama hedefinin arayüz karşılığını göstermektedir. Kullanıcı katalog veya AI üretim yönlerinden birini seçerek aynı platform içinde ilerleyebilmektedir.",
    )
    add_screen_figure(
        doc,
        "catalog_empty.png",
        "Şekil 3.3. Katalog ve filtreleme ekranı",
        "Katalog ekranı arama, kategori ve maksimum fiyat filtreleriyle yapılandırılmıştır. Henüz ürün yok durumunun boş bırakılmaması, satıcı katılımı ve AI model oluşturma akışına yönlendirme yaparak kullanıcıyı sistem içinde tutmaktadır.",
        "Katalog yapısı, ileride ürün sayısı arttığında kullanıcıya arama ve sınıflandırma desteği verecek şekilde tasarlanmıştır. Boş durum ekranının yönlendirici olması, kullanılabilirlik açısından önemli bir tasarım kararıdır.",
    )
    add_screen_figure(
        doc,
        "examples_search.png",
        "Şekil 3.4. Örnek görsel arama ekranı",
        "Örnekler ekranı, kullanıcının üretim fikrine hızlı başlaması için hazır arama önerileri ve görsel sonuçlar sunmaktadır. Görsel arama sonucunun AI üretim sayfasına aktarılabilmesi, fikir aşaması ile model üretimi arasındaki geçişi kısaltmaktadır.",
        "Bu akış, modelleme bilgisi olmayan kullanıcı için fikir bulma eşiğini düşürmektedir. Kullanıcı boş bir metin alanıyla başlamak yerine görsel referans üzerinden üretim niyetini somutlaştırabilmektedir.",
    )
    add_screen_figure(
        doc,
        "ai_generator.png",
        "Şekil 3.5. AI model üretim ekranı",
        "AI model üretim ekranında referans görsel yükleme, model açıklaması yazma ve önceki üretimleri izleme alanları bulunmaktadır. Bu yapı, kullanıcının yalnızca görsel seçip beklemesini değil, üretim amacını açıklayarak daha kontrollü bir sonuç hedeflemesini sağlamaktadır.",
        "AI üretim ekranının ayrı bir çalışma alanı olarak tasarlanması, uzun sürebilen üretim görevlerinin katalog ve mesajlaşma ekranlarından bağımsız yönetilmesini sağlar. Böylece sistem yalnızca görsel arama yapan bir sayfa değil, üretim görevi başlatan bir modül haline gelir.",
    )
    add_screen_figure(
        doc,
        "login.png",
        "Şekil 3.6. Giriş ekranı",
        "Giriş ekranı, katalog, teklif ve AI model akışlarına devam etmek için kimlik doğrulama adımını sunmaktadır. Bu ekran, kullanıcıya ait favori, sepet, mesaj ve sipariş bilgilerinin korunabilmesi için rol tabanlı oturum yapısının başlangıç noktasıdır.",
        "Kimlik doğrulama yapısının sade tutulması, kullanıcıyı temel iş akışından koparmadan hesap güvenliğini sağlamayı hedeflemektedir. Bu ekran aynı zamanda satıcı ve normal kullanıcı rollerinin ayrıştırılacağı oturum bilgisini üretir.",
    )
    add_screen_figure(
        doc,
        "seller_products.png",
        "Şekil 3.7. Satıcı ürün yönetimi ekranı",
        "Satıcı ürün yönetimi ekranı, satıcının katalogda yayınladığı ürünleri tek noktadan görmesini ve yeni ürün ekleme akışına geçmesini sağlar. Bu ekran, platformun yalnızca alıcı tarafı için değil, satıcı operasyonları için de tasarlandığını göstermektedir.",
        "Satıcı panelinin bulunması, projenin pazaryeri boyutunu güçlendirmektedir. Ürün yönetimi alıcı arayüzünden ayrıldığı için satıcı, katalog içeriğini daha düzenli biçimde kontrol edebilmektedir.",
    )
    add_screen_figure(
        doc,
        "seller_add_product.png",
        "Şekil 3.8. Satıcı ürün ekleme ekranı",
        "Ürün ekleme ekranı ürün adı, açıklama, kategori, fiyat ve görsel yükleme alanlarını içerir. Yayın önizlemesinin aynı sayfada bulunması, satıcının ürünün katalogda nasıl görüneceğini yayın öncesinde değerlendirmesini sağlar.",
        "Form yapısında fiyatın sabit değer olarak alınması, projenin pazarlık yerine net fiyat ve mesajlaşma akışını önceleyen yaklaşımıyla uyumludur. Görsel yükleme alanı ise ürünün katalogda güven verici biçimde sunulması için gereklidir.",
    )
    add_screen_figure(
        doc,
        "favorites.png",
        "Şekil 3.9. Favoriler ekranı",
        "Favoriler ekranı, kullanıcının ilgilendiği ürünleri daha sonra incelemek üzere saklamasına olanak tanımaktadır. Bu özellik, katalog inceleme davranışını tek oturumluk bir işlem olmaktan çıkarıp kullanıcı hesabına bağlı kalıcı bir deneyime dönüştürmektedir.",
        "Favoriler modülü, kullanıcı karar verme sürecini destekleyen yardımcı bir bileşendir. Özellikle çok sayıda katalog ürünü olduğunda kullanıcı, satın alma veya mesajlaşma kararını erteleyebilmekte ve ilgilendiği ürünleri kaybetmemektedir.",
    )
    add_screen_figure(
        doc,
        "messages.png",
        "Şekil 3.10. Mesajlar ekranı",
        "Mesajlar ekranı satıcı ve müşteri konuşmalarını sipariş ekranından ayrı olarak yönetmektedir. Bu ayrım, ürün hakkında bilgi alma, teklif netleştirme ve üretim sürecini konuşma gibi iletişim ihtiyaçlarının düzenli takip edilmesini sağlar.",
        "Mesajlaşma modülü, AI üretim veya katalog ürünü sonrasında satıcıyla temas kurulmasını sağlayan ana bağlantıdır. Siparişten ayrı tutulması, bilgi alma konuşmaları ile aktif üretim süreçlerinin karışmasını azaltmaktadır.",
    )
    add_screen_figure(
        doc,
        "orders.png",
        "Şekil 3.11. Sipariş takip ekranı",
        "Sipariş takip ekranı, satıcının gelen siparişleri ve sipariş durumu değişikliklerini görmesi için tasarlanmıştır. Mesaj ekranından ayrılan bu yapı, iletişim ile üretim/sipariş takibinin karışmasını azaltır.",
        "Sipariş ekranı, projenin yalnızca model fikri üretmekle kalmayıp hizmet sürecini takip edebilecek bir pazaryeri altyapısına yöneldiğini göstermektedir. İleride ödeme, kargo ve üretim aşaması bildirimleri bu ekran üzerinden genişletilebilir.",
    )

    doc.add_page_break()
    bt.add_section_heading(doc, "3.3. Derleme ve Teknik Doğrulama Bulguları")
    bt.add_para(
        doc,
        "Projenin teknik doğrulaması için backend ve frontend üretim derlemeleri çalıştırılmıştır. Backend katmanında TypeScript derlemesi hata vermeden tamamlanmıştır. Frontend katmanında Next.js üretim derlemesi tamamlanmış ve uygulama sayfaları üretim çıktısı içinde listelenmiştir. Derleme sırasında görülen önbellek uyarıları uygulamanın çalışmasını engelleyen işlevsel hata olarak değerlendirilmemiştir.",
    )
    bt.add_figure(doc, diagrams["fig_3_2"], "Şekil 3.12. Derleme ve doğrulama özeti", 14.0)

    bt.add_section_heading(doc, "3.4. Kullanıcı Akışlarına İlişkin Bulgular")
    bt.add_para(
        doc,
        "Kullanıcı akışları incelendiğinde PrintForge'un fikir aşamasından satıcıyla iletişim aşamasına kadar kesintisiz bir deneyim sunduğu görülmektedir. Kullanıcı örnekler sayfasında bir görsel arayabilir, seçilen görselden AI üretim ekranına geçebilir, katalog ürünlerini inceleyebilir, favori veya sepet listesi oluşturabilir ve satıcıya mesaj gönderebilir. Satıcı ise ürünlerini yayınlayabilir ve sipariş takibini ayrı bir ekranda gerçekleştirebilir.",
    )
    rows = [
        ["Akış", "Gerçekleşen davranış", "Kullanıcı katkısı"],
        ["Örnekten AI üretime geçiş", "Görsel ve bağlam AI üretim ekranına taşınır.", "Kullanıcı boş prompt alanıyla başlamaz."],
        ["Katalog", "Ürünler arama, kategori ve fiyat filtresiyle incelenir.", "Hazır üretim seçeneklerine hızlı erişim sağlanır."],
        ["Satıcı paneli", "Ürün görseli, kategori, açıklama ve fiyat yönetilir.", "Satıcı teknik destek almadan katalog oluşturabilir."],
        ["Mesajlaşma", "Satıcı ve kullanıcı konuşmaları ayrı ekranda tutulur.", "Üretim öncesi iletişim izlenebilir hale gelir."],
        ["Sipariş takibi", "Gelen siparişler durum akışıyla izlenir.", "Üretim sürecinin takibi kolaylaşır."],
    ]
    bt.add_table(
        doc,
        "Tablo 3.2. Gerçekleştirilen kullanıcı akışları",
        rows,
        [3.2, 5.3, 7.0],
        "Frontend sayfaları ve kullanıcı senaryoları birlikte incelenerek hazırlanmıştır.",
        9,
    )

    bt.add_section_heading(doc, "3.5. Güvenlik, Sınırlılıklar ve Geliştirme Önerileri")
    bt.add_para(
        doc,
        "Güvenlik açısından elde edilen en önemli bulgu, hassas servis anahtarlarının frontend tarafına taşınmadan backend üzerinden yönetilmesidir. Görsel arama sağlayıcıları, AI üretim servisleri ve veritabanı bağlantısı ortam değişkenleriyle ayrılmıştır. Kimlik doğrulama akışı JWT ve bcrypt ile desteklenmekte, korunan rotalarda token doğrulaması yapılmaktadır.",
    )
    bt.add_para(
        doc,
        "Sistemin temel sınırlılıkları AI üretim kalitesinin harici servise bağlı olması, modelin baskıya uygunluğunun otomatik mühendislik analiziyle doğrulanmaması ve ödeme/kargo entegrasyonlarının bitirme projesi kapsamı dışında bırakılmasıdır. Buna karşın mevcut mimari, çoklu AI sağlayıcı desteği, maliyet tahmini, baskı uygunluk analizi ve ödeme entegrasyonu gibi geliştirmeler için genişletilebilir bir temel sunmaktadır.",
    )
    rows = [
        ["Sınırlılık", "Etkisi", "Geliştirme önerisi"],
        ["AI servis bağımlılığı", "Üretim süresi ve çıktı kalitesi dış servise bağlıdır.", "Birden fazla sağlayıcı için yedekleme mekanizması kurulabilir."],
        ["Baskı uygunluk analizi yok", "Modelin hacim, duvar kalınlığı ve destek ihtiyacı otomatik yorumlanmaz.", "STL/GLB geometri analizi ve uygunluk puanı eklenebilir."],
        ["Telif/lisans kontrolü sınırlı", "Harici görsellerin kullanım hakkı kullanıcı sorumluluğundadır.", "Lisans filtresi ve kaynak uyarıları geliştirilebilir."],
        ["Ödeme ve lojistik entegrasyonu yok", "Sipariş akışı iletişim ve durum takibi düzeyindedir.", "Ödeme, kargo ve fatura modülleri eklenebilir."],
        ["Maliyet tahmini yok", "Kullanıcı üretim maliyetini satıcıyla konuşarak öğrenir.", "Malzeme, hacim ve doluluk oranına dayalı tahmin yapılabilir."],
    ]
    bt.add_table(
        doc,
        "Tablo 3.3. Sınırlılıklar ve geliştirme önerileri",
        rows,
        [3.2, 5.0, 7.3],
        "Proje kapsamı ve test bulguları esas alınarak hazırlanmıştır.",
        9,
    )

    bt.add_section_heading(doc, "3.6. Tartışma")
    bt.add_para(
        doc,
        "Elde edilen bulgular, PrintForge'un araştırma sorusuna olumlu yanıt verdiğini göstermektedir. Teknik modelleme bilgisi sınırlı kullanıcılar için referans görsel seçimi, AI üretim görevi ve satıcıyla iletişim tek bir web uygulamasında birleştirilebilmiştir. Bu yönüyle proje, 3D baskının yaygınlaşmasında yalnızca üretim donanımının değil, kullanıcıyı doğru hizmet akışına yönlendiren yazılım arayüzlerinin de önemli olduğunu göstermektedir.",
    )
    bt.add_para(
        doc,
        "Literatürde eklemeli imalatın dağıtık üretim ve kişiselleştirme potansiyeli vurgulanırken kalite kontrol ve standartlaşma sorunlarının sürdüğü görülmektedir (Ford ve Despeisse,2016:1573-1587; Ngo vd.,2018:172-196). PrintForge bu sorunları tamamen çözmemekte, ancak kullanıcı-satıcı iletişimi ve AI destekli başlangıç akışıyla 3D baskıya erişim bariyerlerini azaltan uygulanabilir bir ara katman sunmaktadır.",
    )
    bt.add_para(
        doc,
        "Bu nedenle projenin katkısı, yeni bir 3D üretim algoritması önermekten çok, mevcut AI ve web teknolojilerini 3D baskı hizmet sürecine bütünleşik şekilde uygulamasıdır. Tez kapsamında sunulan ekran görüntüleri, sistemin gerçek arayüz akışlarını desteklediğini ve projenin savunmada gösterilebilir bir yazılım çıktısı niteliği taşıdığını ortaya koymaktadır.",
    )

    bt.add_section_heading(doc, "3.7. Demo ve Kabul Değerlendirmesi")
    bt.add_para(
        doc,
        "Bitirme projesinin kabul edilebilirliği açısından çalışan demo, tez metnindeki mimari açıklamaların gerçek sistem üzerinde karşılığının bulunduğunu göstermelidir. Bu nedenle PrintForge için önerilen demo akışı yalnızca ekranların sırasıyla açılmasından ibaret değildir; her ekranın hangi araştırma sorusuna, hangi modüle ve hangi kullanıcı ihtiyacına karşılık geldiği açıklanmalıdır.",
    )
    bt.add_para(
        doc,
        "Savunmada ilk gösterilmesi gereken ekran ana sayfadır. Ana sayfa, platformun iki ana değer önerisini aynı anda sunmaktadır: kullanıcı hazır katalog ürünlerini inceleyebilir veya AI ile yeni model fikri oluşturabilir. Bu ekran, tezin problem tanımında belirtilen dağınık hizmet akışlarını tek platformda toplama hedefini görsel olarak kanıtlar.",
    )
    bt.add_para(
        doc,
        "İkinci adımda örnekler ve görsel arama ekranı gösterilmelidir. Bu ekran, teknik modelleme bilgisi olmayan kullanıcının üretim fikri bulmasını kolaylaştırır. Kullanıcının 'vazo' gibi basit bir arama ile referans görsellere ulaşması ve bu görsellerin AI üretim sürecine bağlanabilmesi, çalışmanın kullanılabilirlik katkısını somutlaştırır.",
    )
    bt.add_para(
        doc,
        "Üçüncü adım AI model üretim ekranıdır. Bu ekranda referans görsel, model açıklaması ve önceki üretimler birlikte görülür. Demo sırasında üretim servisinin anlık olarak sonuç vermesi zorunlu değildir; önemli olan görev tabanlı üretim mantığının, referans görsel kullanımının ve üretim geçmişi yaklaşımının açıklanmasıdır. Böylece dış servis gecikmesi projenin ana değerini gölgelememiş olur.",
    )
    bt.add_para(
        doc,
        "Dördüncü adım katalog ve satıcı ürün ekleme ekranlarıdır. Katalog ekranı, ürünlerin arama, kategori ve fiyat bilgisiyle sunulduğunu gösterir. Satıcı ürün ekleme ekranı ise ürün adı, açıklama, kategori, fiyat ve görsellerin tek form üzerinden yönetilebildiğini kanıtlar. Bu iki ekran birlikte gösterildiğinde platformun yalnızca AI üretim aracı değil, aynı zamanda pazaryeri altyapısı olduğu anlaşılır.",
    )
    bt.add_para(
        doc,
        "Beşinci adım mesajlaşma ve sipariş takibi ekranlarıdır. Bu ekranlar, kullanıcının bir ürün veya AI modeli hakkında satıcıyla iletişime geçebildiğini ve üretim sürecinin sipariş durumlarıyla takip edilebildiğini gösterir. Mesajlar ekranının sipariş ekranından ayrılması, bilgi alma konuşmaları ile aktif üretim takibinin birbirine karışmasını azaltan bilinçli bir tasarım kararıdır.",
    )
    bt.add_para(
        doc,
        "Demo değerlendirmesinde güvenlik kararları da ayrıca açıklanmalıdır. API anahtarlarının frontend tarafına taşınmaması, JWT tabanlı oturum yönetimi, Zod doğrulama şemaları ve model dosyası erişim kontrolü projenin yalnızca görsel olarak çalışan bir arayüz değil, temel güvenlik ilkelerini dikkate alan bir yazılım çıktısı olduğunu gösterir.",
    )
    bt.add_para(
        doc,
        "Kabul açısından bir diğer önemli nokta sınırlılıkların açıkça belirtilmesidir. PrintForge gerçek ödeme, kargo ve otomatik baskı uygunluk analizi entegrasyonlarını bu çalışma kapsamında tamamlamamıştır. Ancak bu sınırlılıklar saklanmamış; sonuç ve tartışma bölümlerinde gelecek çalışma önerileriyle birlikte açıklanmıştır. Akademik değerlendirmede bu açıklık, projenin kapsamının gerçekçi biçimde belirlendiğini gösterir.",
    )
    bt.add_para(
        doc,
        "Son olarak demo sırasında kullanılan ekran görüntüleri ve tezdeki şekiller birbirini desteklemelidir. Tezde yer alan mimari şemalar sistemin arka planını, ekran görüntüleri ise kullanıcıya görünen yüzünü açıklar. Bu iki kanıt türünün birlikte verilmesi, çalışmanın hem teknik hem de kullanılabilirlik açısından savunulabilir olmasını sağlar.",
    )
    rows = [
        ["Demo adımı", "Gösterilecek ekran", "Kanıtladığı tez iddiası"],
        ["1", "Ana sayfa", "Platform katalog ve AI üretim akışlarını tek girişte birleştirir."],
        ["2", "Örnekler ve görsel arama", "Kullanıcı modelleme bilgisi olmadan üretim fikri bulabilir."],
        ["3", "AI model üretim ekranı", "Referans görsel ve açıklama üretim görevine dönüştürülebilir."],
        ["4", "Katalog ve satıcı ürün ekleme", "Sistem pazaryeri işlevi ve satıcı operasyonu destekler."],
        ["5", "Mesajlar ve siparişler", "Alıcı-satıcı iletişimi ve sipariş durumu izlenebilir."],
        ["6", "Mimari ve veri modeli açıklaması", "Arayüzde görülen akışların backend ve veritabanı karşılığı vardır."],
    ]
    bt.add_table(
        doc,
        "Tablo 3.4. Savunma demosu ve kanıt matrisi",
        rows,
        [2.2, 4.8, 8.8],
        "Tez bulguları, ekran görüntüleri ve önerilen savunma akışı esas alınarak hazırlanmıştır.",
        9,
    )
    bt.add_para(
        doc,
        "Bu demo matrisi, savunma sırasında anlatımın dağılmasını önlemek için hazırlanmıştır. Bitirme projelerinde en sık görülen sorunlardan biri, çalışan ekranların gösterilmesine rağmen bu ekranların tezdeki problem tanımıyla ilişkilendirilememesidir. PrintForge için önerilen sıra, önce problemin kullanıcıya görünen yüzünü, sonra üretim ve pazaryeri akışını, en sonunda da teknik mimariyi açıklayacak biçimde düzenlenmiştir.",
    )
    bt.add_para(
        doc,
        "Jüri değerlendirmesinde ana sayfa ve katalog ekranı, projenin kullanıcı deneyimi yönünü kanıtlar. Örnekler ve AI üretim ekranları, çalışmanın yapay zeka destekli 3D model üretimi boyutunu gösterir. Satıcı paneli, ürün ekleme, mesajlar ve siparişler ekranları ise projenin yalnızca tek kullanıcılı bir prototip olmadığını, satıcı-alıcı etkileşimini destekleyen iki taraflı bir platform olarak tasarlandığını ortaya koyar.",
    )
    bt.add_para(
        doc,
        "Teknik mimari açıklanırken frontend ve backend ayrımının neden yapıldığı özellikle vurgulanmalıdır. Frontend kullanıcı etkileşimini, yönlendirmeyi, 3D önizlemeyi ve durum gösterimini üstlenirken; backend kimlik doğrulama, veri doğrulama, servis anahtarlarını koruma, harici servis çağrıları ve dosya erişim kontrolü gibi güvenlik açısından kritik işleri yürütür. Bu ayrım, projenin yalnızca görsel bir arayüz değil, katmanlı bir yazılım sistemi olduğunu gösterir.",
    )
    bt.add_para(
        doc,
        "Kabul olgunluğu açısından projenin güçlü tarafı, her modülün başka bir modülle ilişkili olmasıdır. Örnek görsel araması AI üretim ekranına, AI çıktısı satıcıyla mesajlaşmaya, katalog ürünü sepet ve favorilere, mesajlaşma ise sipariş durumuna bağlanmaktadır. Bu bağlantılar sayesinde sistemde izole ekranlar yerine uçtan uca bir hizmet akışı oluşmuştur.",
    )
    bt.add_para(
        doc,
        "Projenin sınırlılıklarının açıkça yazılması kabul açısından olumsuz değil, aksine olgun bir mühendislik yaklaşımıdır. Otomatik baskı uygunluk analizi, ödeme, kargo ve gelişmiş maliyet tahmini gibi alanların gelecekte yapılacak çalışmalar olarak belirtilmesi, çalışmanın kapsamının gerçekçi çizildiğini gösterir. Böylece jüri, eksik bırakılan alanların fark edilmediği için değil, bitirme projesi kapsamı dışında tutulduğu için tamamlanmadığını anlayabilir.",
    )
    bt.add_para(
        doc,
        "Sonuç olarak demo ve tez metni birlikte değerlendirildiğinde PrintForge, problem tanımı, yöntem, uygulama, bulgular ve sonuç bölümleri arasında tutarlı bir çizgi kurmaktadır. Kullanıcı ihtiyacı giriş bölümünde açıklanmış, literatürle desteklenmiş, materyal ve metot bölümünde yazılım mimarisiyle karşılanmış, bulgular bölümünde gerçek ekran görüntüleriyle gösterilmiş ve sonuç bölümünde geliştirilebilir yönleriyle birlikte değerlendirilmiştir.",
    )


def build(total_pages: int = 48) -> None:
    configure_lists()
    page_map_path = WORKSPACE / "thesis_work" / "page_map_45.json"
    page_map_path.write_text(json.dumps(PAGE_MAP, ensure_ascii=False, indent=2), encoding="utf-8")

    diagrams = bt.make_diagrams()
    doc = Document()
    bt.configure_styles(doc)
    bt.setup_section(doc.sections[0], footer=False)
    bt.add_cover_pages(doc, total_pages)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    bt.setup_section(front, footer=True, start=5, fmt="lowerRoman")
    bt.add_abstracts(doc, total_pages)
    bt.add_front_matter(doc, PAGE_MAP)

    main = doc.add_section(WD_SECTION.NEW_PAGE)
    bt.setup_section(main, footer=True, start=1, fmt="decimal")
    bt.add_intro(doc)
    bt.add_chapter_one(doc, diagrams)
    bt.add_chapter_two(doc, diagrams)
    add_extended_technical_design_compact(doc)
    add_scenario_and_quality_sections(doc)
    add_chapter_three(doc, diagrams)
    bt.add_conclusion_refs_appendices(doc)
    bt.enable_update_fields_on_open(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
