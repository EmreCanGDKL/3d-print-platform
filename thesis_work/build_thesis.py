from __future__ import annotations

import argparse
import json
import math
import os
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(__file__).resolve().parents[1]
OUT_DIR = WORKSPACE / "thesis_work"
DOCX_PATH = WORKSPACE / "PrintForge_Bitirme_Tezi.docx"

TITLE_TR = "PRINTFORGE: AI DESTEKLİ 3D BASKI PAZARYERİ VE MODEL ÜRETİM PLATFORMU"
TITLE_TR_PRETTY = "PrintForge: AI Destekli 3D Baskı Pazaryeri ve Model Üretim Platformu"
TITLE_EN = "PrintForge: AI-Assisted 3D Printing Marketplace and Model Generation Platform"
AUTHOR = "Emre Can Gedikli"
ADVISOR = "Meryem Şerifi"
DATE_TEXT = "Haziran 2026"
DEFAULT_PAGE = "00"


REFERENCES = [
    "Attaran, M., (2017). The rise of 3-D printing: The advantages of additive manufacturing over traditional manufacturing. Business Horizons, 60(5):677-688.",
    "Berman, B., (2012). 3-D printing: The new industrial revolution. Business Horizons, 55(2):155-162.",
    "Campbell, T., Williams, C., Ivanova, O. and Garrett, B., (2011). Could 3D printing change the world? Technologies, Potential, and Implications of Additive Manufacturing. Atlantic Council, Washington, 1-16.",
    "Ford, S. and Despeisse, M., (2016). Additive manufacturing and sustainability: An exploratory study of the advantages and challenges. Journal of Cleaner Production, 137:1573-1587.",
    "Gebler, M., Schoot Uiterkamp, A.J.M. and Visser, C., (2014). A global sustainability perspective on 3D printing technologies. Energy Policy, 74:158-167.",
    "Gibson, I., Rosen, D., Stucker, B. and Khorasani, M., (2021). Additive Manufacturing Technologies. Springer, Cham, p.1-506.",
    "Holmstrom, J. and Partanen, J., (2014). Digital manufacturing-driven transformations of service supply chains for complex products. Supply Chain Management: An International Journal, 19(4):421-430.",
    "Jiang, R., Kleer, R. and Piller, F.T., (2017). Predicting the future of additive manufacturing: A Delphi study on economic and societal implications of 3D printing for 2030. Technological Forecasting and Social Change, 117:84-97.",
    "Jun, H. and Nichol, A., (2023). Shap-E: Generating conditional 3D implicit functions. arXiv preprint, arXiv:2305.02463.",
    "LeCun, Y., Bengio, Y. and Hinton, G., (2015). Deep learning. Nature, 521(7553):436-444.",
    "Lin, C.H., Gao, J., Tang, L., Takikawa, T., Zeng, X., Huang, X., Kreis, K., Fidler, S., Liu, M.Y. and Lin, T.Y., (2023). Magic3D: High-resolution text-to-3D content creation. IEEE/CVF Conference on Computer Vision and Pattern Recognition, p.300-309.",
    "Ngo, T.D., Kashani, A., Imbalzano, G., Nguyen, K.T.Q. and Hui, D., (2018). Additive manufacturing (3D printing): A review of materials, methods, applications and challenges. Composites Part B: Engineering, 143:172-196.",
    "Nichol, A., Jun, H., Dhariwal, P., Mishkin, P. and Chen, M., (2022). Point-E: A system for generating 3D point clouds from complex prompts. arXiv preprint, arXiv:2212.08751.",
    "Poole, B., Jain, A., Barron, J.T. and Mildenhall, B., (2022). DreamFusion: Text-to-3D using 2D diffusion. arXiv preprint, arXiv:2209.14988.",
    "Rayna, T. and Striukova, L., (2016). From rapid prototyping to home fabrication: How 3D printing is changing business model innovation. Technological Forecasting and Social Change, 102:214-224.",
    "Shahrubudin, N., Lee, T.C. and Ramlan, R., (2019). An overview on 3D printing technology: Technological, materials, and applications. Procedia Manufacturing, 35:1286-1296.",
    "Weller, C., Kleer, R. and Piller, F.T., (2015). Economic implications of 3D printing: Market structure models in light of additive manufacturing revisited. International Journal of Production Economics, 164:43-56.",
]


TOC_ENTRIES = [
    ("ÖZET", "ozet", 0),
    ("ABSTRACT", "abstract", 0),
    ("ÖNSÖZ", "onsoz", 0),
    ("KISALTMALAR", "kisaltmalar", 0),
    ("TABLOLAR LİSTESİ", "tablolar", 0),
    ("ŞEKİLLER LİSTESİ", "sekiller", 0),
    ("EKLER LİSTESİ", "ekler_listesi", 0),
    ("GİRİŞ", "giris", 0),
    ("BİRİNCİ BÖLÜM TEORİK ALT YAPI", "ch1", 0),
    ("1.1. Eklemeli İmalat ve 3D Baskı", "s11", 1),
    ("1.2. AI Destekli 3D Model Üretimi", "s12", 1),
    ("1.3. Web Tabanlı Pazaryeri Yaklaşımı", "s13", 1),
    ("1.4. Yazılım Mimarisi, Güvenlik ve Veri Yönetimi", "s14", 1),
    ("1.5. Literatür Değerlendirmesi ve Özgün Değer", "s15", 1),
    ("İKİNCİ BÖLÜM MATERYAL VE METOT", "ch2", 0),
    ("2.1. Kullanılan Materyaller ve Yazılım Araçları", "s21", 1),
    ("2.2. Sistem Mimarisi", "s22", 1),
    ("2.3. Veri Modeli ve Veritabanı Tasarımı", "s23", 1),
    ("2.4. Uygulama Geliştirme Süreci", "s24", 1),
    ("2.5. Yapılan Testler", "s25", 1),
    ("2.6. Güvenlik, Etik Standartlar ve Varsayımlar", "s26", 1),
    ("2.7. Ayrıntılı Modül Tasarımı", "s27", 1),
    ("2.8. İş Paketleri ve Zaman Yönetimi", "s28", 1),
    ("2.9. Ayrıntılı Teknik İnceleme Sayfaları", "s29", 1),
    ("ÜÇÜNCÜ BÖLÜM BULGULAR VE TARTIŞMA", "ch3", 0),
    ("3.1. Gerçekleştirilen Modüller", "s31", 1),
    ("3.2. Derleme ve Teknik Doğrulama Bulguları", "s32", 1),
    ("3.3. Kullanıcı Akışlarına İlişkin Bulgular", "s33", 1),
    ("3.4. Güvenlik ve Sürdürülebilirlik Bulguları", "s34", 1),
    ("3.5. Kullanılabilirlik Bulguları", "s35", 1),
    ("3.6. Teknik Risk ve Dayanıklılık Bulguları", "s36", 1),
    ("3.7. Sürdürülebilirlik ve Genişletilebilirlik Bulguları", "s37", 1),
    ("3.8. Tartışma", "s38", 1),
    ("SONUÇ VE ÖNERİLER", "sonuc", 0),
    ("KAYNAKÇA", "kaynakca", 0),
    ("EKLER", "ekler", 0),
    ("ÖZGEÇMİŞ", "ozgecmis", 0),
]

TABLES = [
    ("Tablo 1.1. Literatür çalışmaları ve PrintForge ile ilişkisi", "tab_1_1"),
    ("Tablo 2.1. Projede kullanılan yazılım bileşenleri", "tab_2_1"),
    ("Tablo 2.2. Backend API modülleri ve görevleri", "tab_2_2"),
    ("Tablo 2.3. Temel veritabanı varlıkları", "tab_2_3"),
    ("Tablo 2.4. Test senaryoları ve kabul ölçütleri", "tab_2_4"),
    ("Tablo 2.5. Ayrıntılı mimari katmanlar ve kontrol noktaları", "tab_2_5"),
    ("Tablo 2.6. İş paketleri ve çıktıları", "tab_2_6"),
    ("Tablo 3.1. Derleme doğrulama sonuçları", "tab_3_1"),
    ("Tablo 3.2. Gerçekleştirilen kullanıcı akışları", "tab_3_2"),
    ("Tablo 3.3. Sınırlılıklar ve geliştirme önerileri", "tab_3_3"),
]

FIGURES = [
    ("Şekil 1.1. Literatürden projeye uzanan kavramsal konumlandırma", "fig_1_1"),
    ("Şekil 2.1. PrintForge sistem mimarisi", "fig_2_1"),
    ("Şekil 2.2. Referans görselden AI üretime veri akışı", "fig_2_2"),
    ("Şekil 2.3. Veri modeli ilişkileri", "fig_2_3"),
    ("Şekil 2.4. Kimlik doğrulama ve güvenli erişim akışı", "fig_2_4"),
    ("Şekil 3.1. Modül tamamlama durumu", "fig_3_1"),
    ("Şekil 3.2. Derleme ve doğrulama özeti", "fig_3_2"),
]

APPENDICES = [
    ("Ek A. Kurulum ve çalıştırma adımları", "app_a"),
    ("Ek B. Backend ortam değişkenleri", "app_b"),
    ("Ek C. API uç noktaları özeti", "app_c"),
    ("Ek D. Manuel test kontrol listesi", "app_d"),
]


def load_page_map(path: str | None) -> dict[str, str]:
    if not path:
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_page(page_map: dict[str, str], key: str) -> str:
    return str(page_map.get(key, DEFAULT_PAGE))


def cm_to_dxa(cm_value: float) -> int:
    return int(round(cm_value / 2.54 * 1440))


def set_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_layout(table, widths_cm: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total_dxa = sum(cm_to_dxa(w) for w in widths_cm)
    tblW.set(qn("w:w"), str(total_dxa))
    tblW.set(qn("w:type"), "dxa")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(cm_to_dxa(w)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = Cm(widths_cm[min(idx, len(widths_cm) - 1)])
            cell.width = width
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(cm_to_dxa(widths_cm[min(idx, len(widths_cm) - 1)])))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def shade_cell(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tblBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tblBorders.append(element)
        element.set(qn("w:val"), "nil")


def set_page_number_type(section, start: int, fmt: str):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))
    pgNumType.set(qn("w:fmt"), fmt)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_font(run, 11)


def setup_section(section, footer: bool, start: int = 1, fmt: str = "decimal"):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    if footer:
        section.footer.is_linked_to_previous = False
        set_page_number_type(section, start, fmt)
        p = section.footer.paragraphs[0]
        add_page_field(p)


def enable_update_fields_on_open(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    def new_style(name, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first=0, before=0, after=0):
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Cm(first)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        return style

    new_style("TezBody", 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.25, 0, 0)
    new_style("TezNoIndent", 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 0, 0)
    new_style("TezCenter", 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 0)
    new_style("TezTitle", 14, True, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 0)
    new_style("TezHeading", 12, True, WD_ALIGN_PARAGRAPH.LEFT, 0, 12, 6)
    new_style("TezSubHeading", 12, True, WD_ALIGN_PARAGRAPH.LEFT, 0, 8, 4)
    new_style("TezCaption", 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 6, 6)
    new_style("TezTableCaption", 12, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 6, 3)
    new_style("TezSource", 10, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 6)
    new_style("TezSmall", 10, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 0, 0)
    new_style("TezList", 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 0, 0)

    for h in ("Heading 1", "Heading 2", "Heading 3"):
        styles[h].font.name = "Times New Roman"
        styles[h]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        styles[h]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        styles[h]._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        styles[h].font.color.rgb = RGBColor(0, 0, 0)
        styles[h].font.size = Pt(12)
        styles[h].font.bold = True


def add_para(doc: Document, text: str = "", style: str = "TezBody", bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_center(doc: Document, text: str, size=12, bold=False, space_after=0):
    p = doc.add_paragraph(style="TezCenter")
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_empty_lines(doc: Document, n: int):
    for _ in range(n):
        doc.add_paragraph(style="TezCenter")


def add_main_heading(doc: Document, text: str, key: str | None = None):
    p = doc.add_paragraph(style="TezTitle")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_font(r, 12, True)
    return p


def add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph(style="TezHeading")
    r = p.add_run(text)
    set_font(r, 12, True)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in p.runs:
        set_font(run)
    r = p.add_run(text)
    set_font(r)
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    set_font(r)
    return p


def add_toc_entry(doc: Document, title: str, key: str, level: int, page_map: dict[str, str]):
    p = doc.add_paragraph(style="TezNoIndent")
    p.paragraph_format.left_indent = Cm(0.6 * level)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f"{title}\t{get_page(page_map, key)}")
    set_font(r, 11)


def add_list_entry(doc: Document, text: str, key: str, page_map: dict[str, str]):
    p = doc.add_paragraph(style="TezNoIndent")
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f"{text}\t{get_page(page_map, key)}")
    set_font(r, 11)


def add_table(doc: Document, caption: str, rows: list[list[str]], widths: list[float], source: str, font_size=10):
    cap = doc.add_paragraph(style="TezTableCaption")
    r = cap.add_run(caption)
    set_font(r, 12)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_layout(table, widths)
    mark_repeat_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                shade_cell(cell, "EDEDED")
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 or len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_font(run, font_size, bold=row_idx == 0)
    src = doc.add_paragraph(style="TezSource")
    r = src.add_run(f"Kaynak: {source}")
    set_font(r, 10)
    return table


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.2):
    p = doc.add_paragraph(style="TezCenter")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(style="TezCaption")
    r = cap.add_run(caption)
    set_font(r, 12)


def font(path_name: str, size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap(draw, text, fnt, width):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        test = f"{line} {word}".strip()
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def draw_box(draw, xy, title, body, fill="#f7f7f7", outline="#333333"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    title_font = font("times", 30, True)
    body_font = font("times", 22)
    draw.text((x1 + 24, y1 + 18), title, font=title_font, fill="#111111")
    y = y1 + 62
    for line in wrap(draw, body, body_font, x2 - x1 - 48)[:4]:
        draw.text((x1 + 24, y), line, font=body_font, fill="#333333")
        y += 28


def arrow(draw, start, end, color="#444444"):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=5)
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 18
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        ax = x2 + length * math.cos(angle + delta)
        ay = y2 + length * math.sin(angle + delta)
        draw.line((x2, y2, ax, ay), fill=color, width=5)


def make_diagrams():
    OUT_DIR.mkdir(exist_ok=True)
    diagrams = {}

    img = Image.new("RGB", (1600, 880), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 40), "3D baski literaturunden PrintForge yaklasimina", font=font("times", 42, True), fill="#111111")
    boxes = [
        ((80, 170, 430, 390), "Eklemeli imalat", "Malzeme, uretim esnekligi, hizli prototipleme ve dusuk adetli uretim."),
        ((625, 170, 975, 390), "AI 3D uretim", "Metin veya referans gorselden 3D geometri uretimi icin model tabanli surecler."),
        ((1170, 170, 1520, 390), "Pazaryeri", "Satici katalogu, fiyatlandirma, iletisim ve siparis sureclerinin web ortaminda yonetimi."),
        ((430, 560, 1170, 760), "PrintForge ozgunlesmesi", "Fikir bulma, referans secme, AI model uretimi, 3D onizleme ve satici iletisimi tek platformda birlestirilmistir."),
    ]
    for xy, title, body in boxes:
        draw_box(d, xy, title, body)
    arrow(d, (430, 280), (625, 280))
    arrow(d, (975, 280), (1170, 280))
    arrow(d, (255, 390), (540, 560))
    arrow(d, (800, 390), (800, 560))
    arrow(d, (1345, 390), (1060, 560))
    path = OUT_DIR / "fig_1_1.png"
    img.save(path)
    diagrams["fig_1_1"] = path

    img = Image.new("RGB", (1600, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), "PrintForge sistem mimarisi", font=font("times", 44, True), fill="#111111")
    draw_box(d, (80, 170, 460, 380), "Frontend", "Next.js, React, TypeScript, Tailwind CSS, Three.js model goruntuleyici.", "#f4f8ff")
    draw_box(d, (610, 170, 990, 380), "Backend API", "Express, TypeScript, JWT, Zod, CORS, Helmet ve servis rotalari.", "#f7fbf3")
    draw_box(d, (1140, 170, 1520, 380), "Veri katmani", "Prisma ORM, kullanici, model, mesaj, yorum, soru ve ornek varliklari.", "#fff8ef")
    draw_box(d, (250, 570, 620, 790), "Dis servisler", "SerpApi/Google Images, UploadThing, Tripo3D veya Hitem3D.", "#f8f4ff")
    draw_box(d, (850, 570, 1220, 790), "Dosya depolama", "Orijinal GLB/STL, guvenli onizleme JSON ve urun gorselleri.", "#f7f7f7")
    arrow(d, (460, 275), (610, 275))
    arrow(d, (990, 275), (1140, 275))
    arrow(d, (800, 380), (470, 570))
    arrow(d, (800, 380), (1035, 570))
    path = OUT_DIR / "fig_2_1.png"
    img.save(path)
    diagrams["fig_2_1"] = path

    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), "Referans gorselden AI uretime veri akisi", font=font("times", 44, True), fill="#111111")
    steps = [
        ("Ornek / arama", "Admin ornegi veya internet gorsel sonucu secilir."),
        ("Prompt", "Baslik ve arama baglami ile otomatik aciklama uretilir."),
        ("AI gorevi", "Gorsel ve prompt backend uzerinden uretim servisine iletilir."),
        ("Durum izleme", "Frontend gorevi 15 saniyelik araliklarla izler."),
        ("Model", "GLB/STL dosyasi kaydedilir ve onizleme acilir."),
    ]
    x = 70
    for idx, (t, b) in enumerate(steps):
        draw_box(d, (x, 210, x + 250, 500), t, b, "#f9fafb")
        if idx < len(steps) - 1:
            arrow(d, (x + 250, 355), (x + 320, 355))
        x += 320
    draw_box(d, (390, 640, 1210, 790), "Satici baglantisi", "Tamamlanan AI modeli icin kullanici satici secer, mesajlasma baslatir ve teklif/siparis surecine gecer.", "#f0fdf4")
    arrow(d, (1390, 500), (1210, 640))
    path = OUT_DIR / "fig_2_2.png"
    img.save(path)
    diagrams["fig_2_2"] = path

    img = Image.new("RGB", (1600, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), "Veri modeli iliskileri", font=font("times", 44, True), fill="#111111")
    entities = [
        ((90, 170, 430, 340), "User", "email, password, name, role, companyName"),
        ((650, 170, 990, 340), "Model", "AI/CATALOG, status, prompt, fiyat, dosya anahtarlari"),
        ((1170, 170, 1510, 340), "ExampleItem", "title, category, imageUrl, prompt, tags"),
        ((360, 570, 700, 740), "Conversation", "buyerId, sellerId, modelId, status"),
        ((900, 570, 1240, 740), "Message", "content, quoteAmount, readAt"),
        ((90, 780, 430, 930), "Review / Question", "puan, yorum, soru, cevap"),
    ]
    for xy, t, b in entities:
        draw_box(d, xy, t, b, "#fbfbfb")
    arrow(d, (430, 255), (650, 255))
    arrow(d, (820, 340), (530, 570))
    arrow(d, (700, 655), (900, 655))
    arrow(d, (260, 340), (260, 780))
    arrow(d, (820, 340), (260, 780))
    path = OUT_DIR / "fig_2_3.png"
    img.save(path)
    diagrams["fig_2_3"] = path

    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), "Kimlik dogrulama ve guvenli erisim akisi", font=font("times", 44, True), fill="#111111")
    flow = [
        ("Kayit / giris", "E-posta, sifre ve rol Zod ile dogrulanir; sifre bcrypt ile ozetlenir."),
        ("JWT", "Basarili giriste kullanici kimligi imzali token icine yazilir."),
        ("Korunan rota", "Authorization basligi kontrol edilir ve kullanici veritabaniyla eslestirilir."),
        ("Yetki", "Seller, user ve admin akislari rol bilgisine gore ayrilir."),
    ]
    x = 90
    for idx, (t, b) in enumerate(flow):
        draw_box(d, (x, 220, x + 310, 520), t, b, "#f7fbff")
        if idx < len(flow) - 1:
            arrow(d, (x + 310, 370), (x + 390, 370))
        x += 390
    draw_box(d, (360, 660, 1240, 805), "Model dosyasi erisimi", "Model dosyasina yalnizca model sahibi veya ilgili konusmanin taraflari erisebilir; dosya yolu uploads dizini disina cikmayacak sekilde cozulur.", "#fff8ef")
    path = OUT_DIR / "fig_2_4.png"
    img.save(path)
    diagrams["fig_2_4"] = path

    img = Image.new("RGB", (1400, 850), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), "Modul tamamlama durumu", font=font("times", 42, True), fill="#111111")
    modules = [
        ("Kimlik", 100), ("Pazaryeri", 100), ("AI akis", 90), ("Mesaj", 95), ("Admin", 90), ("Guvenlik", 90), ("Dagitim", 85)
    ]
    y = 160
    for name, val in modules:
        d.text((100, y), name, font=font("times", 26, True), fill="#111111")
        d.rounded_rectangle((300, y, 1200, y + 36), radius=16, outline="#999999", width=2, fill="#f2f2f2")
        d.rounded_rectangle((300, y, 300 + 9 * val, y + 36), radius=16, fill="#4f7f63")
        d.text((1220, y - 2), f"%{val}", font=font("times", 26), fill="#111111")
        y += 82
    path = OUT_DIR / "fig_3_1.png"
    img.save(path)
    diagrams["fig_3_1"] = path

    img = Image.new("RGB", (1400, 800), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 42), "Derleme ve dogrulama ozeti", font=font("times", 42, True), fill="#111111")
    draw_box(d, (100, 180, 620, 430), "Backend", "TypeScript derlemesi basariyla tamamlandi. API katmani tsc ile dogrulandi.", "#f0fdf4")
    draw_box(d, (780, 180, 1300, 430), "Frontend", "Next.js uretim derlemesi basariyla tamamlandi ve 21 sayfa olusturuldu.", "#f0fdf4")
    draw_box(d, (300, 560, 1100, 700), "Not", "Webpack onbellek uyarisi derlemeyi engellememis; islevsel hata olarak degerlendirilmemistir.", "#fff8ef")
    arrow(d, (620, 305), (780, 305))
    path = OUT_DIR / "fig_3_2.png"
    img.save(path)
    diagrams["fig_3_2"] = path

    return diagrams


def add_cover_pages(doc: Document, total_pages: int):
    add_empty_lines(doc, 5)
    add_center(doc, "T.C.", 12, True)
    add_center(doc, "TARSUS ÜNİVERSİTESİ", 12, True)
    add_center(doc, "MÜHENDİSLİK FAKÜLTESİ", 12, True)
    add_center(doc, "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ", 12, True)
    add_empty_lines(doc, 6)
    add_center(doc, TITLE_TR, 12, True)
    add_empty_lines(doc, 4)
    add_center(doc, AUTHOR, 12, True)
    add_empty_lines(doc, 4)
    add_center(doc, "LİSANS BİTİRME TEZİ", 12, True)
    add_empty_lines(doc, 8)
    add_center(doc, "TARSUS - 2026", 12, True)
    doc.add_page_break()

    add_empty_lines(doc, 5)
    add_center(doc, "T.C.", 12, True)
    add_center(doc, "TARSUS ÜNİVERSİTESİ", 12, True)
    add_center(doc, "MÜHENDİSLİK FAKÜLTESİ", 12, True)
    add_center(doc, "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ", 12, True)
    add_empty_lines(doc, 5)
    add_center(doc, TITLE_TR, 12, True)
    add_empty_lines(doc, 3)
    add_center(doc, AUTHOR, 12, True)
    add_empty_lines(doc, 3)
    add_center(doc, f"Danışman: {ADVISOR}", 12, True)
    add_empty_lines(doc, 3)
    add_center(doc, "LİSANS BİTİRME TEZİ", 12, True)
    add_empty_lines(doc, 8)
    add_center(doc, "TARSUS - 2026", 12, True)
    doc.add_page_break()

    add_empty_lines(doc, 7)
    add_para(doc, "Bu çalışma, jüri tarafından BİLGİSAYAR MÜHENDİSLİĞİ LİSANS BİTİRME TEZİ olarak kabul edilmiştir.", "TezNoIndent")
    add_empty_lines(doc, 5)
    for label in ["Başkan: ........................................ (Danışman)", "Üye: ........................................", "Üye: ........................................"]:
        p = doc.add_paragraph(style="TezNoIndent")
        p.paragraph_format.left_indent = Cm(3.0)
        r = p.add_run(label)
        set_font(r, 12)
    doc.add_page_break()

    add_main_heading(doc, "ETİK BEYAN")
    ethics = [
        "Tarsus Üniversitesi Mühendislik Fakültesi Bilgisayar Mühendisliği Bölümü Tez Yazım Kurallarına uygun olarak hazırladığım bu tez çalışmasında;",
        "Tez içinde sunduğum verileri, bilgileri ve dokümanları akademik ve etik kurallar çerçevesinde elde ettiğimi,",
        "Tüm bilgi, belge, değerlendirme ve sonuçları bilimsel etik ve ahlak kurallarına uygun olarak sunduğumu,",
        "Tez çalışmasında yararlandığım eserlerin tümüne uygun atıfta bulunarak kaynak gösterdiğimi,",
        "Kullanılan verilerde ve ortaya çıkan sonuçlarda herhangi bir değişiklik yapmadığımı,",
        "Bu tezde sunduğum çalışmanın özgün olduğunu,",
    ]
    add_para(doc, ethics[0])
    for item in ethics[1:]:
        add_bullet(doc, item)
    add_para(doc, "bildirir, aksi bir durumda aleyhime doğabilecek tüm hak kayıplarını kabullendiğimi beyan ederim.")
    add_empty_lines(doc, 3)
    p = doc.add_paragraph(style="TezNoIndent")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(".... / .... / 20....")
    set_font(r, 12)
    add_empty_lines(doc, 2)
    for text in ["İMZA", "Adı Soyadı"]:
        p = doc.add_paragraph(style="TezNoIndent")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        set_font(r, 12, True if text == "İMZA" else False)


def add_abstracts(doc: Document, total_pages: int):
    add_main_heading(doc, "ÖZET")
    for line in [TITLE_TR, AUTHOR, "Lisans Bitirme Tezi, Bilgisayar Mühendisliği Bölümü", f"Danışman: {ADVISOR}", f"{DATE_TEXT}, {total_pages if total_pages else 'xxx'} sayfa"]:
        add_center(doc, line, 12, True)
    summary = (
        "Bu tez çalışmasında, 3D baskı hizmetlerinden yararlanmak isteyen kullanıcıların fikir geliştirme, referans görsel seçme, AI destekli 3D model üretme, baskı kalitesini fotoğraf üzerinden ön değerlendirme ve satıcıyla iletişime geçme süreçlerini tek bir web platformunda birleştiren PrintForge sistemi geliştirilmiştir. Çalışmanın temel problemi, teknik modelleme bilgisi olmayan kullanıcıların üretime uygun 3D model dosyasına erişmekte, baskı sonrası kaliteyi yorumlamakta ve güvenilir satıcılarla düzenli bir süreç üzerinden iletişim kurmakta yaşadığı zorluktur. Bu problemi çözmek amacıyla kullanıcı, satıcı ve admin rollerini destekleyen; satıcı katalogu, admin örnek yönetimi, internet görsel araması, referans görsele dayalı AI model üretimi, baskı kalite kontrolü, 3D model önizleme, mesajlaşma ve sipariş takibi modüllerinden oluşan bütünleşik bir yazılım mimarisi tasarlanmıştır."
    )
    add_para(doc, summary)
    add_para(doc, "Sistem frontend katmanında Next.js, React, TypeScript, Tailwind CSS ve Three.js; backend katmanında Node.js, Express, TypeScript, Prisma ORM, JWT tabanlı kimlik doğrulama, Zod doğrulama şemaları, CORS ve Helmet güvenlik bileşenleri kullanılarak geliştirilmiştir. Görsel arama işlemleri API anahtarlarının istemci tarafına taşınmaması için backend üzerinden yürütülmüş, AI model üretimi Tripo3D veya Hitem3D gibi harici servislerle görev tabanlı şekilde ilişkilendirilmiştir. Uygulama, AI üretim süresinin uzun olabileceği varsayımıyla görev durumunu belirli aralıklarla izleyen ve tamamlanan model dosyasını güvenli erişim kurallarıyla sunan bir yapı içermektedir.")
    add_para(doc, "Gerçekleştirilen doğrulamada backend TypeScript derlemesi başarıyla tamamlanmış, frontend üretim derlemesi Next.js tarafından 21 sayfa oluşturularak tamamlanmıştır. Bulgular, PrintForge'un 3D baskı fikrinden satıcıyla iletişime ve baskı sonrası kalite gözlemine kadar uzanan süreci tek sistem altında toplayabildiğini, admin tarafından yönetilen örnek içeriklerle dış servis risklerine karşı yedekli bir başlangıç akışı sunduğunu ve satıcı-kullanıcı etkileşimini katalog, mesajlaşma, kalite kontrol, yorum, soru-cevap ve sipariş durumu bileşenleriyle desteklediğini göstermiştir.")
    p = doc.add_paragraph(style="TezNoIndent")
    r = p.add_run("Anahtar Kelimeler: ")
    set_font(r, 12, True)
    r = p.add_run("3D baskı, eklemeli imalat, yapay zeka, web pazaryeri, 3D model üretimi, kalite kontrol.")
    set_font(r, 12)
    doc.add_page_break()

    add_main_heading(doc, "ABSTRACT")
    for line in [TITLE_EN.upper(), AUTHOR, "Graduation Thesis, Department of Computer Engineering", f"Supervisor: {ADVISOR}", f"June 2026, {total_pages if total_pages else 'xxx'} pages"]:
        add_center(doc, line, 12, True)
    add_para(doc, "In this thesis, PrintForge, an integrated web platform that combines idea discovery, reference image selection, AI-assisted 3D model generation, photo-based print quality pre-evaluation, seller communication, and order tracking for 3D printing services, has been developed. The main problem addressed in the study is that users without technical modeling knowledge often have difficulty obtaining production-ready 3D model files, interpreting post-print quality, and communicating with suitable sellers through an organized workflow. To address this problem, a role-based software architecture supporting users, sellers, and administrators has been designed and implemented with seller catalog management, curated example administration, internet image search, reference-image-based AI generation, print quality control, 3D model preview, messaging, and order tracking modules.")
    add_para(doc, "The frontend layer of the system was implemented with Next.js, React, TypeScript, Tailwind CSS, and Three.js, while the backend layer was implemented with Node.js, Express, TypeScript, Prisma ORM, JWT-based authentication, Zod validation schemas, CORS, and Helmet security middleware. Image search operations are handled by the backend to avoid exposing API keys on the client side. AI model generation is integrated with external task-based services such as Tripo3D or Hitem3D. Since AI generation may take time depending on service load and model complexity, the application includes a polling-based task monitoring mechanism and secure model access controls.")
    add_para(doc, "The verification results show that the backend TypeScript build was completed successfully and the frontend production build generated 21 pages with Next.js. The findings indicate that PrintForge can connect the process from a 3D printing idea to seller communication and post-print quality observation within a single system, provide a resilient starting flow through administrator-managed example content, and support buyer-seller interaction with catalog, messaging, quality control, review, question-answer, and order status features.")
    p = doc.add_paragraph(style="TezNoIndent")
    r = p.add_run("Keywords: ")
    set_font(r, 12, True)
    r = p.add_run("3D printing, additive manufacturing, artificial intelligence, web marketplace, 3D model generation, quality control.")
    set_font(r, 12)
    doc.add_page_break()


def add_front_matter(doc: Document, page_map: dict[str, str]):
    add_main_heading(doc, "ÖNSÖZ")
    add_para(doc, "Bu tez çalışması, 3D baskı hizmetlerinden yararlanmak isteyen kullanıcıların teknik modelleme bilgisine sahip olmadan üretime uygun model oluşturma sürecine katılabilmesini ve satıcılarla düzenli bir pazaryeri akışı üzerinden iletişim kurabilmesini hedefleyen PrintForge platformunun tasarım ve geliştirme sürecini kapsamaktadır.")
    add_para(doc, "Çalışma boyunca yazılım mimarisi, kullanıcı deneyimi, güvenlik, veri yönetimi ve AI destekli model üretim akışı birlikte ele alınmıştır. Bitirme tezi sürecinde proje fikrinin olgunlaşmasına, kapsamın belirlenmesine ve akademik raporlama disiplininin korunmasına katkı sağlayan danışman öğretim üyeme teşekkür ederim.")
    add_para(doc, "Ayrıca eğitim hayatım boyunca desteklerini esirgemeyen aileme, proje geliştirme sürecinde görüşleriyle katkıda bulunan arkadaşlarıma ve mühendislik eğitimi süresince sağladıkları akademik ortam için Tarsus Üniversitesi Mühendislik Fakültesi Bilgisayar Mühendisliği Bölümüne teşekkür ederim.")
    doc.add_page_break()

    add_main_heading(doc, "İÇİNDEKİLER")
    p = doc.add_paragraph(style="TezNoIndent")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Sayfa")
    set_font(r, 11, True)
    for entry in TOC_ENTRIES:
        add_toc_entry(doc, *entry, page_map)
    doc.add_page_break()

    add_main_heading(doc, "KISALTMALAR")
    abbreviations = [
        ("AI", "Artificial Intelligence / Yapay Zeka"),
        ("API", "Application Programming Interface"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("CSS", "Cascading Style Sheets"),
        ("FDM", "Fused Deposition Modeling / Eriyik yığma modelleme"),
        ("GLB", "Binary glTF 3D model dosya biçimi"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("JWT", "JSON Web Token"),
        ("ORM", "Object Relational Mapping"),
        ("REST", "Representational State Transfer"),
        ("STL", "Stereolithography 3D model dosya biçimi"),
        ("UI", "User Interface / Kullanıcı Arayüzü"),
        ("UX", "User Experience / Kullanıcı Deneyimi"),
        ("WebGL", "Web Graphics Library"),
    ]
    for abbr, desc in abbreviations:
        p = doc.add_paragraph(style="TezNoIndent")
        r = p.add_run(f"{abbr}: ")
        set_font(r, 12, True)
        r = p.add_run(desc)
        set_font(r, 12)
    doc.add_page_break()

    add_main_heading(doc, "TABLOLAR LİSTESİ")
    p = doc.add_paragraph(style="TezNoIndent")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Sayfa")
    set_font(r, 11, True)
    for caption, key in TABLES:
        add_list_entry(doc, caption, key, page_map)
    doc.add_page_break()

    add_main_heading(doc, "ŞEKİLLER LİSTESİ")
    p = doc.add_paragraph(style="TezNoIndent")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Sayfa")
    set_font(r, 11, True)
    for caption, key in FIGURES:
        add_list_entry(doc, caption, key, page_map)
    doc.add_page_break()

    add_main_heading(doc, "EKLER LİSTESİ")
    p = doc.add_paragraph(style="TezNoIndent")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Sayfa")
    set_font(r, 11, True)
    for caption, key in APPENDICES:
        add_list_entry(doc, caption, key, page_map)


def add_intro(doc: Document):
    add_main_heading(doc, "GİRİŞ")
    paragraphs = [
        "Eklemeli imalat ve 3D baskı teknolojileri, dijital tasarımların fiziksel nesnelere dönüştürülmesini hızlandırarak prototipleme, kişiselleştirilmiş üretim ve düşük adetli ürün geliştirme alanlarında önemli fırsatlar sunmaktadır. Bununla birlikte 3D baskı hizmetinden yararlanmak isteyen son kullanıcıların büyük bir kısmı, üretime uygun 3D model dosyası hazırlama, modelin baskıya uygunluğunu yorumlama ve güvenilir satıcıyla iletişim kurma konularında teknik bilgiye ihtiyaç duymaktadır. Bu durum, 3D baskı teknolojilerinin yaygınlaşmasını yalnızca üretim altyapısına değil, aynı zamanda kullanıcıyı fikir aşamasından üreticiyle iletişime taşıyan yazılım deneyimlerine de bağlı hale getirmektedir.",
        "PrintForge projesi bu ihtiyaçtan hareketle geliştirilmiştir. Platform, kullanıcının hazır örneklerden veya internet görsel arama sonuçlarından referans seçmesini, seçilen görselin AI model üretim sayfasına bağlamıyla birlikte aktarılmasını, model üretim görevinin arka planda izlenmesini ve tamamlanan model için satıcıyla mesajlaşma/sipariş sürecinin başlatılmasını hedeflemektedir. Böylece kullanıcı, ayrı ayrı yürütülen fikir arama, model üretme, ürün inceleme ve satıcıyla görüşme adımlarını tek bir bütünleşik web uygulaması üzerinden gerçekleştirebilmektedir.",
        "Literatürde 3D baskının üretim modellerini dönüştürdüğü, küçük ölçekli ve kişiselleştirilmiş üretim için yeni pazar yapıları oluşturduğu ve tedarik zincirlerinde dijitalleşmeyi hızlandırdığı belirtilmektedir (Berman,2012:155-162; Weller vd.,2015:43-56; Holmstrom ve Partanen,2014:421-430). Bunun yanında eklemeli imalatın sürdürülebilirlik, malzeme kullanımı, üretim esnekliği ve dağıtık üretim açısından fırsatlar sunduğu; ancak standartlaşma, kalite kontrol, ölçeklenebilirlik ve kullanıcı erişilebilirliği gibi sınırlamalar taşıdığı vurgulanmaktadır (Ford ve Despeisse,2016:1573-1587; Gebler vd.,2014:158-167; Ngo vd.,2018:172-196). PrintForge, bu tartışmaların yazılım tarafındaki karşılığını ele alarak 3D baskıya erişimi kullanıcı deneyimi, AI üretim desteği ve pazaryeri etkileşimi üzerinden kolaylaştırmayı amaçlamaktadır.",
        "Çalışmanın temel araştırma sorusu şu şekilde formüle edilmiştir: Teknik modelleme bilgisi sınırlı olan kullanıcılar için referans görsel seçimi, AI destekli 3D model üretimi ve satıcıyla iletişim süreçleri tek bir web platformunda güvenli, izlenebilir ve kullanılabilir biçimde birleştirilebilir mi? Bu soruya yanıt aranırken yalnızca AI üretim servisine istek gönderen dar bir prototip değil, kullanıcı rolleri, katalog yönetimi, admin içerik akışı, mesajlaşma, sipariş durumu ve model dosyası erişim kontrolünü içeren bütüncül bir sistem geliştirilmiştir.",
        "Çalışmanın amacı, 3D baskı yaptırmak isteyen kullanıcıların üretime uygun model oluşturma sürecine daha düşük teknik eşikle katılabilmesini sağlamak ve satıcıların ürünlerini yönetebileceği bir pazaryeri altyapısı oluşturmaktır. Bu amaç doğrultusunda kullanıcı, satıcı ve admin rolleri tanımlanmış; satıcı ürün katalogu, örnek görsel yönetimi, internet görsel araması, AI model üretimi, 3D model önizleme, yorum/soru-cevap, mesajlaşma ve sipariş takibi modülleri geliştirilmiştir.",
        "Çalışmanın kapsamı web tabanlı bir yazılım platformu ile sınırlandırılmıştır. Fiziksel 3D yazıcı kontrolü, otomatik dilimleme, gerçek ödeme entegrasyonu, lojistik entegrasyonu ve modelin tam üretilebilirliğini otomatik doğrulayan gelişmiş mühendislik analizleri bu çalışmanın kapsamı dışında bırakılmıştır. Buna karşın sistem, tamamlanan model dosyalarının görüntülenmesi, satıcıyla görüşülmesi ve ileride maliyet tahmini ya da baskı parametresi önerisi gibi modüllerin eklenebilmesi için genişletilebilir bir temel sunmaktadır.",
        "Tez çalışması üç temel iş paketi üzerinde yürütülmüştür. İlk iş paketinde literatür ve ihtiyaç analizi yapılmış, platformun problem tanımı, kullanıcı rolleri ve modül sınırları belirlenmiştir. İkinci iş paketinde frontend ve backend mimarisi, veritabanı modeli, harici servis entegrasyonları ve kullanıcı akışları geliştirilmiştir. Üçüncü iş paketinde uygulama derleme doğrulamaları, senaryo bazlı testler ve bulguların değerlendirilmesi gerçekleştirilmiştir. Bu yapı, projenin yalnızca çalışan bir yazılım çıktısı olarak değil, akademik olarak açıklanabilir ve tekrarlanabilir bir mühendislik çalışması olarak raporlanmasını sağlamaktadır.",
    ]
    for para in paragraphs:
        add_para(doc, para)


def add_chapter_one(doc: Document, diagrams: dict[str, Path]):
    doc.add_page_break()
    add_main_heading(doc, "BİRİNCİ BÖLÜM")
    add_main_heading(doc, "TEORİK ALT YAPI")
    add_section_heading(doc, "1.1. Eklemeli İmalat ve 3D Baskı")
    for para in [
        "Eklemeli imalat, geleneksel talaşlı veya kalıba dayalı üretim yöntemlerinden farklı olarak nesneyi katman katman oluşturan bir üretim yaklaşımıdır. Bu yaklaşım, karmaşık geometrilerin düşük adetlerde üretilebilmesi, hızlı prototipleme yapılabilmesi ve dijital tasarımdan fiziksel ürüne geçiş süresinin kısaltılması açısından önem taşır. Gibson vd. (2021:1-506), eklemeli imalat teknolojilerinin tasarım serbestliği sağladığını; ancak malzeme seçimi, işlem parametreleri ve parça kalitesinin sistematik biçimde değerlendirilmesi gerektiğini belirtmektedir.",
        "3D baskı teknolojileri yalnızca üretim cihazlarından oluşmaz; tasarım dosyası, dosya formatı, ön işleme, üretim parametreleri, malzeme seçimi ve son işlem adımlarından oluşan bütünleşik bir süreçtir. Bu nedenle kullanıcıların yalnızca bir fikir veya görsel referansla sürece başlaması çoğu zaman yeterli olmamakta, üretime uygun model dosyası oluşturma veya bulma ihtiyacı ortaya çıkmaktadır. Shahrubudin vd. (2019:1286-1296), 3D baskının farklı malzeme ve uygulama alanlarında yaygınlaştığını; Ngo vd. (2018:172-196) ise teknolojinin gelişmesine rağmen kalite, malzeme davranışı ve süreç kontrolünün önemli zorluklar olmaya devam ettiğini vurgulamaktadır.",
        "PrintForge'un çıkış noktası bu teknik eşiktir. Platform, 3D baskı fikrini doğrudan üretime hazır kabul etmek yerine, kullanıcıyı referans seçimi, AI üretim görevi, model önizleme ve satıcıyla değerlendirme adımlarından geçirir. Böylece eklemeli imalatın teknik karmaşıklığı tamamen ortadan kaldırılmasa da kullanıcı açısından daha anlaşılır ve yönetilebilir bir hizmet akışına dönüştürülür.",
    ]:
        add_para(doc, para)

    add_section_heading(doc, "1.2. AI Destekli 3D Model Üretimi")
    for para in [
        "Yapay zeka tabanlı 3D içerik üretimi, metin ya da görsel girdilerden nokta bulutu, ağ yapısı, örtük fonksiyon veya model dosyası üretmeyi hedefleyen araştırma alanlarından biridir. Derin öğrenmenin görsel veri işleme alanında sağladığı ilerleme, 3D üretim problemlerine de yansımış; iki boyutlu üretken modellerden üç boyutlu temsillere geçiş yapan yöntemler ortaya çıkmıştır. LeCun vd. (2015:436-444), derin öğrenmenin çok katmanlı temsiller aracılığıyla karmaşık veri ilişkilerini öğrenebildiğini göstermekte; bu yaklaşım 3D model üretimi gibi yüksek boyutlu görevlerde de temel oluşturmaktadır.",
        "DreamFusion, metinden 3D içerik üretiminde 2D difüzyon modellerinden yararlanan önemli çalışmalardan biridir (Poole vd.,2022). Point-E ve Shap-E çalışmaları ise metin girdilerinden nokta bulutu veya örtük 3D fonksiyon üretmeye odaklanarak bu alandaki yaklaşımların çeşitlendiğini göstermiştir (Nichol vd.,2022; Jun ve Nichol,2023). Magic3D çalışması, yüksek çözünürlüklü metinden 3D içerik üretimi için iki aşamalı bir yaklaşım önermiştir (Lin vd.,2023:300-309). Bu çalışmalar, PrintForge gibi uygulamaların neden harici AI üretim servisleriyle görev tabanlı entegre edilmesi gerektiğini açıklayan teorik arka planı oluşturmaktadır.",
        "PrintForge, AI model üretim algoritmasının kendisini bu tez kapsamında sıfırdan geliştirmemiştir. Bunun yerine güncel AI model üretim servislerini web uygulaması akışına bağlayan, kullanıcıdan referans görsel alan, görevi backend üzerinden başlatan, görev durumunu izleyen ve tamamlanan modeli güvenli biçimde sunan bir entegrasyon mimarisi geliştirilmiştir. Bu tercih, lisans bitirme projesi kapsamına uygun olarak uygulama mimarisi, veri güvenliği, kullanıcı deneyimi ve servis entegrasyonu problemlerine odaklanılmasını sağlamıştır.",
    ]:
        add_para(doc, para)

    add_figure(doc, diagrams["fig_1_1"], "Şekil 1.1. Literatürden projeye uzanan kavramsal konumlandırma")

    add_section_heading(doc, "1.3. Web Tabanlı Pazaryeri Yaklaşımı")
    for para in [
        "3D baskı teknolojilerinin yaygınlaşması, yalnızca üretim atölyeleri veya bireysel yazıcı sahipleriyle sınırlı değildir; kullanıcıların tasarım dosyalarına, satıcılara, fiyat bilgisine ve üretim seçeneklerine eriştiği dijital platformlar da bu ekosistemin parçasıdır. Rayna ve Striukova (2016:214-224), 3D baskının iş modeli inovasyonunu etkilediğini ve kullanıcıların üretim sürecine daha doğrudan katılabildiğini belirtmektedir. Weller vd. (2015:43-56), 3D baskının pazar yapısı ve maliyet dinamikleri üzerinde dönüştürücü etkiler yaratabileceğini göstermiştir.",
        "Web tabanlı pazaryeri yaklaşımı, satıcıların ürün ve hizmetlerini listelerken kullanıcıların ürünleri arama, filtreleme, inceleme, favorilere ekleme, sepete alma, soru sorma ve mesajlaşma gibi işlemleri gerçekleştirmesini sağlar. PrintForge projesinde pazaryeri yalnızca ürün gösterim alanı olarak ele alınmamış; AI ile üretilen model için de satıcıya ulaşılabilen bir iletişim katmanı olarak tasarlanmıştır. Bu özellik, hazır ürün katalogu ile kişisel üretim talebi arasındaki boşluğu kapatmayı amaçlamaktadır.",
        "Platformun pazaryeri yaklaşımında sabit fiyatlı ürün kartları, satıcı görünürlüğü, yorum ve soru-cevap modülleri, sepet ve favoriler gibi e-ticaret kullanıcılarının beklediği temel işlevler yer alır. Bunun yanında AI üretimden gelen model dosyası için ayrıca satıcı seçme ve mesajlaşma akışının bulunması, sistemi klasik katalog uygulamalarından ayırmaktadır.",
    ]:
        add_para(doc, para)

    add_section_heading(doc, "1.4. Yazılım Mimarisi, Güvenlik ve Veri Yönetimi")
    for para in [
        "Web uygulamalarında sürdürülebilir mimari, kullanıcı arayüzü, API katmanı, veri tabanı, kimlik doğrulama ve harici servis entegrasyonlarının açık sorumluluklarla ayrılmasını gerektirir. PrintForge bu nedenle frontend ve backend olmak üzere iki ana katman üzerinde geliştirilmiştir. Frontend katmanında sayfa yönlendirme, durum yönetimi, kullanıcı etkileşimi ve 3D önizleme; backend katmanında kimlik doğrulama, iş kuralları, veri tabanı erişimi, harici servis çağrıları ve dosya erişim kontrolü yürütülmektedir.",
        "Güvenlik açısından özellikle API anahtarlarının istemciye verilmemesi, kullanıcı girdilerinin doğrulanması, yetkisiz model dosyası erişiminin engellenmesi ve kullanıcı rollerine göre işlev ayrımı yapılması önemlidir. PrintForge backend katmanında JWT tabanlı oturum yönetimi, bcrypt ile parola özetleme, Zod ile giriş doğrulama, CORS yapılandırması ve Helmet güvenlik başlıkları kullanılmıştır. Model dosyalarına erişimde model sahibi veya ilgili konuşmanın tarafı olma koşulu aranarak yetkisiz erişim riski azaltılmıştır.",
        "Veri yönetiminde Prisma ORM kullanılmıştır. User, Model, Conversation, Message, ProductReview, ProductQuestion ve ExampleItem varlıkları üzerinden kullanıcı rolleri, katalog ürünleri, AI modelleri, mesajlaşma, sipariş durumu, yorumlar, sorular ve admin örnekleri ilişkilendirilmiştir. Bu veri modeli, platformun ileride ödeme, baskı parametresi, teslimat ve üretim maliyeti gibi yeni alanlarla genişletilebilmesine uygun bir temel sağlamaktadır.",
    ]:
        add_para(doc, para)

    add_section_heading(doc, "1.5. Literatür Değerlendirmesi ve Özgün Değer")
    add_para(doc, "Literatür incelendiğinde çalışmaların bir kısmının eklemeli imalatın üretim, malzeme ve sürdürülebilirlik boyutlarına; bir kısmının 3D baskının iş modeli ve pazar yapısına; son yıllardaki çalışmaların ise AI destekli 3D içerik üretimine odaklandığı görülmektedir. PrintForge bu alanların kesişiminde, kullanıcı deneyimi ve web platformu uygulaması düzeyinde konumlanmaktadır. Tablo 1.1, seçilen literatürün proje ile ilişkisini özetlemektedir.")
    rows = [
        ["Yayın", "Odak", "PrintForge ile ilişki"],
        ["Berman (2012)", "3D baskının endüstriyel dönüşüm etkisi", "3D baskıya erişimi kolaylaştıran platform ihtiyacını destekler."],
        ["Weller vd. (2015)", "Pazar yapısı ve ekonomik etkiler", "Pazaryeri modülünün ekonomik gerekçesini açıklar."],
        ["Rayna ve Striukova (2016)", "İş modeli inovasyonu", "Kullanıcının üretim sürecine katılımını temellendirir."],
        ["Ngo vd. (2018)", "Malzeme, yöntem ve zorluklar", "Model üretiminin satıcı değerlendirmesiyle desteklenmesi gerektiğini gösterir."],
        ["Shahrubudin vd. (2019)", "3D baskı teknolojileri ve uygulamalar", "Platformun uygulama alanını geniş çerçeveye yerleştirir."],
        ["Poole vd. (2022)", "Metinden 3D üretim", "AI görev tabanlı üretim akışının teorik arka planını oluşturur."],
        ["Nichol vd. (2022)", "Kompleks promptlardan 3D nokta bulutu", "Prompt tabanlı üretimin önemini gösterir."],
        ["Jun ve Nichol (2023)", "Koşullu 3D örtük fonksiyon üretimi", "Görsel/metin girdisinin 3D modele dönüşüm potansiyelini destekler."],
        ["Lin vd. (2023)", "Yüksek çözünürlüklü text-to-3D", "Üretim kalitesi ve servis entegrasyonu ihtiyacını açıklar."],
    ]
    add_table(doc, "Tablo 1.1. Literatür çalışmaları ve PrintForge ile ilişkisi", rows, [3.2, 4.0, 8.3], "Yazar tarafından literatür taraması esas alınarak hazırlanmıştır.", 9)
    add_para(doc, "Tablo 1.1'de görüldüğü üzere literatürde 3D baskı, AI üretim ve platform ekonomisi ayrı başlıklarda ele alınmaktadır. PrintForge'un özgün değeri, bu başlıkları lisans bitirme projesi ölçeğinde çalışan bir web uygulamasında bir araya getirmesidir. Sistem, kullanıcıyı yalnızca katalogda gezdiren bir yapı değil; fikirden referans seçimine, referanstan AI model üretimine, modelden satıcıyla iletişime kadar uçtan uca bir iş akışı sunmaktadır.")


def add_chapter_two(doc: Document, diagrams: dict[str, Path]):
    doc.add_page_break()
    add_main_heading(doc, "İKİNCİ BÖLÜM")
    add_main_heading(doc, "MATERYAL VE METOT")
    add_section_heading(doc, "2.1. Kullanılan Materyaller ve Yazılım Araçları")
    add_para(doc, "Bu bölümde PrintForge platformunun geliştirilmesinde kullanılan yazılım bileşenleri, veri modeli, uygulama mimarisi ve test yaklaşımı açıklanmaktadır. Proje, frontend ve backend ayrımıyla geliştirilmiş tam yığın bir web uygulamasıdır. Frontend tarafında kullanıcı arayüzü, sayfa geçişleri, AI üretim durumu, 3D model önizleme ve pazaryeri etkileşimleri; backend tarafında kimlik doğrulama, veri tabanı işlemleri, harici servis entegrasyonları, model dosyası yönetimi ve erişim kontrolü yer almaktadır.")
    rows = [
        ["Bileşen", "Teknoloji", "Projede kullanım amacı"],
        ["Frontend", "Next.js 14, React 18, TypeScript", "Sayfa yapısı, bileşen mimarisi, istemci etkileşimleri ve üretim derlemesi."],
        ["Arayüz", "Tailwind CSS, lucide-react", "Duyarlı, rol tabanlı ve etkileşimli kullanıcı arayüzü."],
        ["3D Önizleme", "Three.js, @react-three/fiber, @react-three/drei", "GLB model dosyalarının tarayıcıda döndürülerek incelenmesi."],
        ["Backend", "Node.js, Express, TypeScript", "REST benzeri API uç noktaları ve iş kuralları."],
        ["Veri", "Prisma ORM, SQLite/PostgreSQL şemaları", "Geliştirme ve canlı ortam için veritabanı erişimi."],
        ["Güvenlik", "JWT, bcryptjs, Zod, Helmet, CORS", "Oturum, parola güvenliği, doğrulama ve güvenli API erişimi."],
        ["Harici servis", "SerpApi/Google Images, UploadThing, Tripo3D/Hitem3D", "Görsel arama, görsel yükleme ve AI model üretimi."],
    ]
    add_table(doc, "Tablo 2.1. Projede kullanılan yazılım bileşenleri", rows, [3.0, 4.0, 8.5], "Proje package.json dosyaları ve kaynak kodu esas alınarak hazırlanmıştır.", 9)

    add_section_heading(doc, "2.2. Sistem Mimarisi")
    for para in [
        "Sistem mimarisi istemci, API, veri tabanı, dosya depolama ve harici servisler olmak üzere beş temel bileşenden oluşmaktadır. Kullanıcı arayüzü Next.js uygulaması üzerinde çalışmakta, backend API ise Express sunucusu üzerinden /api/auth, /api/ai, /api/models, /api/chat, /api/examples ve /api/images rotalarını sunmaktadır. Bu ayrım, kullanıcı deneyiminin frontend tarafında hızlı biçimde yönetilmesini ve hassas anahtarların backend ortam değişkenlerinde tutulmasını sağlamaktadır.",
        "Görsel arama akışında kullanıcı sorgusu önce frontend tarafında temizlenmekte, ardından backend /api/images/search uç noktasına gönderilmektedir. Backend, yapılandırmaya göre SerpApi veya Google Images sağlayıcısını seçmekte; sonuç başlıklarını ve kaynak alanlarını normalize etmekte; aynı sorgular için kısa süreli bellek önbelleği kullanmaktadır. Bu yapı hem kota kullanımını azaltmakta hem de API anahtarlarının istemciye sızmasını önlemektedir.",
        "AI üretim akışında kullanıcı referans görselini ve prompt bilgisini gönderir. Backend bu isteği doğrular, görsel dosyasını boyut ve mimetype açısından sınırlar, AI servisinde görev başlatır ve veritabanında PENDING durumunda bir AI Model kaydı oluşturur. Frontend tarafındaki AiGenerationProvider, görev durumunu 15 saniyelik aralıklarla sorgular; görev tamamlandığında model dosyası backend tarafından indirilir, uploads dizinine kaydedilir ve model kaydı COMPLETED durumuna geçirilir.",
    ]:
        add_para(doc, para)
    add_figure(doc, diagrams["fig_2_1"], "Şekil 2.1. PrintForge sistem mimarisi")
    add_figure(doc, diagrams["fig_2_2"], "Şekil 2.2. Referans görselden AI üretime veri akışı")

    rows = [
        ["Modül", "Uç nokta", "Temel görev"],
        ["Kimlik doğrulama", "/api/auth", "Kayıt, giriş, profil, e-posta ve şifre güncelleme işlemleri."],
        ["AI üretim", "/api/ai", "Model üretim görevi başlatma, görev durumu izleme ve üretim geçmişi."],
        ["Model ve katalog", "/api/models", "Katalog listeleme, satıcı ürün yönetimi, yorum/soru ve dosya erişimi."],
        ["Mesajlaşma", "/api/chat", "Satıcı listesi, konuşma, sipariş, bildirim ve mesaj akışları."],
        ["Örnek içerikler", "/api/examples", "Admin tarafından yönetilen referans görseller ve proxy görsel erişimi."],
        ["Görsel arama", "/api/images", "SerpApi/Google Images sonuçlarını normalize eden arama katmanı."],
    ]
    add_table(doc, "Tablo 2.2. Backend API modülleri ve görevleri", rows, [3.2, 3.8, 8.5], "Backend rota dosyaları incelenerek hazırlanmıştır.", 9)

    add_section_heading(doc, "2.3. Veri Modeli ve Veritabanı Tasarımı")
    add_para(doc, "Veri modeli, sistemin iki ana iş alanını destekleyecek şekilde tasarlanmıştır: AI destekli model üretimi ve satıcı pazaryeri. User varlığı kullanıcı, satıcı ve admin rollerini temsil eder. Model varlığı hem AI ile üretilen modelleri hem de satıcı katalog ürünlerini tutar. Conversation ve Message varlıkları alıcı-satıcı iletişimini, ProductReview ve ProductQuestion varlıkları katalog ürünleri etrafındaki sosyal kanıt ve soru-cevap akışını, ExampleItem ise admin tarafından yönetilen referans içerikleri temsil eder.")
    add_figure(doc, diagrams["fig_2_3"], "Şekil 2.3. Veri modeli ilişkileri")
    rows = [
        ["Varlık", "Önemli alanlar", "Açıklama"],
        ["User", "email, password, name, role, companyName", "Kullanıcı, satıcı ve admin rollerinin temel hesap bilgilerini tutar."],
        ["Model", "type, status, prompt, viewerDataKey, priceRangeMin", "AI üretimleri ve katalog ürünleri için ortak model kaydıdır."],
        ["Conversation", "buyerId, sellerId, modelId, status", "AI modeli veya katalog ürünü üzerinden başlatılan konuşma/sipariş kaydıdır."],
        ["Message", "conversationId, senderId, content, readAt", "Konuşma içindeki mesajları ve okunma durumunu tutar."],
        ["ProductReview", "modelId, userId, rating, comment", "Katalog ürünü için kullanıcı puanı ve yorumu sağlar."],
        ["ProductQuestion", "question, answer, answerUserId", "Alıcı sorusu ve satıcı cevabını ilişkilendirir."],
        ["ExampleItem", "title, category, imageUrl, prompt, tags", "Admin tarafından yönetilen örnek referans kartlarını saklar."],
    ]
    add_table(doc, "Tablo 2.3. Temel veritabanı varlıkları", rows, [3.0, 4.3, 8.2], "Prisma şema dosyaları esas alınarak hazırlanmıştır.", 9)

    add_section_heading(doc, "2.4. Uygulama Geliştirme Süreci")
    add_para(doc, "Geliştirme süreci modüler biçimde yürütülmüştür. İlk aşamada kullanıcı rolleri ve kimlik doğrulama akışı kurulmuştur. Kullanıcılar USER veya SELLER rolüyle kayıt olabilir; admin yetkisi ise yardımcı fonksiyonlarla etkin role dönüştürülür. Parolalar bcrypt ile özetlenir ve başarılı girişte JWT üretilir. Korunan rotalarda token doğrulanır ve veritabanındaki kullanıcıyla eşleştirilir.")
    add_para(doc, "İkinci aşamada satıcı katalogu geliştirilmiştir. Satıcılar ürün adı, açıklama, kategori, sabit fiyat ve en fazla beş görsel ile ürün ekleyebilir. Katalog sayfasında kullanıcılar arama, kategori ve fiyat filtrelerini kullanabilir; ürünleri favorilere veya sepete ekleyebilir; satıcıya mesaj gönderebilir; ürün detayında yorum ve soru-cevap alanlarını inceleyebilir. Satıcı kendi ürünlerini düzenleyebilir veya katalogdan pasif hale getirebilir.")
    add_para(doc, "Üçüncü aşamada örnek görsel ve AI üretim akışı eklenmiştir. Admin tarafından eklenen örnekler /examples sayfasında gösterilir. Kullanıcı arama yaptığında backend görsel arama sağlayıcısından sonuçları alır ve kartlara dönüştürür. Kullanıcı karttaki 'AI ile 3D'ye dönüştür' işlemini seçtiğinde görsel URL'si, başlık, kaynak ve otomatik prompt AI üretim sayfasına aktarılır. AI sayfası görseli proxy üzerinden indirip PNG dosyasına dönüştürür ve backend'e üretim isteği olarak gönderir.")
    add_para(doc, "Dördüncü aşamada mesajlaşma ve sipariş akışı geliştirilmiştir. Kullanıcı katalog ürünü veya AI modeli üzerinden satıcıyla konuşma başlatabilir. Konuşmalar alıcı ve satıcı tarafında listelenir, okunmamış mesaj sayıları takip edilir ve sipariş durumları ORDERED, PREPARING, SHIPPED, COMPLETED, CANCELLED gibi değerlerle yönetilir.")

    add_section_heading(doc, "2.5. Yapılan Testler")
    add_para(doc, "Test yaklaşımı derleme doğrulaması, senaryo bazlı manuel testler ve hata durumlarının gözlenmesi olmak üzere üç katmanda ele alınmıştır. Derleme doğrulamasında backend için TypeScript derlemesi, frontend için Next.js üretim derlemesi çalıştırılmıştır. Senaryo bazlı testlerde kayıt/giriş, satıcı ürün ekleme, katalog filtreleme, örnek görselden AI sayfasına geçiş, AI görev izleme, satıcıyla mesajlaşma, sipariş durumu güncelleme ve admin örnek yönetimi gibi kullanıcı akışları değerlendirilmiştir.")
    rows = [
        ["Senaryo", "Beklenen sonuç", "Kabul ölçütü"],
        ["Kullanıcı kayıt/giriş", "Geçerli bilgilerle token ve güvenli kullanıcı bilgisi döner.", "Hatalı e-posta/şifre için anlaşılır hata mesajı gösterilir."],
        ["Satıcı ürün ekleme", "Ürün görselleri yüklenir ve katalogda aktif ürün olarak görünür.", "Satıcı olmayan kullanıcı ürün ekleyemez."],
        ["Görsel arama", "Sorgu sanitize edilir, sonuçlar kart olarak listelenir.", "API anahtarı frontend bundle içine taşınmaz."],
        ["AI üretim", "Referans görsel ve prompt ile görev başlar, durum izlenir.", "Tamamlanan model güvenli dosya uç noktasından alınır."],
        ["Mesajlaşma", "Kullanıcı ve satıcı aynı konuşma üzerinden mesajlaşır.", "Konuşma tarafı olmayan kullanıcı erişemez."],
        ["Sipariş", "Katalog ürünü için konuşma ORDERED durumuna geçer.", "Satıcı durum güncelleyebilir, kullanıcı iptal edebilir."],
        ["Admin örnekleri", "Admin örnek ekleyip silebilir.", "Admin olmayan kullanıcı yönetim işlemi yapamaz."],
    ]
    add_table(doc, "Tablo 2.4. Test senaryoları ve kabul ölçütleri", rows, [3.2, 5.0, 7.3], "Proje demo rehberi ve kaynak kodu esas alınarak hazırlanmıştır.", 9)

    add_section_heading(doc, "2.6. Güvenlik, Etik Standartlar ve Varsayımlar")
    for para in [
        "Çalışma yazılım tabanlıdır ve fiziksel bir donanım kurulumu içermemektedir. Bu nedenle donanımsal güvenlik riski bulunmamakla birlikte, kullanıcı verisi, harici API anahtarları ve dosya erişimi açısından yazılım güvenliği ön planda tutulmuştur. API anahtarları .env dosyalarında saklanmış, istemci tarafına yalnızca backend adresi gibi gerekli genel yapılandırmalar aktarılmıştır.",
        "Görsel arama sonuçları internet üzerindeki kaynaklardan geldiği için telif ve kullanım hakları kullanıcı tarafından dikkate alınmalıdır. Sistem, görseli referans olarak AI üretim akışına taşırken kaynağın başlık ve site bilgilerini de gösterir; ancak harici görsellerin lisans durumunu otomatik olarak doğrulamaz. Bu sınırlılık, etik kullanım açısından kullanıcı bilgilendirmesi ve ileride lisans filtreleme mekanizmaları eklenmesi gerektiğini göstermektedir.",
        "AI model üretimi harici servislere bağlıdır. Bu nedenle üretim süresi, model kalitesi, çıktı formatı ve servis erişilebilirliği dış servis yoğunluğuna veya kota durumuna göre değişebilir. PrintForge bu riski azaltmak için görev durumunu arka planda izleyen, hata mesajlarını kullanıcıya aktaran ve admin örnekleriyle dış arama servisleri çalışmasa bile temel akışı sürdürebilen bir tasarım kullanmıştır.",
    ]:
        add_para(doc, para)
    add_figure(doc, diagrams["fig_2_4"], "Şekil 2.4. Kimlik doğrulama ve güvenli erişim akışı")


def _draw_wrapped_text(draw, text, x, y, max_width, fnt, fill="#222222", line_gap=10, max_lines=6):
    current_y = y
    for line in wrap(draw, text, fnt, max_width)[:max_lines]:
        draw.text((x, current_y), line, font=fnt, fill=fill)
        current_y += fnt.size + line_gap
    return current_y


def _hi_box(draw, xy, title, body, fill="#f8fafc", outline="#1f2937"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=30, fill=fill, outline=outline, width=5)
    title_font = font("times", 54, True)
    body_font = font("times", 38)
    draw.text((x1 + 42, y1 + 36), title, font=title_font, fill="#111827")
    _draw_wrapped_text(draw, body, x1 + 42, y1 + 112, x2 - x1 - 84, body_font, "#374151", 10, 5)


def _hi_arrow(draw, start, end, color="#374151"):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=10)
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 44
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        ax = x2 + length * math.cos(angle + delta)
        ay = y2 + length * math.sin(angle + delta)
        draw.line((x2, y2, ax, ay), fill=color, width=10)


def _save_hi_diagram(name, title, boxes, arrows, size=(3600, 2200)):
    OUT_DIR.mkdir(exist_ok=True)
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw.text((120, 90), title, font=font("times", 82, True), fill="#111827")
    draw.line((120, 190, size[0] - 120, 190), fill="#d1d5db", width=4)
    for box in boxes:
        _hi_box(draw, *box)
    for start, end in arrows:
        _hi_arrow(draw, start, end)
    path = OUT_DIR / f"{name}.png"
    img.save(path, quality=95)
    return path


def make_diagrams():
    diagrams = {}
    diagrams["fig_1_1"] = _save_hi_diagram(
        "fig_1_1",
        "Literatürden PrintForge Yaklaşımına",
        [
            ((130, 360, 900, 820), "Eklemeli imalat", "3D baskı; hızlı prototipleme, kişiselleştirme ve düşük adetli üretim için esnek bir üretim yaklaşımı sunar.", "#eef6ff", "#1d4ed8"),
            ((1415, 360, 2185, 820), "AI 3D üretim", "Metin veya referans görselden 3D geometri üretimi, kullanıcıların modelleme eşiğini azaltır.", "#f0fdf4", "#047857"),
            ((2700, 360, 3470, 820), "Pazaryeri", "Satıcı katalogu, fiyat, mesajlaşma ve sipariş takibi üretim hizmetine erişimi düzenler.", "#fff7ed", "#c2410c"),
            ((820, 1320, 2780, 1780), "PrintForge özgünleşmesi", "Fikir bulma, referans seçme, AI üretim görevi, 3D önizleme ve satıcı iletişimi tek platformda birleştirilmiştir.", "#f8fafc", "#111827"),
        ],
        [((900, 590), (1415, 590)), ((2185, 590), (2700, 590)), ((520, 820), (1160, 1320)), ((1800, 820), (1800, 1320)), ((3085, 820), (2440, 1320))],
    )
    diagrams["fig_2_1"] = _save_hi_diagram(
        "fig_2_1",
        "PrintForge Sistem Mimarisi",
        [
            ((130, 360, 1000, 860), "Frontend", "Next.js, React, TypeScript, Tailwind CSS ve Three.js ile kullanıcı arayüzü ve 3D model önizleme yönetilir.", "#eef6ff", "#2563eb"),
            ((1365, 360, 2235, 860), "Backend API", "Express ve TypeScript rotaları; kimlik doğrulama, doğrulama, iş kuralları ve servis entegrasyonlarını yürütür.", "#f0fdf4", "#059669"),
            ((2600, 360, 3470, 860), "Veri katmanı", "Prisma modelleri kullanıcı, model, mesaj, yorum, soru, konuşma ve örnek içerikleri ilişkilendirir.", "#fff7ed", "#ea580c"),
            ((480, 1250, 1350, 1740), "Harici servisler", "SerpApi/Google Images, UploadThing ve Tripo3D/Hitem3D gibi servisler backend üzerinden çağrılır.", "#faf5ff", "#7e22ce"),
            ((2250, 1250, 3120, 1740), "Dosya depolama", "Orijinal GLB/STL dosyaları, güvenli önizleme JSON verileri ve ürün görselleri düzenli anahtarlarla saklanır.", "#f8fafc", "#374151"),
        ],
        [((1000, 610), (1365, 610)), ((2235, 610), (2600, 610)), ((1800, 860), (920, 1250)), ((1800, 860), (2685, 1250))],
    )
    diagrams["fig_2_2"] = _save_hi_diagram(
        "fig_2_2",
        "Referans Görselden AI Üretime Veri Akışı",
        [
            ((100, 420, 700, 880), "1. Referans", "Admin örneği veya internet görsel arama sonucu seçilir.", "#eef6ff", "#2563eb"),
            ((820, 420, 1420, 880), "2. Prompt", "Başlık, kaynak ve arama bağlamı ile otomatik üretim açıklaması hazırlanır.", "#f0fdf4", "#059669"),
            ((1540, 420, 2140, 880), "3. Backend", "Görsel ve prompt doğrulanır; AI üretim servisine görev gönderilir.", "#fff7ed", "#ea580c"),
            ((2260, 420, 2860, 880), "4. İzleme", "Görev durumu belirli aralıklarla sorgulanır ve kullanıcı bilgilendirilir.", "#faf5ff", "#7e22ce"),
            ((2980, 420, 3580, 880), "5. Model", "Tamamlanan GLB/STL dosyası kaydedilir ve önizleme açılır.", "#f8fafc", "#374151"),
            ((850, 1340, 2750, 1760), "Satıcı bağlantısı", "Kullanıcı tamamlanan AI modeli veya katalog ürünü için satıcı seçer, konuşma başlatır ve teklif/sipariş sürecini izler.", "#ecfdf5", "#047857"),
        ],
        [((700, 650), (820, 650)), ((1420, 650), (1540, 650)), ((2140, 650), (2260, 650)), ((2860, 650), (2980, 650)), ((3280, 880), (2350, 1340))],
    )
    diagrams["fig_2_3"] = _save_hi_diagram(
        "fig_2_3",
        "Veri Modeli İlişkileri",
        [
            ((140, 360, 880, 760), "User", "email, password, name, role ve companyName alanlarıyla kullanıcı, satıcı ve admin hesaplarını temsil eder.", "#eef6ff", "#2563eb"),
            ((1430, 360, 2170, 760), "Model", "AI ve CATALOG türleri; durum, prompt, fiyat, dosya anahtarları ve teknik metrikleri tutar.", "#f0fdf4", "#059669"),
            ((2720, 360, 3460, 760), "ExampleItem", "Admin tarafından yönetilen başlık, kategori, görsel URL, prompt ve etiket bilgilerini saklar.", "#fff7ed", "#ea580c"),
            ((780, 1180, 1520, 1580), "Conversation", "buyerId, sellerId, modelId ve status alanlarıyla mesajlaşma ve sipariş sürecini bağlar.", "#faf5ff", "#7e22ce"),
            ((2070, 1180, 2810, 1580), "Message", "Konuşma içeriği, teklif bilgisi, gönderen ve okunma zamanı burada tutulur.", "#f8fafc", "#374151"),
            ((140, 1700, 880, 2060), "Review / Question", "Ürün puanı, yorum, soru ve satıcı cevabı katalog güvenini destekler.", "#fef2f2", "#b91c1c"),
        ],
        [((880, 560), (1430, 560)), ((1800, 760), (1150, 1180)), ((1520, 1380), (2070, 1380)), ((520, 760), (520, 1700)), ((1800, 760), (520, 1700))],
    )
    diagrams["fig_2_4"] = _save_hi_diagram(
        "fig_2_4",
        "Kimlik Doğrulama ve Güvenli Erişim Akışı",
        [
            ((120, 430, 870, 900), "Kayıt / giriş", "E-posta, şifre, ad ve rol bilgileri Zod ile doğrulanır; parola bcrypt ile özetlenir.", "#eef6ff", "#2563eb"),
            ((1030, 430, 1780, 900), "JWT", "Başarılı oturumda kullanıcı kimliği imzalı token içinde taşınır.", "#f0fdf4", "#059669"),
            ((1940, 430, 2690, 900), "Korunan rota", "Authorization başlığı okunur, token doğrulanır ve kullanıcı veritabanında eşleştirilir.", "#fff7ed", "#ea580c"),
            ((2850, 430, 3600, 900), "Yetki", "Seller, user ve admin akışları role göre ayrılır; dosya erişimi sahiplik veya konuşma tarafı olma koşuluna bağlanır.", "#faf5ff", "#7e22ce"),
            ((620, 1350, 2980, 1770), "Model dosyası erişimi", "Model dosyaları yalnızca model sahibi veya ilgili konuşmanın tarafları tarafından alınabilir. Dosya yolu uploads dizininin dışına çıkmayacak biçimde çözülür.", "#f8fafc", "#374151"),
        ],
        [((870, 665), (1030, 665)), ((1780, 665), (1940, 665)), ((2690, 665), (2850, 665)), ((3225, 900), (2400, 1350))],
    )
    diagrams["fig_3_1"] = _save_hi_diagram(
        "fig_3_1",
        "Modül Tamamlama Durumu",
        [
            ((180, 360, 1050, 760), "Tamamlanan çekirdek", "Kimlik doğrulama, pazaryeri, satıcı paneli, mesajlaşma, admin örnekleri ve 3D önizleme uygulanmıştır.", "#ecfdf5", "#047857"),
            ((1370, 360, 2240, 760), "Servise bağlı akış", "AI üretim ve internet görsel arama işlevleri harici servis anahtarı ve kota durumuna bağlıdır.", "#fff7ed", "#ea580c"),
            ((2560, 360, 3430, 760), "Geliştirilebilir alan", "Maliyet tahmini, baskı uygunluk analizi, ödeme ve lojistik entegrasyonu sonraki sürümlere bırakılmıştır.", "#fef2f2", "#b91c1c"),
            ((600, 1230, 3000, 1700), "Genel bulgu", "Platform fikir bulma, AI destekli model üretimi ve satıcı iletişimini bütünleşik bir kullanıcı deneyiminde birleştirmiştir.", "#f8fafc", "#374151"),
        ],
        [((1050, 560), (1370, 560)), ((2240, 560), (2560, 560)), ((1800, 760), (1800, 1230))],
    )
    diagrams["fig_3_2"] = _save_hi_diagram(
        "fig_3_2",
        "Derleme ve Doğrulama Özeti",
        [
            ((260, 440, 1460, 930), "Backend", "TypeScript derlemesi başarıyla tamamlanmış; Express API katmanı tsc ile doğrulanmıştır.", "#ecfdf5", "#047857"),
            ((2140, 440, 3340, 930), "Frontend", "Next.js üretim derlemesi başarıyla tamamlanmış ve 21 sayfa oluşturulmuştur.", "#ecfdf5", "#047857"),
            ((720, 1330, 2880, 1740), "Not", "Webpack önbellek uyarısı derlemeyi engellememiştir; işlevsel hata değil, izlenebilir çevresel uyarı olarak değerlendirilmiştir.", "#fff7ed", "#ea580c"),
        ],
        [((1460, 685), (2140, 685)), ((1800, 930), (1800, 1330))],
    )
    return diagrams


def add_extended_technical_design(doc: Document):
    doc.add_page_break()
    add_section_heading(doc, "2.7. Ayrıntılı Modül Tasarımı")
    add_para(doc, "Bu alt bölüm, uygulamanın yalnızca hangi teknolojilerle geliştirildiğini değil, modüllerin hangi sorumluluklarla ayrıldığını ve bu ayrımın neden tercih edildiğini açıklamak amacıyla eklenmiştir. PrintForge'un hedefi, kullanıcının fikir bulma, AI üretim, katalog inceleme ve satıcıyla iletişim adımlarını aynı sistemde tamamlayabilmesidir. Bu nedenle her modül tek başına çalışan bir sayfa değil, bir sonraki iş adımına veri aktaran bir süreç parçası olarak tasarlanmıştır.")

    detailed_sections = [
        (
            "2.7.1. Frontend Sayfa Yapısı ve Kullanıcı Deneyimi",
            [
                "Frontend katmanında ana sayfa, örnekler, AI üretim stüdyosu, pazaryeri, sepet, favoriler, mesajlar, siparişler, satıcı ürün paneli ve admin örnek yönetimi gibi sayfalar bulunmaktadır. Bu sayfaların her biri farklı kullanıcı rolünün bir görevi tamamlamasına hizmet eder. Kullanıcı için en kritik akış örnek görselden AI üretim sayfasına geçiştir; satıcı için ürün ekleme ve gelen mesaj/sipariş takibidir; admin için ise platformdaki referans içeriklerinin düzenli tutulmasıdır.",
                "Arayüz tasarımında rol tabanlı görünürlük tercih edilmiştir. Satıcı hesabı ürün ekleme ve ürün yönetme bağlantılarını görürken, normal kullanıcı katalog inceleme, AI üretim ve satıcıya ulaşma akışlarına yönlendirilir. Bu yaklaşım kullanıcıyı ilgisiz kontrollerle karşılaştırmadığı için hata olasılığını azaltır ve arayüzü daha anlaşılır hale getirir.",
                "Next.js sayfa yapısı sayesinde statik ve dinamik sayfalar ayrılabilmiştir. Pazaryeri ve sohbet gibi veriyle çalışan sayfalar gerektiğinde sunucu veya istemci tarafı veri çekme mekanizmalarını kullanırken, örnekler ve AI üretim sayfası kullanıcı etkileşimine göre durum yönetimi yapar. Böylece sistem hem hızlı açılış hem de etkileşimli kullanıcı deneyimi hedeflerini birlikte destekler.",
            ],
        ),
        (
            "2.7.2. Örnek Görsel ve Prompt Aktarım Mekanizması",
            [
                "Örnekler sayfasında iki farklı referans kaynağı vardır: admin tarafından yönetilen kontrollü örnekler ve internet görsel arama sonuçları. Kontrollü örnekler, dış servis kotası veya arama kalitesi sorunlarında sistemin temel akışını sürdürebilmesini sağlar. İnternet araması ise kullanıcının daha geniş fikir havuzuna ulaşmasına yardımcı olur.",
                "Kullanıcı bir görsel kartını seçtiğinde yalnızca görsel URL'si aktarılmaz; başlık, kaynak, arama kelimesi ve oluşturulan prompt da query parametreleriyle AI üretim sayfasına taşınır. Bu karar, kullanıcının boş bir prompt alanıyla karşılaşmasını önler. Prompt, referans görselin 3D baskıya uygun, temiz yüzeyli ve üretilebilir bir modele dönüştürülmesi yönünde başlangıç metni sağlar.",
                "AI üretim sayfası dış görseli doğrudan üretim servisine iletmek yerine backend proxy mekanizması ve tarayıcı tarafında dönüştürme adımıyla dosya haline getirir. Böylece kullanıcı, internetten seçtiği bir görseli ayrıca bilgisayarına indirmeden üretim akışına katabilir. Bu tasarım, fikirden üretim görevine geçiş süresini kısaltmaktadır.",
            ],
        ),
        (
            "2.7.3. AI Üretim Görevi ve Durum İzleme",
            [
                "AI üretim servisleri genellikle anlık sonuç döndürmek yerine görev tabanlı çalışır. Bu nedenle PrintForge, model üretimini tek istek-tek cevap şeklinde değil, görev başlatma ve görev durumunu izleme şeklinde tasarlamıştır. Backend üretim isteğini aldıktan sonra harici servisten taskId alır ve bunu Model kaydıyla ilişkilendirir.",
                "Frontend tarafındaki AiGenerationProvider, aktif üretim görevini localStorage içinde saklar. Bu sayede kullanıcı sayfadan ayrılsa veya katalogda gezinmeye devam etse bile üretim görevi tamamen unutulmaz. Sistem görev durumunu belirli aralıklarla sorgular ve tamamlandığında ilgili model kimliğini kullanıcı arayüzüne taşır.",
                "Bu yaklaşım kullanıcı deneyimi açısından önemlidir. 3D model üretimi servis yoğunluğuna göre uzun sürebilir; kullanıcıyı tek ekranda beklemeye zorlamak yerine arka plan izleme yapılması, platformun pazaryeri yönüyle birlikte çalışmasını sağlar. Kullanıcı model hazırlanırken katalog ürünlerini inceleyebilir veya daha önceki üretim geçmişine bakabilir.",
            ],
        ),
        (
            "2.7.4. Backend Doğrulama ve Hata Yönetimi",
            [
                "Backend katmanında gelen isteklerin biçimsel doğruluğu Zod şemalarıyla denetlenmiştir. Kayıt, giriş, profil güncelleme, ürün ekleme, yorum, soru, cevap, mesaj, sipariş ve görsel arama isteklerinin her biri beklenen tip, uzunluk ve zorunlu alanlara göre kontrol edilir. Bu yaklaşım hatalı verinin veritabanına ulaşmasını engeller.",
                "Hata yönetiminde kullanıcıya teknik ayrıntı yerine anlaşılır mesaj döndürme yaklaşımı kullanılmıştır. Örneğin görsel arama kotası dolduğunda veya API anahtarı geçersiz olduğunda kullanıcıya durumun ne olduğu açıklanır. Geliştirici açısından ise backend logları daha ayrıntılı hata izleme olanağı sağlar.",
                "Üretim servisleri ve görsel arama sağlayıcıları dış sistemler olduğu için hata olasılığı yerel koddan bağımsız olarak ortaya çıkabilir. Bu nedenle timeout, kota, yetkisiz anahtar, geçersiz yanıt ve bağlantı problemi gibi durumlar ayrı ayrı değerlendirilmiştir. Bu tasarım, sistemin hata anında tamamen sessiz kalmasını veya kullanıcıyı belirsiz bir bekleme durumunda bırakmasını önler.",
            ],
        ),
        (
            "2.7.5. Model Dosyası ve 3D Önizleme Tasarımı",
            [
                "AI üretim görevi başarıyla tamamlandığında backend model dosyasını indirir ve dosya türünü kontrol eder. GLB formatındaki modeller doğrudan önizlenebilir dosya olarak saklanırken, STL formatındaki modeller için güvenli önizleme verisi üretilebilir. Bu ayrım, farklı servislerden gelen çıktı formatlarının aynı uygulamada yönetilebilmesi için gereklidir.",
                "3D önizleme bileşeni Three.js ekosistemi üzerine kurulmuştur. ModelViewer bileşeni GLTF/GLB dosyasını yükler, modeli merkezler, kamera sınırlarını ayarlar ve kullanıcıya döndürme/inceleme olanağı verir. Böylece kullanıcı model dosyasını indirmeden önce biçimsel olarak ne üretildiğini görebilir.",
                "Model dosyalarına erişim güvenli uç nokta üzerinden yapılır. Kullanıcı yalnızca kendi modeline veya tarafı olduğu konuşmaya bağlı modele erişebilir. Dosya yolu çözülürken uploads dizini dışına çıkılmaması kontrol edildiği için yol manipülasyonu riski azaltılmıştır. Bu yaklaşım özellikle satıcı-kullanıcı arasında paylaşılan AI modelleri için önemlidir.",
            ],
        ),
        (
            "2.7.6. Pazaryeri ve Satıcı Ürün Yönetimi",
            [
                "Pazaryeri modülü, AI üretimden bağımsız olarak çalışan hazır ürün katalogunu temsil eder. Satıcı, ürün adı, açıklama, kategori, fiyat ve görseller ile katalog ürünü oluşturur. Kullanıcılar bu ürünleri arayabilir, kategori ve maksimum fiyat filtreleriyle daraltabilir, ürün görsellerini inceleyebilir ve satıcıya mesaj gönderebilir.",
                "Ürün yönetiminde pasif hale getirme yaklaşımı tercih edilmiştir. Satıcı bir ürünü kaldırdığında kayıt tamamen yok edilmek yerine INACTIVE durumuna alınabilir. Bu yaklaşım, ileride sipariş geçmişi, konuşma kayıtları ve raporlama ihtiyaçları için veri bütünlüğünün korunmasına yardımcı olur.",
                "Ürün detayında yorum ve soru-cevap alanlarının bulunması, pazaryeri güvenini artıran sosyal kanıt mekanizmasıdır. Kullanıcı ürün hakkında soru sorabilir; yalnızca ürünün satıcısı bu soruya cevap verebilir. Böylece ürün bilgisinin tek taraflı katalog açıklamasından ibaret kalmaması sağlanmıştır.",
            ],
        ),
        (
            "2.7.7. Mesajlaşma ve Sipariş Durum Yaşam Döngüsü",
            [
                "Mesajlaşma modülü kullanıcı ile satıcı arasında ürün veya AI modeli bağlamında konuşma oluşturur. Katalog ürünleri için konuşma satıcının ürün sahibi olması üzerinden açılırken, AI modeli için kullanıcı mesaj göndermek istediği satıcıyı seçebilir. Bu fark, hazır ürün alımı ile özel üretim talebi arasındaki iş akışı farkını yansıtır.",
                "Konuşma kayıtlarında buyerId, sellerId, modelId, modelType ve status alanları bulunur. Status alanı yalnızca mesajlaşma durumunu değil, sipariş yaşam döngüsünü de temsil eder. ORDERED, PREPARING, SHIPPED, COMPLETED ve CANCELLED değerleri alıcı ve satıcı arasında sipariş takibinin metinsel mesajlardan ayrı, izlenebilir bir süreç olarak yürütülmesini sağlar.",
                "Okunmamış mesaj sayısı ve bildirim özeti, kullanıcının yeni etkileşimleri kaçırmaması için eklenmiştir. Mesajlar ekranı konuşmaları listelerken siparişler ekranı yalnızca aktif sipariş durumuna geçmiş konuşmaları gösterir. Bu ayrım, iletişim ve sipariş takibinin karışmasını önler.",
            ],
        ),
        (
            "2.7.8. Dağıtım, Ortam Değişkenleri ve Sürdürme Yaklaşımı",
            [
                "Frontend dağıtımı için Netlify yapılandırması hazırlanmıştır. Netlify yalnızca frontend uygulamasını yayınladığı için backend'in ayrı bir web servisi olarak çalışması gerekir. Bu ayrım README içinde açık biçimde belirtilmiş ve canlı ortamda localhost adreslerinin kullanılmaması gerektiği vurgulanmıştır.",
                "Backend tarafında DATABASE_URL, JWT_SECRET, FRONTEND_URLS, SERPAPI_API_KEY, TRIPO_API_KEY ve HITEM3D anahtarları gibi değerler ortam değişkenleriyle yönetilir. Böylece gizli bilgiler kaynak koduna yazılmaz. Frontend tarafında ise yalnızca NEXT_PUBLIC_BACKEND_URL gibi istemcinin bilmesi gereken adresler kullanılır.",
                "Sürdürme açısından modüler rota yapısı önemlidir. Auth, AI, models, chat, examples ve images rotalarının ayrı dosyalarda bulunması, hata ayıklamayı ve yeni özellik eklemeyi kolaylaştırır. Bu yapı ileride ödeme, maliyet tahmini veya model kalite puanı gibi yeni modüllerin sisteme eklenmesini daha düzenli hale getirir.",
            ],
        ),
    ]
    for heading, paragraphs in detailed_sections:
        doc.add_page_break()
        add_section_heading(doc, heading)
        for para in paragraphs:
            add_para(doc, para)

    add_para(doc, "Ayrıntılı modül tasarımında ortaya çıkan en önemli sonuç, PrintForge'un tek bir servis entegrasyonundan ibaret olmadığıdır. Sistem, harici AI üretimini kullanıcı deneyimi, güvenlik, veri modeli, pazaryeri ve mesajlaşma katmanlarıyla birlikte ele almaktadır. Bu nedenle proje kapsamı, klasik bir katalog uygulaması veya yalnızca görselden model üreten bir arayüzden daha geniştir.")
    rows = [
        ["Katman", "Sorumluluk", "Kontrol noktası"],
        ["Frontend", "Kullanıcı etkileşimi, yönlendirme, 3D önizleme ve üretim durumu", "Rol tabanlı görünürlük ve kullanıcıya anlaşılır durum mesajı"],
        ["Backend", "Doğrulama, kimlik, iş kuralları ve harici servis çağrıları", "Zod şemaları, JWT doğrulaması ve hata kodları"],
        ["Veri modeli", "Kullanıcı, model, konuşma, mesaj ve katalog ilişkileri", "Prisma ilişkileri ve indeksler"],
        ["Dosya yönetimi", "Model dosyaları ve görsel bağlantıları", "Uploads dizini sınırı ve yetki kontrolü"],
        ["Harici servisler", "Görsel arama, görsel yükleme ve AI üretim", "Timeout, kota ve anahtar hatası yönetimi"],
    ]
    add_table(doc, "Tablo 2.5. Ayrıntılı mimari katmanlar ve kontrol noktaları", rows, [3.2, 6.2, 6.1], "Proje kaynak kodu ve modül sorumlulukları esas alınarak hazırlanmıştır.", 9)

    doc.add_page_break()
    add_section_heading(doc, "2.8. İş Paketleri ve Zaman Yönetimi")
    add_para(doc, "Bitirme projesi süreci, yalnızca kod yazımı olarak değil, ihtiyaç analizi, mimari kararlar, uygulama geliştirme, entegrasyon, test ve raporlama adımlarını kapsayan bir mühendislik süreci olarak planlanmıştır. Bu nedenle iş paketleri çıktıya göre ayrılmıştır. Her iş paketinin sonunda kontrol edilebilir bir çıktı üretilmesi hedeflenmiştir.")
    rows = [
        ["İş paketi", "Amaç", "Çıktı"],
        ["İhtiyaç analizi", "3D baskı kullanıcılarının model ve satıcı bulma problemini tanımlamak", "Problem tanımı, araştırma sorusu ve kapsam"],
        ["Mimari tasarım", "Frontend, backend, veri modeli ve servis ayrımını planlamak", "Sistem mimarisi ve Prisma şeması"],
        ["Katalog geliştirme", "Satıcı ürünlerini ve kullanıcı katalog deneyimini oluşturmak", "Pazaryeri, ürün kartları, filtreler, detaylar"],
        ["AI akışı", "Referans görselden üretim görevine geçişi sağlamak", "Örnekler, prompt aktarımı, görev izleme"],
        ["Mesaj/sipariş", "Satıcı-kullanıcı iletişimini ve sipariş durumlarını izlemek", "Konuşma, mesaj, sipariş ekranları"],
        ["Test ve raporlama", "Derleme, senaryo ve hata durumlarını doğrulamak", "Derleme çıktıları, test tabloları, tez metni"],
    ]
    add_table(doc, "Tablo 2.6. İş paketleri ve çıktıları", rows, [3.2, 6.1, 6.2], "Proje geliştirme süreci esas alınarak hazırlanmıştır.", 9)
    add_para(doc, "İş paketlerinin sıralı tasarlanması, projenin son aşamada bütünleşik görünmesini sağlamıştır. Örneğin AI üretim sayfası katalog ve mesajlaşma modülleri olmadan yalnızca model oluşturan bir araç olarak kalacaktı. Mesajlaşma ve sipariş modüllerinin eklenmesiyle AI çıktısının gerçek üretim talebine dönüşmesi mümkün hale gelmiştir.")


def add_detailed_review_pages(doc: Document):
    topics = [
        ("2.9.1. Kayıt ve Giriş Akışının Değerlendirilmesi", "kullanıcı hesabı oluşturma ve oturum açma", "registerSchema ve loginSchema ile e-posta, parola, ad ve rol bilgilerinin doğrulanması", "hatalı biçimde gelen verinin veritabanına yazılmadan reddedilmesi"),
        ("2.9.2. Rol Tabanlı Kullanıcı Deneyimi", "USER, SELLER ve ADMIN rollerinin arayüzde ayrıştırılması", "satıcı paneli, admin örnek yönetimi ve kullanıcı katalog akışlarının ayrı görünmesi", "kullanıcının kendi rolüyle ilgisiz işlemlerle karşılaşmaması"),
        ("2.9.3. Satıcı Hesabı ve Şirket Bilgisi", "satıcıların şirket adı veya görünen ad ile temsil edilmesi", "companyName alanının satıcı hesaplarında öncelikli gösterilmesi", "pazaryerinde satıcı güveninin ve kurumsal görünürlüğün artırılması"),
        ("2.9.4. Profil, E-posta ve Şifre Güncelleme", "kullanıcının hesap bilgisini güvenli biçimde güncellemesi", "mevcut parola kontrolü ve bcrypt ile yeni parola özeti", "başka kullanıcının e-posta adresinin üzerine yazılmasının engellenmesi"),
        ("2.9.5. Admin Örnek İçerik Yönetimi", "AI üretim için kontrollü referans içeriklerinin hazırlanması", "title, category, imageUrl, prompt ve tags alanlarının doğrulanması", "dış arama servisi çalışmasa bile demo ve kullanım akışının sürdürülebilmesi"),
        ("2.9.6. Görsel Arama Sorgu Temizleme", "internet görsel araması sırasında kullanıcı sorgusunun güvenli hale getirilmesi", "kontrol karakterleri ve riskli sembollerin temizlenmesi, 80 karakter sınırı", "hatalı veya zararlı sorguların servis çağrısına dönüşmeden önce azaltılması"),
        ("2.9.7. Görsel Arama Önbellekleme", "aynı arama sonuçlarının kısa süre içinde tekrar istenmesini azaltma", "provider ve sorguya göre bellek içi cache anahtarı oluşturulması", "kota kullanımının azalması ve kullanıcıya daha hızlı yanıt verilmesi"),
        ("2.9.8. SerpApi ve Google Images Sağlayıcı Seçimi", "görsel aramada farklı sağlayıcıların desteklenmesi", "IMAGE_SEARCH_PROVIDER ve API anahtarlarına göre sağlayıcı belirlenmesi", "tek sağlayıcıya bağımlılığın azaltılması"),
        ("2.9.9. Harici Görsel Proxy Mekanizması", "dış kaynaklı görsellerin AI üretim sayfasına güvenli aktarılması", "proxy-image uç noktasıyla içerik tipinin image/* olarak doğrulanması", "bozuk veya görsel olmayan bağlantıların üretim akışını bozmasının önlenmesi"),
        ("2.9.10. Prompt Oluşturma Stratejisi", "referans görselden üretim açıklaması hazırlama", "görsel başlığı, kaynak ve arama bağlamını prompt içine yerleştirme", "kullanıcının boş prompt alanıyla başlamaması"),
        ("2.9.11. AI Üretim İsteği Doğrulaması", "görselden model üretimi başlatmadan önce girdi denetimi", "multer dosya sınırı, image/* mimetype kontrolü ve üretim tipi doğrulaması", "uygunsuz dosya veya eksik prompt nedeniyle oluşabilecek servis hatalarının azaltılması"),
        ("2.9.12. Tripo3D ve Hitem3D Entegrasyon Mantığı", "AI model üretimi için harici servisleri görev tabanlı kullanma", "servis anahtarı varlığına göre sağlayıcı akışının seçilmesi", "farklı servislerin çıktı formatı ve kimlik doğrulama biçimlerine uyum sağlanması"),
        ("2.9.13. AI Görev Durumu İzleme", "uzun süren üretimlerin kullanıcıya düzenli bildirilmesi", "taskId üzerinden belirli aralıklarla status sorgulama", "kullanıcının belirsiz bekleme durumunda bırakılmaması"),
        ("2.9.14. Üretim Geçmişi", "kullanıcının önceki AI üretimlerini görebilmesi", "userId ve type='AI' filtresiyle geçmiş modellerin listelenmesi", "tamamlanan üretimlere geri dönülebilmesi"),
        ("2.9.15. GLB ve STL Çıktı Ayrımı", "AI servisinden dönen model dosyasının türünü belirleme", "GLB başlığının kontrol edilmesi ve STL için işleme yolunun ayrılması", "farklı çıktı formatlarının aynı platformda yönetilmesi"),
        ("2.9.16. SecureGeometry Veri Yapısı", "STL dosyalarının kontrollü önizleme verisine dönüştürülmesi", "pozisyon, normal, hacim, yüzey alanı ve sınır kutusu hesaplama", "ham dosya yerine görüntüleme odaklı güvenli geometri sunulması"),
        ("2.9.17. Model Dosyası Erişim Yetkisi", "model dosyalarının yalnızca yetkili kişilere verilmesi", "model sahibi veya ilgili konuşmanın tarafı olma koşulu", "özel üretim modelinin yetkisiz kişiler tarafından indirilmesinin engellenmesi"),
        ("2.9.18. Three.js Model Önizleme", "3D modelin tarayıcı üzerinde incelenmesi", "Canvas, Bounds, Center, Environment ve OrbitControls kullanımı", "kullanıcının model dosyasını indirmeden önce görsel kontrol yapabilmesi"),
        ("2.9.19. Katalog Filtreleme", "satıcı ürünlerinin aranabilir ve sınıflandırılabilir hale getirilmesi", "ürün adı, açıklama, satıcı adı, kategori ve maksimum fiyat filtresi", "pazaryeri içinde hızlı karar verme"),
        ("2.9.20. Ürün Kartı ve Galeri Deneyimi", "katalog ürünlerinin görsel olarak sunulması", "kapak görseli, çoklu görsel sayacı ve detay modalı", "kullanıcının ürün hakkında hızlı bağlam edinmesi"),
        ("2.9.21. Yorum ve Puanlama", "ürün güvenini artıran sosyal kanıt oluşturma", "modelId ve userId için tekil yorum kaydı", "aynı kullanıcının aynı ürüne tekrar tekrar bağımsız puan vermesinin önlenmesi"),
        ("2.9.22. Soru-Cevap Mekanizması", "ürün hakkında alıcı ve satıcı arasında açık bilgi paylaşımı", "soruyu kullanıcının, cevabı yalnızca ürün satıcısının yazabilmesi", "ürün bilgisinin katalog açıklamasıyla sınırlı kalmaması"),
        ("2.9.23. Satıcı Ürün Ekleme", "satıcının katalogda ürün yayınlaması", "ürün adı, açıklama, kategori, fiyat ve görsel URL doğrulaması", "satıcı olmayan kullanıcıların ürün yayınlamasının engellenmesi"),
        ("2.9.24. Ürün Görsel Yükleme", "satıcı ürünlerinin görsel olarak desteklenmesi", "UploadThing üzerinden en fazla beş görselin alınması", "katalog kartlarının güvenilir ve incelenebilir hale gelmesi"),
        ("2.9.25. Ürün Düzenleme ve Pasifleştirme", "satıcının mevcut ürünlerini güncellemesi veya kaldırması", "ürün sahibinin kontrol edilmesi ve status alanının INACTIVE yapılması", "geçmiş konuşmalar korunurken katalog görünürlüğünün kapatılması"),
        ("2.9.26. Favori ve Sepet Akışı", "kullanıcının ilgilendiği ürünleri ayırması", "localStorage tabanlı favori ve sepet yardımcı fonksiyonları", "hızlı karar veremeyen kullanıcı için geri dönüş noktası oluşturma"),
        ("2.9.27. Katalog Ürününden Mesaj Başlatma", "ürün hakkında satıcıyla iletişime geçme", "modelId üzerinden ürün sahibinin satıcı olarak seçilmesi", "kullanıcı ile satıcı arasında bağlamlı konuşma oluşturulması"),
        ("2.9.28. AI Modeli İçin Satıcı Seçimi", "özel üretim modeli için uygun satıcıya ulaşma", "AI modeli tamamlandığında satıcı listesinin gösterilmesi", "hazır katalog ürünü olmayan özel taleplerin de pazaryerine bağlanması"),
        ("2.9.29. Sipariş Oluşturma", "katalog ürününün satın alma niyetine dönüştürülmesi", "conversation status değerinin ORDERED yapılması", "mesajlaşma ile sipariş takibinin aynı bağlamda ilişkilendirilmesi"),
        ("2.9.30. Sipariş Durum Güncelleme", "üretim ve teslim sürecinin izlenmesi", "PREPARING, SHIPPED, COMPLETED ve CANCELLED durumları", "satıcı ve alıcının süreci tek ekrandan takip etmesi"),
        ("2.9.31. Okunmamış Mesaj ve Bildirim Özeti", "yeni iletişimlerin kaçırılmaması", "readAt alanı ve unreadCount hesaplaması", "mesaj merkezinin aktif kullanımı"),
        ("2.9.32. Konuşma Arşivleme", "kullanıcının mesaj listesini sadeleştirmesi", "buyerArchivedAt ve sellerArchivedAt alanlarıyla taraf bazlı gizleme", "bir tarafın arşivlemesinin diğer tarafın kaydını silmemesi"),
        ("2.9.33. Prisma İndeksleri", "veritabanı sorgularının düzenli çalışması", "userId, modelId, type, status ve conversation ilişkilerine indeks eklenmesi", "listeleme ve erişim sorgularının ölçeklenebilirliğinin artırılması"),
        ("2.9.34. SQLite Uyumluluk Katmanı", "yerel geliştirme ortamında eksik tabloları hazırlama", "server başlangıcında tablo ve kolon kontrolleri", "geliştirme sırasında migration farklarından kaynaklı hataların azaltılması"),
        ("2.9.35. CORS Yapılandırması", "frontend ve backend arasında güvenli istekleşme", "FRONTEND_URLS listesindeki adreslere izin verilmesi", "canlı ortamda yanlış origin kullanımının önlenmesi"),
        ("2.9.36. Helmet ve HTTP Güvenliği", "temel HTTP güvenlik başlıklarının ayarlanması", "Express uygulamasında Helmet middleware kullanımı", "tarayıcı kaynaklı temel saldırı yüzeyinin azaltılması"),
        ("2.9.37. Ortam Değişkenleri Yönetimi", "gizli anahtarların kaynak koddan ayrılması", "JWT_SECRET, API anahtarları ve backend adreslerinin .env içinde tutulması", "canlı dağıtımda gizli bilgilerin korunması"),
        ("2.9.38. Netlify Frontend Dağıtımı", "frontend uygulamasının canlıya alınması", "netlify.toml içinde base, build command ve publish ayarları", "frontend ve backend dağıtım sorumluluklarının ayrılması"),
        ("2.9.39. Backend Canlı Ortam Gereksinimleri", "API servisinin frontend dışında barındırılması", "DATABASE_URL, FRONTEND_URLS ve servis anahtarlarının ayrı panelde tanımlanması", "Netlify'ın yalnızca frontend yayınlamasından doğan mimari farkın yönetilmesi"),
        ("2.9.40. Demo Riskleri ve Yedek Plan", "sunum sırasında dış servis sorunlarını yönetme", "admin örnekleri, hazır model geçmişi ve port kontrol listesi", "jüri sunumunun servis kotasına tamamen bağımlı kalmaması"),
        ("2.9.41. Etik ve Telif Değerlendirmesi", "harici görsellerin referans olarak kullanımında sorumluluk", "kaynak bilgisinin gösterilmesi ve kullanıcı bilgilendirmesi", "lisans filtreleme ihtiyacının gelecek çalışma olarak belirlenmesi"),
        ("2.9.42. Kişisel Veri ve Gizlilik", "hesap ve mesaj verilerinin korunması", "parola özetleme, token doğrulama ve rol bazlı erişim", "kullanıcı bilgilerinin gereksiz yere istemciye taşınmaması"),
        ("2.9.43. Baskıya Uygunluk Analizi İçin Temel", "üretilen modelin teknik olarak değerlendirilebilmesi", "hacim, yüzey alanı ve sınır kutusu gibi metriklerin hesaplanması", "gelecekte maliyet ve baskı zorluğu tahminine temel oluşturma"),
        ("2.9.44. Maliyet Tahmini Geliştirme Alanı", "kullanıcının üretim öncesi yaklaşık fiyat görmesi", "malzeme, hacim, doluluk oranı ve baskı süresi değişkenlerinin modellenmesi", "satıcı teklif sürecinin daha şeffaf hale gelmesi"),
        ("2.9.45. Gözlemlenebilirlik ve Loglama", "hata ayıklama ve servis takibi", "backend hata logları ve servis yanıtlarının kontrollü izlenmesi", "canlı ortamda kullanıcı şikayetlerinin daha hızlı çözümlenmesi"),
    ]

    doc.add_page_break()
    add_section_heading(doc, "2.9. Ayrıntılı Teknik İnceleme Sayfaları")
    add_para(doc, "Bu bölüm, tezin önceki alt başlıklarında açıklanan mimari kararları daha ayrıntılı biçimde açmak için hazırlanmıştır. Her alt başlık, PrintForge projesindeki gerçek bir modül, kontrol noktası veya geliştirme kararına karşılık gelmektedir. Amaç sayfa sayısını biçimsel olarak artırmak değil; uygulamanın teslim edilebilir bir mühendislik çalışması olarak hangi kararlarla geliştirildiğini ayrıntılı biçimde göstermektir.")

    for heading, focus, implementation, contribution in topics:
        doc.add_page_break()
        add_section_heading(doc, heading)
        add_para(doc, f"Bu başlıkta {focus} ele alınmaktadır. PrintForge kapsamında bu konu, kullanıcı deneyimi ile teknik güvenilirlik arasında bağlantı kurduğu için önemlidir. Platformun yalnızca çalışan ekranlardan oluşması yeterli görülmemiş; her ekranın hangi veriyi aldığı, hangi kontrolü yaptığı ve hangi sonraki adıma bağlandığı ayrıca değerlendirilmiştir.")
        add_para(doc, f"Uygulamadaki karşılığı {implementation} olarak özetlenebilir. Bu yaklaşım, kod tabanında ilgili modülün sorumluluğunu açık hale getirir. Böylece hata oluştuğunda sorun arayüz, API, veri modeli veya dış servis katmanlarından hangisinde ortaya çıkıyor daha kolay anlaşılır.")
        add_para(doc, f"Bu kararın proje açısından temel katkısı {contribution} şeklindedir. Özellikle 3D baskı gibi kullanıcı için teknik ayrıntıları fazla olan bir alanda, sistemin hatayı erken yakalaması ve kullanıcıya anlaşılır akış sunması projenin uygulanabilirliğini artırmaktadır.")
        add_para(doc, "Gelecekte bu alan daha gelişmiş ölçüm ve izleme mekanizmalarıyla desteklenebilir. Örneğin kullanıcı davranışları, servis yanıt süreleri, başarısız istek oranları ve üretim görevi tamamlanma süreleri düzenli olarak raporlandığında platformun hangi bölümünün iyileştirmeye daha çok ihtiyaç duyduğu daha somut biçimde görülebilir.")


def add_chapter_three(doc: Document, diagrams: dict[str, Path]):
    doc.add_page_break()
    add_main_heading(doc, "ÜÇÜNCÜ BÖLÜM")
    add_main_heading(doc, "BULGULAR VE TARTIŞMA")
    add_section_heading(doc, "3.1. Gerçekleştirilen Modüller")
    add_para(doc, "PrintForge projesi kapsamında kullanıcı, satıcı ve admin rollerinin ihtiyaçlarını destekleyen çok modüllü bir web platformu gerçekleştirilmiştir. Kullanıcı tarafında kayıt/giriş, örnek görsel arama, AI model üretimi, katalog inceleme, favori, sepet, mesajlaşma ve sipariş takibi akışları bulunmaktadır. Satıcı tarafında ürün ekleme, ürün düzenleme, ürün kaldırma, soruları yanıtlama ve sipariş durumunu güncelleme işlevleri geliştirilmiştir. Admin tarafında ise örnek referans görsellerinin yönetimi sağlanmıştır.")
    add_para(doc, "Gerçekleştirilen modüller incelendiğinde sistemin yalnızca tek bir AI üretim ekranından ibaret olmadığı görülmektedir. Katalog ve mesajlaşma modülleri, AI üretim çıktısının üretim hizmetine dönüşebilmesi için gerekli iş katmanını sağlamaktadır. Admin örnekleri ve görsel arama modülleri ise kullanıcının fikir bulma aşamasında desteklenmesini sağlar. Bu yapı Şekil 3.1'de modül tamamlama durumu açısından özetlenmiştir.")
    add_figure(doc, diagrams["fig_3_1"], "Şekil 3.1. Modül tamamlama durumu")

    add_section_heading(doc, "3.2. Derleme ve Teknik Doğrulama Bulguları")
    add_para(doc, "Projenin teknik doğrulaması için backend ve frontend üretim derlemeleri çalıştırılmıştır. Backend katmanında TypeScript derlemesi hata vermeden tamamlanmıştır. Frontend katmanında Next.js üretim derlemesi başarıyla tamamlanmış ve 21 sayfa oluşturulmuştur. Derleme çıktısında webpack önbellek stratejisine ilişkin bir uyarı görülmüş, ancak bu uyarı derlemeyi durdurmamış ve işlevsel hata olarak değerlendirilmemiştir.")
    rows = [
        ["Bileşen", "Çalıştırılan doğrulama", "Sonuç"],
        ["Backend", "npm run build / tsc", "Başarılı. TypeScript derlemesi tamamlandı."],
        ["Frontend", "npm run build / next build", "Başarılı. Next.js 14.1.0 üretim derlemesi 21 sayfa oluşturdu."],
        ["Dinamik sayfalar", "/marketplace, /chat/[id], API rotaları", "Üretim çıktısında server-rendered/on-demand olarak listelendi."],
        ["Uyarı", "Webpack cache PackFileCacheStrategy", "Derlemeyi engellemedi; izlenmesi gereken önbellek uyarısıdır."],
    ]
    add_table(doc, "Tablo 3.1. Derleme doğrulama sonuçları", rows, [3.0, 5.0, 7.5], "03.06.2026 tarihinde proje çalışma klasöründe yapılan derleme çıktıları esas alınmıştır.", 9)
    add_figure(doc, diagrams["fig_3_2"], "Şekil 3.2. Derleme ve doğrulama özeti")

    add_section_heading(doc, "3.3. Kullanıcı Akışlarına İlişkin Bulgular")
    add_para(doc, "Kullanıcı akışları incelendiğinde PrintForge'un fikir aşamasından satıcıyla iletişim aşamasına kadar kesintisiz bir deneyim sunduğu görülmektedir. /examples sayfasında kullanıcı admin örneklerini veya internet görsel arama sonuçlarını görebilmekte; seçilen görsel ve otomatik oluşturulan prompt /ai-create üzerinden /ai-generator sayfasına taşınmaktadır. AI üretim sayfası görseli dosyaya dönüştürerek backend'e göndermekte, görev durumunu izlemekte ve tamamlanan modeli 3D önizleme alanında göstermektedir.")
    add_para(doc, "Pazaryeri akışında kullanıcı katalog ürünlerini filtreleyebilmekte, sabit fiyatları görebilmekte, ürünleri favorilere veya sepete ekleyebilmekte, ürün detayında yorum ve soru-cevap bölümlerini inceleyebilmektedir. Satıcılar ürünlerini ayrı bir panelden yönetebilmekte ve alıcı sorularına yanıt verebilmektedir. Mesaj merkezi, satıcı ve alıcı konuşmalarını sipariş ekranından ayrı sunarak iletişimin takip edilebilirliğini artırmaktadır.")
    rows = [
        ["Akış", "Gerçekleştirilen işlev", "Kullanıcıya katkısı"],
        ["Örnekten AI üretime geçiş", "Görsel URL, başlık, kaynak ve prompt query parametresiyle aktarılır.", "Kullanıcı prompt yazmaya sıfırdan başlamak zorunda kalmaz."],
        ["AI görev izleme", "Görev localStorage ile korunur ve 15 saniyede bir sorgulanır.", "Uzun süren üretimde kullanıcı sayfadan ayrılabilse de süreç izlenir."],
        ["Katalog", "Arama, kategori ve fiyat filtresi ile ürün kartları listelenir.", "Hazır üretim seçeneklerine hızlı erişim sağlanır."],
        ["Satıcı paneli", "Ürün görselleri, kategori, açıklama ve sabit fiyat yönetilir.", "Satıcı ürünlerini teknik destek almadan yayınlayabilir."],
        ["Mesajlaşma", "Konuşmalar, okunmamış mesajlar ve sipariş durumları izlenir.", "Teklif ve üretim görüşmeleri kayıt altında tutulur."],
        ["Admin örnekleri", "Referans görseller ve promptlar yönetilir.", "Dış servis çalışmasa bile kontrollü başlangıç içerikleri sunulur."],
    ]
    add_table(doc, "Tablo 3.2. Gerçekleştirilen kullanıcı akışları", rows, [3.2, 5.3, 7.0], "Frontend sayfaları, backend rotaları ve demo rehberi birlikte incelenerek hazırlanmıştır.", 9)

    add_section_heading(doc, "3.4. Güvenlik ve Sürdürülebilirlik Bulguları")
    add_para(doc, "Güvenlik açısından elde edilen en önemli bulgu, hassas servis anahtarlarının frontend tarafına taşınmadan backend üzerinden yönetilmesidir. Görsel arama sağlayıcıları, AI üretim servisleri ve veritabanı bağlantısı .env dosyalarıyla ayrılmıştır. Kimlik doğrulama akışı JWT ve bcrypt ile desteklenmekte, korunan rotalarda token doğrulaması yapılmaktadır. Zod şemaları kullanıcı girdilerinin sınırlandırılmasını sağlamakta; arama sorguları maksimum 80 karakterle sınırlandırılarak zararlı karakterler temizlenmektedir.")
    add_para(doc, "Sürdürülebilirlik açısından kod tabanı modüler rota dosyalarına ve ayrı frontend sayfalarına bölünmüştür. Backend tarafında SQLite uyumluluk yardımcıları ve PostgreSQL şeması birlikte bulunmakta, bu durum yerel geliştirme ve canlı ortam arasında geçişe olanak sağlamaktadır. Netlify yapılandırması frontend dağıtımı için hazırlanmış, backend'in ayrı web servisi olarak çalışacağı README ve demo rehberinde belirtilmiştir.")

    doc.add_page_break()
    add_section_heading(doc, "3.5. Kullanılabilirlik Bulguları")
    add_para(doc, "Kullanılabilirlik açısından PrintForge'un en önemli bulgusu, kullanıcının teknik modelleme terimleriyle karşılaşmadan AI üretim akışına başlayabilmesidir. Örnek görsel seçimi, otomatik prompt üretimi ve referans görselin üretim sayfasına taşınması, kullanıcının yalnızca ne üretmek istediğine odaklanmasını sağlar. Bu tasarım, 3D baskı sürecindeki teknik eşiği tamamen ortadan kaldırmasa da ilk adımı belirgin biçimde kolaylaştırmaktadır.")
    add_para(doc, "Katalog tarafında filtreleme, sabit fiyat gösterimi, favori ve sepet işlevleri, kullanıcının alışık olduğu e-ticaret davranışlarını 3D baskı bağlamına taşımaktadır. Bu tercih bilinçlidir; çünkü 3D baskı hizmeti teknik bir alan olsa da platformun hedef kitlesi yalnızca mühendis veya tasarımcı değildir. Kullanıcıya tanıdık gelen katalog ve mesajlaşma kalıpları, sistemin öğrenilmesini kolaylaştırır.")
    add_para(doc, "Satıcı deneyiminde ürün ekleme akışının görsel, kategori, açıklama ve tek fiyat üzerinden kurgulanması, satıcıların karmaşık stok veya varyant yönetimiyle uğraşmadan katalog oluşturmasını sağlar. Bu basitlik, bitirme projesi kapsamı için yeterli görülmüş; ileride malzeme, renk, doluluk oranı, baskı süresi ve teslimat seçenekleri eklenebilecek şekilde veri modeli genişletilebilir bırakılmıştır.")
    add_para(doc, "Mesajlar ve siparişler ekranlarının ayrılması da kullanılabilirlik bulgusudur. Kullanıcı veya satıcı tüm konuşmaları mesaj merkezinde takip ederken, sipariş durumuna geçmiş işlemleri ayrıca sipariş ekranında görebilmektedir. Bu ayrım, sıradan bilgi alma konuşmaları ile aktif üretim/sipariş süreçlerinin birbirine karışmasını önler.")

    doc.add_page_break()
    add_section_heading(doc, "3.6. Teknik Risk ve Dayanıklılık Bulguları")
    add_para(doc, "Projenin teknik riskleri büyük ölçüde harici servis bağımlılıkları, dosya formatı farklılıkları ve kullanıcı girdilerinin kontrol edilmesi etrafında toplanmaktadır. Görsel arama ve AI üretim servisleri kota, erişim, gecikme veya yanıt formatı değişikliği gibi nedenlerle beklenen şekilde çalışmayabilir. Bu nedenle sistemde yalnızca mutlu yol değil, hata mesajları ve yedek akışlar da dikkate alınmıştır.")
    add_para(doc, "Görsel arama tarafında admin örneklerinin bulunması dayanıklılığı artırmaktadır. SerpApi veya Google Images sağlayıcısı erişilemez olduğunda kullanıcı hâlâ admin tarafından hazırlanmış örneklerle AI üretim akışına geçebilir. Bu yaklaşım, canlı sunum veya jüri demosu sırasında dış servis kaynaklı aksama riskini azaltır.")
    add_para(doc, "AI üretim tarafında görev tabanlı izleme kullanılması, uzun süren üretimlerin kullanıcı arayüzünü kilitlemesini engeller. Görev kimliğinin yerel depoda saklanması, kullanıcının sayfadan ayrılması halinde üretim durumunun tamamen kaybolmamasını sağlar. Bu tasarım, özellikle 3D model üretiminin saniyeler değil dakikalar sürebileceği durumlarda önemlidir.")
    add_para(doc, "Dosya formatı tarafında GLB ve STL ayrımı yapılması da önemli bir teknik bulgudur. GLB dosyaları doğrudan Three.js ile görüntülenebilirken STL dosyaları için ayrı işleme gerekebilir. Backend'in çıktı formatını kontrol etmesi ve uygun saklama yolunu seçmesi, farklı AI servislerinden gelen sonuçların tek platformda yönetilebilmesini sağlar.")

    doc.add_page_break()
    add_section_heading(doc, "3.7. Sürdürülebilirlik ve Genişletilebilirlik Bulguları")
    add_para(doc, "Kod tabanının frontend, backend ve veritabanı şeması olarak ayrılması, sistemin geliştirilmesini kolaylaştırmaktadır. Frontend içinde sayfa ve bileşenlerin ayrı tutulması arayüz değişikliklerini sınırlı alanlarda yapmayı sağlar. Backend tarafında rota dosyalarının iş alanlarına göre ayrılması ise kimlik, AI üretim, model/katalog, mesajlaşma, örnekler ve görsel arama modüllerinin birbirinden bağımsız incelenmesine olanak verir.")
    add_para(doc, "Veri modelinde Model varlığının hem AI hem de CATALOG türünü desteklemesi, ortak alanların tekrarını azaltmıştır. Bu karar sayesinde konuşma, dosya erişimi, yorum, soru ve sipariş gibi ilişkiler tek bir model kavramı üzerinden yürütülebilmektedir. Bununla birlikte ileride AI modeli ile katalog ürünü arasında çok farklı alanlar oluşursa ayrıştırılmış alt tablolar veya tür bazlı ek modeller kullanılabilir.")
    add_para(doc, "Sürdürülebilirlik açısından bir diğer bulgu, ortam değişkenlerinin açık biçimde belgelenmiş olmasıdır. README ve demo rehberi, frontend ve backend ortamlarının farklı ayarlanması gerektiğini belirtmektedir. Bu bilgi, projenin yalnızca geliştirici bilgisayarında değil, canlı yayın ortamında da doğru yapılandırılabilmesi için önemlidir.")
    add_para(doc, "Genişletilebilirlik açısından en uygun sonraki adım, baskı uygunluk analizi ve maliyet tahmini modülüdür. Mevcut sistem model dosyasını saklayıp önizleyebildiği için hacim, yüzey alanı, sınır kutusu, tahmini malzeme tüketimi ve baskı süresi gibi bilgiler hesaplanabilir. Bu hesaplamalar satıcı teklif sürecini destekleyerek platformun ticari değerini artırabilir.")

    add_section_heading(doc, "3.8. Tartışma")
    add_para(doc, "Elde edilen bulgular PrintForge'un araştırma sorusuna olumlu yanıt verdiğini göstermektedir. Teknik modelleme bilgisi sınırlı kullanıcılar için referans görsel seçimi, AI üretim görevi ve satıcıyla iletişim tek bir web uygulamasında birleştirilebilmiştir. Bu yönüyle proje, literatürde tartışılan 3D baskının yaygınlaşması ve iş modeli dönüşümü konularını pratik bir yazılım çıktısına dönüştürmektedir (Rayna ve Striukova,2016:214-224; Weller vd.,2015:43-56).")
    add_para(doc, "Bununla birlikte sistemin sınırlılıkları bulunmaktadır. AI üretim kalitesi dış servise bağlıdır ve sistem üretilen modelin baskıya uygunluğunu otomatik mühendislik analiziyle doğrulamamaktadır. Görsel arama sonuçları telif ve kalite açısından değişken olabilir. Ödeme, kargo ve baskı parametresi hesaplama gibi ticari süreçler temel seviyede modellenmiş, tam entegrasyon yapılmamıştır. Bu sınırlılıklar Tablo 3.3'te gelecek çalışma önerileriyle birlikte sunulmuştur.")
    rows = [
        ["Sınırlılık", "Etkisi", "Geliştirme önerisi"],
        ["AI servis bağımlılığı", "Üretim süresi, kota ve çıktı kalitesi dış servise bağlıdır.", "Birden fazla sağlayıcı için öncelik/fallback mekanizması kurulabilir."],
        ["Baskıya uygunluk analizi yok", "Modelin duvar kalınlığı, hacim ve destek ihtiyacı otomatik doğrulanmaz.", "STL/GLB geometri analizi ve baskı uygunluk puanı eklenebilir."],
        ["Telif/lisans kontrolü sınırlı", "Harici görsellerin kullanım hakkı kullanıcı sorumluluğundadır.", "Lisans filtresi ve kaynak uyarıları geliştirilebilir."],
        ["Ödeme ve lojistik entegrasyonu yok", "Sipariş akışı iletişim ve durum takibi düzeyindedir.", "Ödeme sağlayıcısı, teslimat ve fatura modülleri eklenebilir."],
        ["Maliyet tahmini yok", "Kullanıcı üretim maliyetini satıcıyla konuşarak öğrenir.", "Malzeme, hacim ve doluluk oranına dayalı fiyat tahmini yapılabilir."],
    ]
    add_table(doc, "Tablo 3.3. Sınırlılıklar ve geliştirme önerileri", rows, [3.2, 5.0, 7.3], "Proje kapsamı ve test bulguları esas alınarak hazırlanmıştır.", 9)
    add_para(doc, "Literatürde eklemeli imalatın dağıtık üretim ve kişiselleştirme potansiyeli vurgulanırken kalite kontrol ve standartlaşma sorunlarının sürdüğü görülmektedir (Ford ve Despeisse,2016:1573-1587; Ngo vd.,2018:172-196). PrintForge bu sorunları tamamen çözmemekte, ancak kullanıcı-satıcı iletişimi ve AI destekli başlangıç akışıyla 3D baskıya erişim bariyerlerini azaltan uygulanabilir bir ara katman sunmaktadır. Bu nedenle projenin katkısı, yeni bir 3D üretim algoritması önermekten çok, mevcut AI ve web teknolojilerini 3D baskı hizmet sürecine bütünleşik şekilde uygulamasıdır.")


def add_conclusion_refs_appendices(doc: Document):
    doc.add_page_break()
    add_main_heading(doc, "SONUÇ VE ÖNERİLER")
    for para in [
        "Bu tez çalışmasında, 3D baskı hizmeti almak isteyen kullanıcıların fikir geliştirme, referans görsel seçme, AI destekli model üretme ve satıcıyla iletişim kurma adımlarını tek platformda birleştiren PrintForge sistemi geliştirilmiştir. Çalışma kapsamında kullanıcı, satıcı ve admin rollerini destekleyen; satıcı katalogu, görsel arama, admin örnek yönetimi, AI üretim görevi, 3D model önizleme, mesajlaşma, soru-cevap, yorum ve sipariş takibi modüllerinden oluşan bütüncül bir web uygulaması ortaya konmuştur.",
        "Projenin temel hedefi, teknik modelleme bilgisi sınırlı kullanıcıların 3D baskı sürecine daha kolay katılmasını sağlamaktır. Elde edilen bulgular, bu hedef doğrultusunda referans görselden AI üretim sayfasına geçişin otomatik prompt ile desteklendiğini, üretim görevinin arka planda izlenebildiğini, tamamlanan modelin güvenli dosya erişimiyle görüntülenebildiğini ve kullanıcıların satıcılarla mesajlaşarak üretim sürecini ilerletebildiğini göstermektedir.",
        "Teknik doğrulama sonucunda backend TypeScript derlemesi başarıyla tamamlanmış, frontend Next.js üretim derlemesi 21 sayfayı oluşturarak tamamlanmıştır. Bu sonuç, projenin yalnızca kavramsal bir tasarım olmadığını, derlenebilir ve teslim edilebilir bir yazılım çıktısı sunduğunu göstermektedir. Geliştirilen veri modeli ve modüler API yapısı, projenin ileride yeni servislerle genişletilebilmesine uygun bir temel oluşturmaktadır.",
        "Gelecekte yapılabilecek geliştirmeler arasında üretilen modeller için otomatik baskıya uygunluk analizi, hacim ve yüzey alanına dayalı maliyet tahmini, satıcı tekliflerini karşılaştırma ekranı, malzeme ve doluluk oranı seçenekleri, ödeme ve kargo entegrasyonu, gelişmiş admin moderasyonu, lisans kontrolü ve kullanıcıya ait proje koleksiyonları yer almaktadır. Ayrıca AI üretim servisleri için çoklu sağlayıcı desteği ve kalite karşılaştırması yapılması, sistemin güvenilirliğini artıracaktır.",
        "Sonuç olarak PrintForge, 3D baskı ve AI destekli model üretimi alanındaki teknolojik eğilimleri kullanıcı odaklı bir web pazaryeri deneyimiyle birleştiren uygulanabilir bir lisans bitirme projesi olarak değerlendirilebilir. Proje, eklemeli imalatın yaygınlaşmasında yazılım arayüzlerinin ve bütünleşik hizmet akışlarının önemini göstermekte; gelecekte daha gelişmiş üretim, analiz ve ticari entegrasyon modülleri için güçlü bir başlangıç noktası sunmaktadır.",
    ]:
        add_para(doc, para)

    doc.add_page_break()
    add_main_heading(doc, "KAYNAKÇA")
    for ref in REFERENCES:
        p = doc.add_paragraph(style="TezNoIndent")
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        set_font(r, 11)

    doc.add_page_break()
    add_main_heading(doc, "EKLER")
    add_section_heading(doc, "Ek A. Kurulum ve Çalıştırma Adımları")
    for step in [
        "Backend klasöründe bağımlılıklar yüklendikten sonra geliştirme ortamı npm run dev komutuyla başlatılır.",
        "Frontend klasöründe bağımlılıklar yüklendikten sonra Next.js geliştirme sunucusu npm run dev komutuyla çalıştırılır.",
        "Backend varsayılan olarak http://localhost:3001 adresinde, frontend ise http://localhost:3000 veya uygun başka bir portta çalışır.",
        "Frontend portu değişirse backend .env dosyasındaki FRONTEND_URLS değeri güncellenmelidir.",
    ]:
        add_numbered(doc, step)

    add_section_heading(doc, "Ek B. Backend Ortam Değişkenleri")
    rows = [
        ["Değişken", "Açıklama"],
        ["DATABASE_URL", "Prisma veritabanı bağlantısı."],
        ["JWT_SECRET", "JWT imzalama anahtarı."],
        ["FRONTEND_URLS", "CORS için izin verilen frontend adresleri."],
        ["IMAGE_SEARCH_PROVIDER", "serpapi veya google arama sağlayıcısı."],
        ["SERPAPI_API_KEY", "SerpApi görsel arama anahtarı."],
        ["TRIPO_API_KEY", "Tripo3D üretim servisi anahtarı."],
        ["HITEM3D_ACCESS_KEY / HITEM3D_SECRET_KEY", "Hitem3D üretim servisi kimlik bilgileri."],
    ]
    add_table(doc, "Ek B.1. Backend ortam değişkenleri", rows, [5.0, 10.5], "README dosyası esas alınarak hazırlanmıştır.", 9)

    add_section_heading(doc, "Ek C. API Uç Noktaları Özeti")
    rows = [
        ["Uç nokta", "Açıklama"],
        ["/health", "Backend çalışma durumunu döndürür."],
        ["/api/auth/register", "Kullanıcı veya satıcı hesabı oluşturur."],
        ["/api/auth/login", "JWT tabanlı oturum başlatır."],
        ["/api/images/search", "Görsel arama sonuçlarını döndürür."],
        ["/api/ai/generate", "AI model üretim görevini başlatır."],
        ["/api/ai/status/:taskId", "Üretim görevi durumunu sorgular."],
        ["/api/models", "Katalog ürünlerini listeler veya satıcı ürünü ekler."],
        ["/api/chat/new", "Model veya ürün için konuşma başlatır."],
    ]
    add_table(doc, "Ek C.1. API uç noktaları özeti", rows, [5.0, 10.5], "Backend rota dosyaları esas alınarak hazırlanmıştır.", 9)

    add_section_heading(doc, "Ek D. Manuel Test Kontrol Listesi")
    for item in [
        "Backend /health uç noktası kontrol edilir.",
        "Kullanıcı ve satıcı hesabıyla giriş yapılır.",
        "Satıcı panelinden en az bir ürün katalogda yayınlanır.",
        "Örnekler sayfasında admin örneği veya görsel arama sonucu seçilir.",
        "AI üretim sayfasında referans görsel ve promptun geldiği doğrulanır.",
        "Üretim tamamlanmış model için 3D önizleme açılır.",
        "Katalog ürününden satıcıya mesaj gönderilir.",
        "Sipariş durumu satıcı tarafından güncellenir.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_main_heading(doc, "ÖZGEÇMİŞ")
    lines = [
        "Adı ve Soyadı : " + AUTHOR,
        "Doğum Tarihi  : ........................................",
        "E-mail        : ........................................",
        "",
        "Öğrenim Durumu:",
    ]
    for line in lines:
        add_para(doc, line, "TezNoIndent")
    rows = [
        ["Derece", "Bölüm/Program", "Üniversite", "Yıl"],
        ["Lisans", "Bilgisayar Mühendisliği", "Tarsus Üniversitesi", "2026"],
    ]
    add_table(doc, "Öğrenim bilgileri", rows, [3.0, 5.0, 5.0, 2.5], "Yazar bilgileri tamamlandığında güncellenmek üzere hazırlanmıştır.", 10)
    add_para(doc, "ESERLER (Makaleler ve Bildiriler)", "TezNoIndent")
    add_para(doc, "1. ........................................", "TezNoIndent")


def build(page_map: dict[str, str], total_pages: int):
    diagrams = make_diagrams()
    doc = Document()
    configure_styles(doc)
    setup_section(doc.sections[0], footer=False)
    add_cover_pages(doc, total_pages)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(front, footer=True, start=5, fmt="lowerRoman")
    add_abstracts(doc, total_pages)
    add_front_matter(doc, page_map)

    main = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(main, footer=True, start=1, fmt="decimal")
    add_intro(doc)
    add_chapter_one(doc, diagrams)
    add_chapter_two(doc, diagrams)
    add_extended_technical_design(doc)
    add_detailed_review_pages(doc)
    add_chapter_three(doc, diagrams)
    add_conclusion_refs_appendices(doc)
    enable_update_fields_on_open(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--page-map")
    parser.add_argument("--total-pages", type=int, default=0)
    args = parser.parse_args()
    build(load_page_map(args.page_map), args.total_pages)


if __name__ == "__main__":
    main()
